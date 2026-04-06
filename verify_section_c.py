#!/usr/bin/env python3
"""
Quick Verification - Check all 3 lessons for Section C
"""
from pathlib import Path
from docx import Document


def check_section_c(filepath):
    """Quick check if Section C exists in a DOCX file"""
    
    if not Path(filepath).exists():
        return {
            'exists': False,
            'error': 'File not found'
        }
    
    try:
        doc = Document(filepath)
        
        # Get all text
        all_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Check for Section C
        has_section_c_heading = 'C. Lesson Implementation Framework' in all_text or \
                               'C. LESSON IMPLEMENTATION FRAMEWORK' in all_text
        
        has_implementation = 'Implementation Framework' in all_text
        
        # Count tables
        table_count = len(doc.tables)
        
        # Look for 5-column table with 7+ rows
        section_c_table = None
        for table in doc.tables:
            if len(table.columns) == 5 and len(table.rows) >= 7:
                section_c_table = {
                    'rows': len(table.rows),
                    'cols': len(table.columns)
                }
                break
        
        # Count ROW markers
        row_markers = []
        for i in range(1, 10):
            if f'ROW {i}' in all_text:
                row_markers.append(i)
        
        return {
            'exists': True,
            'has_section_c_heading': has_section_c_heading,
            'has_implementation_text': has_implementation,
            'table_count': table_count,
            'section_c_table': section_c_table,
            'row_markers': row_markers,
            'all_good': has_section_c_heading and section_c_table is not None and len(row_markers) >= 7
        }
        
    except Exception as e:
        return {
            'exists': True,
            'error': str(e)
        }


def main():
    """Check all 3 lesson files"""
    
    print("\n" + "="*70)
    print("VERIFICATION: CHECKING SECTION C IN ALL 3 LESSONS")
    print("="*70 + "\n")
    
    lesson_files = [
        ('Physics', 'data/outputs/docx/Lesson_001_Introduction_to_Current_Electricity.docx'),
        ('Chemistry', 'data/outputs/docx/Lesson_001_Introduction_to_Atoms_and_Elements.docx'),
        ('Biology', 'data/outputs/docx/Lesson_001_Introduction_to_Cells.docx')
    ]
    
    all_passed = True
    
    for subject, filepath in lesson_files:
        print(f"{subject} Lesson:")
        print(f"  File: {Path(filepath).name}")
        
        result = check_section_c(filepath)
        
        if not result.get('exists'):
            print(f"  ✗ ERROR: {result.get('error', 'Unknown error')}")
            all_passed = False
            print()
            continue
        
        if 'error' in result:
            print(f"  ✗ ERROR: {result['error']}")
            all_passed = False
            print()
            continue
        
        # Check results
        if result.get('has_section_c_heading'):
            print(f"  ✓ Section C heading present")
        else:
            print(f"  ✗ Section C heading MISSING")
            all_passed = False
        
        if result.get('section_c_table'):
            table = result['section_c_table']
            print(f"  ✓ Section C table found: {table['rows']} rows × {table['cols']} columns")
        else:
            print(f"  ✗ No 5-column table with 7+ rows found")
            all_passed = False
        
        row_markers = result.get('row_markers', [])
        if len(row_markers) >= 7:
            print(f"  ✓ ROW markers found: {row_markers}")
        else:
            print(f"  ⚠ ROW markers: {row_markers} (expected at least 7)")
        
        if result.get('all_good'):
            print(f"  ✅ ALL CHECKS PASSED")
        else:
            print(f"  ⚠ Some checks failed")
            all_passed = False
        
        print()
    
    print("="*70)
    if all_passed:
        print("🎉 SUCCESS! All 3 lessons have Section C properly formatted!")
    else:
        print("⚠ WARNING: Some lessons may have issues. Check details above.")
    print("="*70 + "\n")


if __name__ == "__main__":
    main()
