#!/usr/bin/env python3
"""
Analyze existing DOCX file to see what was actually generated
"""
import sys
from pathlib import Path
from docx import Document

def analyze_docx(filepath):
    """Analyze a DOCX file's structure"""
    
    print("\n" + "="*70)
    print(f"ANALYZING: {filepath}")
    print("="*70 + "\n")
    
    if not Path(filepath).exists():
        print(f"ERROR: File not found: {filepath}")
        return
    
    doc = Document(filepath)
    
    # Extract all text
    all_text = []
    for para in doc.paragraphs:
        all_text.append(para.text)
    
    full_text = '\n'.join(all_text)
    
    # Save full text
    debug_file = filepath.replace('.docx', '_FULL_TEXT.txt')
    with open(debug_file, 'w') as f:
        f.write(full_text)
    
    print(f"Full text extracted to: {debug_file}")
    print(f"Total length: {len(full_text)} characters")
    print()
    
    # Find headings
    print("HEADINGS FOUND:")
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            print(f"  {para.style.name}: {para.text}")
    
    print()
    
    # Check for sections
    print("SECTION MARKERS:")
    for section in ['A', 'B', 'C', 'D', 'E']:
        if f'Section {section}' in full_text or f'SECTION {section}' in full_text:
            print(f"  ✓ Section {section} mentioned")
        else:
            print(f"  ✗ Section {section} not found")
    
    print()
    
    # Count tables
    print(f"TABLES: {len(doc.tables)}")
    for i, table in enumerate(doc.tables, 1):
        print(f"  Table {i}: {len(table.rows)} rows × {len(table.columns)} columns")
        if len(table.columns) == 5 and len(table.rows) >= 7:
            print(f"    → This looks like Section C! (5 cols, {len(table.rows)} rows)")
    
    print()
    
    # Check for "Implementation Framework"
    if 'Implementation Framework' in full_text or 'IMPLEMENTATION FRAMEWORK' in full_text:
        print("✓ 'Implementation Framework' text found")
    else:
        print("✗ 'Implementation Framework' text not found")
    
    # Check for row markers
    row_markers = []
    for i in range(1, 10):
        if f'ROW {i}' in full_text:
            row_markers.append(i)
    
    if row_markers:
        print(f"✓ Found ROW markers: {row_markers}")
    else:
        print("✗ No ROW markers found")
    
    print("\n" + "="*70)


def main():
    """Main function"""
    
    # Check most recent Chemistry file
    chemistry_files = list(Path('data/outputs/docx').glob('*Atoms*.docx'))
    
    if not chemistry_files:
        print("No Chemistry files found in data/outputs/docx/")
        print("\nUsage: python3 analyze_docx.py <filepath>")
        return
    
    # Analyze most recent
    most_recent = max(chemistry_files, key=lambda p: p.stat().st_mtime)
    analyze_docx(str(most_recent))


if __name__ == "__main__":
    if len(sys.argv) > 1:
        analyze_docx(sys.argv[1])
    else:
        main()
