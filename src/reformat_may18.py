#!/usr/bin/env python3
"""
reformat_may18.py — CBE Document Generator for May 18 Scheme of Work Templates

Generates 12 output files (3 per subject × 4 subjects) from source planning docs.
Generated/AI-suggested content is rendered in italic in all outputs.
Source content (from the Scheme of Work docs) is rendered in normal font.

Subjects:
  biology    → Biology 10 — Sub-Strand 2.2: Transport System in Plants (22 lessons)
  physics    → Physics 10 — Sub-Strand 1.5: Moments and Equilibrium (6 lessons)
  chemistry  → Chemistry 10 — Sub-Strand 1.5: The Periodicity (7 lessons)
  maths      → Mathematics 10 — Sub-Strand 2.1: Similarity and Enlargement (12 lessons)
"""

import sys
import re
from pathlib import Path

# ── Import shared builders from reformat_seavuria ─────────────────────────────
sys.path.insert(0, str(Path(__file__).parent))
import reformat_seavuria as _sv

# Re-export everything we need
_cell_para        = _sv._cell_para
_cell_para_lines  = _sv._cell_para_lines
_cell_para_auto   = _sv._cell_para_auto
_build_table_A    = _sv._build_table_A
_build_table_B    = _sv._build_table_B
_build_table_C_period = _sv._build_table_C_period
_build_table_D    = _sv._build_table_D
_build_table_E    = _sv._build_table_E
_build_table0_overview = _sv._build_table0_overview
_build_section_table   = _sv._build_section_table
_set_page_landscape    = _sv._set_page_landscape
_doc_title        = _sv._doc_title
_doc_subtitle     = _sv._doc_subtitle
_add_page_break   = _sv._add_page_break
_tbl_no_spacing   = _sv._tbl_no_spacing
_new_table        = _sv._new_table
_merge_row        = _sv._merge_row
_shade            = _sv._shade
_col_widths       = _sv._col_widths
_text_to_rich     = _sv._text_to_rich
_prose_to_bullets = _sv._prose_to_bullets
_enrich_content   = _sv._enrich_content
parse_framework_sections = _sv.parse_framework_sections
parse_doc_sections       = _sv.parse_doc_sections

# Color constants
C_NAVY     = _sv.C_NAVY
C_TEAL     = _sv.C_TEAL
C_MED_BLUE = _sv.C_MED_BLUE
C_PURPLE   = _sv.C_PURPLE
C_ORANGE   = _sv.C_ORANGE
C_LT_BLUE  = _sv.C_LT_BLUE
C_TEAL_LT  = _sv.C_TEAL_LT
C_GREEN_LT = _sv.C_GREEN_LT
C_PURPLE_LT = _sv.C_PURPLE_LT
C_ORANGE_LT = _sv.C_ORANGE_LT
C_LT_GRAY  = _sv.C_LT_GRAY
C_WHITE    = _sv.C_WHITE
C_MED_BLUE = _sv.C_MED_BLUE

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent.parent
SRC_DIR  = BASE_DIR / "data" / "raw" / "CBE_Templates_May_18"
OUT_DIR  = BASE_DIR / "data" / "outputs" / "May18_Reformatted"
OUT_DIR.mkdir(parents=True, exist_ok=True)

SOURCE = {
    "biology":   SRC_DIR / "BIOLOGY TRANSPORT SYSTEM IN PLANT  SCHEME OF WORK  (1).docx",
    "physics":   SRC_DIR / "Grade 10 1.5 Moments and equilibrium (1).docx",
    "chemistry": SRC_DIR / "The periodicity.docx",
    "maths":     SRC_DIR / "maths _ Geometry and measurement;Similarity and enlargement.docx",
}

OUTPUTS = {
    "biology": {
        "lesson_seq":  OUT_DIR / "Biology_10_SubStrand2.2_TransportSystemInPlants_CBE_LessonSequence.docx",
        "final_exp":   OUT_DIR / "Biology_10_SubStrand2.2_TransportSystemInPlants_FinalExplanation.docx",
        "summary_tbl": OUT_DIR / "Biology_10_SubStrand2.2_TransportSystemInPlants_SummaryTable.docx",
    },
    "physics": {
        "lesson_seq":  OUT_DIR / "Physics_10_SubStrand1.5_MomentsEquilibrium_CBE_LessonSequence.docx",
        "final_exp":   OUT_DIR / "Physics_10_SubStrand1.5_MomentsEquilibrium_FinalExplanation.docx",
        "summary_tbl": OUT_DIR / "Physics_10_SubStrand1.5_MomentsEquilibrium_SummaryTable.docx",
    },
    "chemistry": {
        "lesson_seq":  OUT_DIR / "Chemistry_10_SubStrand1.5_Periodicity_CBE_LessonSequence.docx",
        "final_exp":   OUT_DIR / "Chemistry_10_SubStrand1.5_Periodicity_FinalExplanation.docx",
        "summary_tbl": OUT_DIR / "Chemistry_10_SubStrand1.5_Periodicity_SummaryTable.docx",
    },
    "maths": {
        "lesson_seq":  OUT_DIR / "Mathematics_10_SubStrand2.1_SimilarityEnlargement_CBE_LessonSequence.docx",
        "final_exp":   OUT_DIR / "Mathematics_10_SubStrand2.1_SimilarityEnlargement_FinalExplanation.docx",
        "summary_tbl": OUT_DIR / "Mathematics_10_SubStrand2.1_SimilarityEnlargement_SummaryTable.docx",
    },
}

# ── Helper: mark content as AI-generated (renders italic) ────────────────────

def _G(text: str) -> list:
    """Convert text to gen_* rich-content tuples (renders italic in output)."""
    return [("gen_" + k, v) for k, v in _text_to_rich(text)]


def _gc(cell, text: str, size_pt: int = 9):
    """Write AI-generated text into a cell (italic)."""
    _cell_para_lines(cell, _G(text), size_pt=size_pt)


# ── Phase row helper ──────────────────────────────────────────────────────────

def _pr(learner, resource, teacher, sensemaking, assessment):
    """5-tuple for one Table C phase row."""
    return (learner, resource, teacher, sensemaking, assessment)


# ═══════════════════════════════════════════════════════════════════════════════
# META — Biology
# ═══════════════════════════════════════════════════════════════════════════════

_BIO_SLO_K = (
    "a) Relate structures of the plant transport system to their functions in plants\n"
    "b) Illustrate the arrangement of vascular tissues in monocotyledonous and dicotyledonous plants\n"
    "c) Demonstrate the uptake of water and mineral salts from roots to leaves\n"
    "d) Demonstrate factors that affect the rate of transpiration in plants\n"
    "e) Describe the translocation of manufactured food in plants"
)
_BIO_SLO_S = (
    "- Conduct investigations on plant transport, transpiration and translocation\n"
    "- Draw and label diagrams of plant transport structures\n"
    "- Collect, record and interpret experimental data\n"
    "- Apply knowledge of transport to real farming challenges"
)
_BIO_SLO_A = (
    "- Appreciate the importance of plant transport systems to food security and farming\n"
    "- Show curiosity and care when handling living plant materials\n"
    "- Value teamwork and honest recording in investigations"
)

def _bio_lesson(num, title, inquiry, know, skills, attit, purpose, materials, safety, p1, p2=None):
    return {
        "number": num, "title": title,
        "inquiry_question": inquiry,
        "slo_knowledge": know, "slo_skills": skills, "slo_attitudes": attit,
        "overview_purpose": purpose,
        "materials": materials, "safety": safety,
        "period1_heading": "Period 1 (40 minutes)",
        "period1_table": p1,
        "period2_heading": "Period 2 (40 minutes)",
        "period2_table": p2 if p2 else [],
        "reflections": [],
    }

_BIO_LESSONS = [
    _bio_lesson(
        1, "Introduction to the Phenomenon",
        "Why do some crops like spinach wilt quickly on hot windy days while mango trees stay green?",  # SOURCE
        "- Recall that plants contain internal structures involved in water movement\n- Identify the unit driving question",
        "- Observe and record plant wilting; draw an initial model of water movement inside a plant",
        "- Show curiosity about how plants respond to their environment",
        "Launches the unit phenomenon. Students encounter wilting crops, generate questions, and produce initial models of water movement inside plants.",
        "Wilted spinach or sukuma wiki; fresh mango branch; Driving Question Board poster; blank A4 paper",
        "Handle live plant material with care; wash hands after touching soil.",
        [
            _pr("Observe wilted spinach alongside fresh mango branch; record observations",
                "Live plant samples or time-lapse video of wilting and recovery",
                "Show samples/video without explanation; ask 'What do you notice? What are you wondering?'",
                "Think–Pair–Share: individual → partner → whole-class share-out",
                "Exit ticket: list 2 observations and 1 question about wilting plants"),
            _pr("Draw 'what is happening inside the plant?' diagram showing roots, stem, leaves and water arrows",
                "Blank A4 paper; large class diagram for reference",
                "Prompt: 'Draw what you think is happening inside the plant when it wilts and when it recovers'",
                "Pair-share diagrams; identify what students agree and disagree on",
                "Collect initial model diagrams for formative reference"),
            _pr("Contribute questions to the class Driving Question Board",
                "Driving Question Board (poster or whiteboard section); sticky notes",
                "Facilitate DQB creation; group student questions by theme",
                "Model 'I notice … I wonder …' sentence stems",
                "Count and categorise student questions to inform lesson sequence"),
            _pr("Discuss: what do we already know about plant structure?",
                "Textbook diagrams or large poster of a whole plant",
                "Elicit prior knowledge; accept all ideas without correction at this stage",
                "Record prior knowledge on a KWL chart (Know–Want to know–Learned)",
                "Oral probe: 'Name one structure inside a plant and guess its job'"),
            _pr("Draw 'My Plant Model — Start of Unit': roots, stem, leaves, water movement arrows",
                "Model template or blank paper",
                "Explain that models will be updated each lesson as understanding grows",
                "Students share initial models with partners; note what is uncertain",
                "Collect models; note common misconceptions for future lessons"),
        ],
    ),
    _bio_lesson(
        2, "What Do Plants Need to Survive?",
        "What substances do plants need to survive, and how do they obtain them?",
        "- List the substances plants need: water, mineral salts, carbon dioxide, light\n- Explain why each substance is essential",
        "- Sort plant needs by category; link each need to a transport process",
        "- Appreciate that plants, like humans, have basic survival needs",
        "Establishes prior knowledge of plant needs and creates the link to why a transport system is necessary.",
        "Cards with plant needs written on them; sorting mats; textbook or chart",
        "No specific hazards.",
        [
            _pr("Brainstorm: what do plants need to survive? Record ideas individually then share",
                "Blank cards; markers for brainstorm",
                "Pose question; accept all responses; record on board without evaluation",
                "Think–Pair–Share; create class mind-map",
                "Quick class poll: thumbs up/down for each suggested need"),
            _pr("Sort plant needs into categories: from soil, from air, from sunlight",
                "Sorting cards and sorting mat",
                "Facilitate sorting activity; ask 'How does the plant get each one?'",
                "Group discussion; justify sorting decisions",
                "Exit ticket: 'Draw an arrow showing where water enters a plant'"),
            _pr("Link each need to a transport process: water via roots, minerals via roots, sugars via phloem",
                "Simple diagram of plant showing roots, stem, leaves",
                "Guided discussion: 'If plants need water at every cell, how does it travel from soil to leaf?'",
                "Connect to driving question: 'This is why plants have a transport system'",
                "Oral check: can students identify which structure transports water vs. food?"),
            _pr("Update Driving Question Board: add question 'How do plants get water from the soil?'",
                "Driving Question Board",
                "Guide students to connect plant needs to the phenomenon (wilting = lack of water)",
                "Re-read driving question together and discuss what we need to find out next",
                "Check that DQB links to upcoming lessons"),
            _pr("Update plant model: add arrows showing where water, minerals and CO₂ enter",
                "Initial model from Lesson 1",
                "Model how to add labels and arrows to an existing diagram",
                "Pair discussion: 'What did you change in your model?'",
                "Observe model updates; note which students include mineral salts vs. only water"),
        ],
    ),
    _bio_lesson(
        3, "Overview of Plant Transport System",
        "What are xylem and phloem, and what does each transport?",
        "- Name xylem and phloem as the two main transport vessels\n- State that xylem transports water and minerals upward; phloem transports sugars",
        "- Label xylem and phloem on a whole-plant diagram\n- Distinguish between the two vessels by function",
        "- Show interest in how microscopic structures enable a whole plant to function",
        "Introduces the two main transport vessels by name and function, providing a framework for the detailed lessons that follow.",
        "Large whole-plant diagram or poster; coloured pencils (blue for xylem, green for phloem)",
        "No specific hazards.",
        [
            _pr("Examine a large labelled diagram of a whole plant; identify xylem and phloem by colour",
                "Large poster or projected diagram of plant cross-section",
                "Introduce xylem and phloem as 'transport tubes'; avoid full structural detail at this stage",
                "Analogy: xylem = water pipes, phloem = sugar delivery system",
                "Oral check: 'Which vessel carries water? Which carries food?'"),
            _pr("Colour-code a blank plant diagram: blue for xylem pathway, green for phloem pathway",
                "Blank plant diagrams; blue and green pencils/pens",
                "Circulate; ensure students trace the correct continuous pathway",
                "Think-aloud: trace a water molecule from root hair to stomata",
                "Collect diagrams; check that pathways are correctly direction-labelled"),
            _pr("Discuss: why does a plant need two separate transport systems?",
                "Completed colour-coded diagrams",
                "Pose question; guide discussion without giving answer; accept hypotheses",
                "Connect to human circulatory system as an analogy (arteries/veins)",
                "Exit ticket: 'Give one reason why xylem and phloem must be separate'"),
            _pr("Update DQB: 'What is the difference between xylem and phloem?'",
                "Driving Question Board",
                "Add teacher-generated question if students have not raised it",
                "Revisit phenomenon: 'When spinach wilts, which vessel is most affected?'",
                "Students predict the answer; record predictions to revisit later"),
            _pr("Update plant model: label xylem and phloem in correct positions",
                "Running plant model diagram",
                "Emphasise: 'Your model should now show two distinct pathways'",
                "Compare Lesson 1 model vs Lesson 3 model; discuss what changed",
                "Note which students now include directionality (up vs. up and down)"),
        ],
    ),
    _bio_lesson(
        4, "Root Structure and Function",
        "How does the structure of root hair cells help plants absorb water and minerals?",
        "- Describe the structure of a root hair cell (large surface area, thin wall, vacuole)\n- Explain how root hairs are adapted for absorption (Obj a)",
        "- Observe real roots or a diagram; identify root hair zone\n- Draw and label a root hair cell",
        "- Appreciate the link between microscopic cell structure and whole-plant function",
        "Examines root structure in detail, establishing the site and mechanism of water entry — essential foundation for Lesson 13.",
        "Germinated bean seeds with visible roots; hand lens or microscope; root diagram",
        "Wash hands after handling soil or roots; handle seeds gently.",
        [
            _pr("Observe germinated bean roots with a hand lens; record observations of root structure",
                "Germinated bean seeds; hand lens; diagram of root tip and root hair zone",
                "Distribute seeds; guide observation: 'Look at the fuzzy zone near the tip — what do you see?'",
                "Think-aloud: 'Each fuzzy strand is a root hair cell — one single cell'",
                "Quick sketch of root with annotations"),
            _pr("Examine diagram of root hair cell; identify: cell wall, cell membrane, large vacuole, cytoplasm",
                "Labelled diagram of root hair cell",
                "Use guided labelling: display unlabelled diagram, students suggest labels",
                "Connect structure to function: 'Why is the vacuole large? Why is the wall thin?'",
                "Exit ticket: list 3 adaptations of root hair cells and their functions"),
            _pr("Calculate: if 1 cm of root has 100 root hairs, how many per metre? Discuss why this matters",
                "Calculators or mental maths",
                "Guide calculation; emphasise that large surface area increases absorption rate",
                "Real-world link: 'After rain, why do plants recover quickly?'",
                "Students explain in writing why root hairs increase absorption efficiency"),
            _pr("Update DQB: 'How does water get from the soil into the root?'",
                "Driving Question Board",
                "Connect today's learning to the phenomenon: wilting is loss of water from cells",
                "Foreshadow Lesson 13: 'We will see exactly how water crosses the root hair membrane'",
                "Students vote on the most important question still unanswered"),
            _pr("Update plant model: add root hair cells at the base of the root with arrows showing water entry",
                "Running plant model diagram",
                "Model how to zoom in on part of a diagram to show cellular detail",
                "Discuss: 'Your model is getting more accurate — what else do you need to show?'",
                "Compare with partner's model; share one improvement idea"),
        ],
    ),
    _bio_lesson(
        5, "Stem Structure",
        "What does the inside of a stem look like, and how is it organised for transport?",
        "- Describe the internal structure of a herbaceous stem\n- Identify epidermis, cortex, vascular bundle (xylem + phloem), pith (Obj a)",
        "- Observe and sketch a stem cross-section\n- Identify vascular bundles in a real or prepared stem",
        "- Show curiosity about the hidden internal organisation of familiar plants",
        "Reveals the internal arrangement of a stem, connecting macroscopic structure to the transport functions already introduced.",
        "Fresh celery stalks or young bean stems; sharp blade (teacher-only); hand lens; red food dye solution",
        "Only teacher cuts with the blade; students handle pre-cut sections only.",
        [
            _pr("Predict: 'Draw what you think the inside of a stem looks like if you cut it across'",
                "Blank paper; pencils",
                "Do NOT show stem cross-section yet; let students predict based on prior lessons",
                "Think-aloud: 'I expect to see … because …'",
                "Collect predictions; note whether students include xylem/phloem from Lesson 3"),
            _pr("Observe pre-cut stem cross-sections with a hand lens; identify the visible 'dots' (vascular bundles)",
                "Pre-cut celery or bean stem sections; hand lens",
                "Circulate; guide students to find the ring of vascular bundles",
                "Compare prediction with observation: 'What is the same? What is different?'",
                "Annotated sketch of observed cross-section"),
            _pr("Compare prediction diagrams with actual cross-section diagram; label epidermis, vascular bundle, pith",
                "Labelled diagram of herbaceous stem cross-section",
                "Guided labelling activity",
                "Connect to previous lessons: 'Which part of the vascular bundle is xylem? Which is phloem?'",
                "Exit ticket: draw and label a stem cross-section from memory"),
            _pr("Dye demonstration: place celery stalk in red food dye; predict where the dye will travel",
                "Celery stalk in red dye solution (set up in advance); pre-cut sections showing dye",
                "Show cross-section stained with dye; ask 'Which tissue absorbed the dye and why?'",
                "Connect to xylem structure: hollow dead cells ideal for water flow",
                "Students record observation and explain using xylem function knowledge"),
            _pr("Update plant model: add clear labels for xylem and phloem within the stem cross-section",
                "Running plant model diagram",
                "Compare model progression from Lesson 1 to Lesson 5",
                "Pair discussion: 'How has your understanding of the stem changed?'",
                "Model update check: does diagram show xylem inside phloem (correct for dicot)?"),
        ],
    ),
    _bio_lesson(
        6, "Leaf Structure and Transport Role",
        "How do the veins and stomata of a leaf connect to the plant's transport system?",
        "- Identify veins (vascular bundles), stomata, guard cells and mesophyll in a leaf\n- Explain how leaf structure supports both transport and gas exchange (Obj a)",
        "- Observe a leaf, locate veins; use a microscope or diagram to identify stomata and guard cells",
        "- Appreciate how multiple structures work together to support the life functions of a leaf",
        "Completes the organ-level structural overview (roots, stem, leaf) and connects leaf structure to both water transport and transpiration.",
        "Fresh leaves; hand lens; diagram of leaf cross-section; microscope or prepared slides (if available)",
        "Handle sharp edges of leaves with care; wash hands after handling plant material.",
        [
            _pr("Hold a leaf up to light; trace the veins; predict what the veins contain",
                "Fresh leaves of different sizes",
                "Guide observation: 'The vein network is the transport highway of the leaf'",
                "Connect to stems: 'Veins are continuous with the vascular bundles in the stem'",
                "Quick sketch of leaf vein network"),
            _pr("Examine diagram of leaf cross-section; identify palisade and spongy mesophyll, upper/lower epidermis, stomata",
                "Labelled leaf cross-section diagram; microscope/prepared slide if available",
                "Guided labelling; emphasise guard cells flanking each stoma",
                "Question: 'Where does water leave the leaf? Where does CO₂ enter?'",
                "Students label an unlabelled leaf cross-section diagram"),
            _pr("Discuss: why are most stomata on the underside of the leaf? Link to transpiration",
                "Completed cross-section diagrams",
                "Pose question; accept hypotheses before explaining",
                "Connect to driving question: 'Stomata are where water vapour escapes — this is transpiration'",
                "Exit ticket: state the function of guard cells"),
            _pr("Update DQB: 'How does water leave the plant through the leaf?'",
                "Driving Question Board",
                "Foreshadow Lesson 16 on transpiration",
                "Revisit phenomenon: 'When wind blows over spinach leaves, what happens at the stomata?'",
                "Students predict whether open or closed stomata cause more wilting"),
            _pr("Update plant model: add leaf detail showing veins connecting to stem xylem, and stomata on underside",
                "Running plant model diagram",
                "Model: 'Zoom in on the leaf — show the last stop before water exits the plant'",
                "Compare models with partner; share one new insight",
                "Note: students should now have a complete organ-level model (root→stem→leaf)"),
        ],
    ),
    _bio_lesson(
        7, "Xylem — Structure and Function",
        "Why is xylem described as 'dead' tissue, and how does this make it ideal for water transport?",
        "- Describe the structure of xylem: hollow dead cells, thickened lignified walls, no end walls\n- Explain how each structural feature is an adaptation for water transport (Obj a)",
        "- Draw a labelled xylem vessel diagram\n- Explain how lignification provides strength while the hollow lumen allows flow",
        "- Appreciate that even dead cells serve vital functions in living organisms",
        "Examines xylem at the cellular level, establishing the structural basis for water transport and cohesion-tension.",
        "Xylem vessel diagrams; cross-section diagram showing lignified walls; model kit or drinking straw analogy",
        "No specific hazards.",
        [
            _pr("Predict: 'What would the ideal tube for transporting water look like?' — sketch your design",
                "Blank paper; pencils",
                "Students design their ideal water transport tube before seeing xylem structure",
                "Think-aloud: 'What properties should it have? Strong walls? Open ends?'",
                "Compare student designs with a partner"),
            _pr("Examine labelled xylem vessel diagram; identify: hollow lumen, lignified wall, no end walls, pits",
                "Labelled xylem vessel diagram",
                "Guided labelling; compare student's predicted design with actual xylem",
                "Question: 'Why is it an advantage that xylem cells are dead?'",
                "Exit ticket: label an unlabelled xylem vessel and state one adaptation + function"),
            _pr("Straw analogy: blow through a drinking straw to feel resistance; discuss how xylem diameter affects flow rate",
                "Drinking straws of different diameters",
                "Demonstrate flow through narrow vs. wide straw; connect to xylem vessel diameter",
                "Introduce term: transpiration stream — continuous column of water from root to leaf",
                "Students write one sentence explaining why wide xylem vessels are more efficient"),
            _pr("Update DQB: 'How does water move continuously up the xylem against gravity?'",
                "Driving Question Board",
                "Foreshadow Lesson 14: cohesion-tension theory",
                "Students vote: 'Is it pressure from below, pull from above, or both?'",
                "Record predictions to revisit in Lesson 14"),
            _pr("Update plant model: show xylem as continuous tube from root to leaf; add lignified wall label",
                "Running plant model diagram",
                "Emphasise the continuous nature of the xylem column",
                "Discuss: 'Your model is now showing why xylem must be continuous — one break stops all flow'",
                "Model update check: does student show xylem as one unbroken pathway?"),
        ],
    ),
    _bio_lesson(
        8, "Phloem — Structure and Function",
        "How does phloem differ from xylem in structure, and why does it transport food rather than water?",
        "- Describe phloem: living sieve tube cells, companion cells, sieve plates\n- Compare xylem and phloem structures and functions (Obj a)",
        "- Draw and compare xylem and phloem diagrams\n- Explain why phloem must be living (active transport of sugars)",
        "- Appreciate that two complementary systems are needed for a plant to function fully",
        "Completes the structural study of vascular tissue by examining phloem, reinforcing the contrast with xylem through comparison.",
        "Phloem diagram; xylem diagram for comparison; comparison table template",
        "No specific hazards.",
        [
            _pr("Recall xylem features from Lesson 7; predict: 'How might the sugar transport tube differ?'",
                "Xylem diagram from Lesson 7",
                "Elicit prior knowledge; focus on: xylem is dead — should phloem be dead or living and why?",
                "Pair discussion: justify prediction",
                "Quick written prediction before revealing phloem diagram"),
            _pr("Examine phloem diagram; identify sieve tube cells, sieve plates, companion cells",
                "Labelled phloem diagram",
                "Guided labelling; emphasise sieve plates (perforated end walls allow flow) and living cells",
                "Question: 'Why do companion cells supply energy to sieve tubes?'",
                "Label an unlabelled phloem diagram"),
            _pr("Complete a comparison table: xylem vs. phloem (cells alive/dead, walls, direction of flow, substance transported)",
                "Blank comparison table template",
                "Pairs complete table; class discussion to agree correct answers",
                "Connect: 'Sugar flows in both directions in phloem — why might this be useful?'",
                "Exit ticket: state two differences between xylem and phloem"),
            _pr("Update DQB: 'Why do plants need to move sugars as well as water?'",
                "Driving Question Board",
                "Connect to translocation (foreshadow Lesson 21)",
                "Real-world link: 'Sucrose from leaves must reach growing roots, fruits, and seeds'",
                "Students predict what would happen if phloem was blocked"),
            _pr("Update plant model: add phloem alongside xylem; show bidirectional arrows for sugar flow",
                "Running plant model diagram",
                "Key addition: phloem arrows point BOTH up and down; xylem arrows point only UP",
                "Compare with previous model; discuss what changed",
                "Check: does student's model distinguish xylem (upward) from phloem (bidirectional)?"),
        ],
    ),
    _bio_lesson(
        9, "Introduction to Monocots and Dicots",
        "What are the key differences between a maize plant (monocot) and a bean plant (dicot)?",
        "- Define monocotyledon and dicotyledon\n- Identify visible differences: leaf venation, root type, number of cotyledons (Obj b)",
        "- Examine and compare monocot and dicot plants; record differences in a table",
        "- Appreciate the diversity of plant life and its organisation into meaningful categories",
        "Introduces the monocot/dicot distinction that will be examined at the tissue level in Lessons 10–12.",
        "Maize and bean seeds (germinated and dry); whole plant specimens or photos; comparison chart",
        "Handle seeds carefully; wash hands after handling plant material.",
        [
            _pr("Examine maize and bean plants side by side; list as many differences as possible",
                "Maize and bean specimens or photos; observation recording sheet",
                "Allow free observation first; then guide: focus on leaves, roots, seeds",
                "Think-Pair-Share observations",
                "Quick list: how many differences did each pair find?"),
            _pr("Focus on leaf venation: trace veins on a monocot leaf (parallel) and dicot leaf (net-like)",
                "Fresh maize and bean leaves; pencils for tracing",
                "Demo: hold leaf to light to show vein pattern clearly",
                "Question: 'What might parallel vs. net venation tell us about internal structure?'",
                "Students sketch both venation patterns with labels"),
            _pr("Focus on seeds: count cotyledons in bean (2) and maize (1); connect name to definition",
                "Bean and maize seeds soaked overnight (easily split)",
                "Students carefully split seeds and count seed leaves",
                "Etymology: 'mono' = one, 'di' = two — helps students remember",
                "Exit ticket: define monocot and dicot using the cotyledon criterion"),
            _pr("Update DQB: 'Are xylem and phloem arranged differently in monocots and dicots?'",
                "Driving Question Board",
                "Foreshadow Lessons 10–11: vascular arrangement in stems and roots",
                "Students predict based on what they already know about vascular bundles",
                "Record predictions for verification in Lessons 10–11"),
            _pr("Update plant model: annotate with M (monocot) and D (dicot) features observed",
                "Running plant model (or new diagram focused on monocot vs. dicot)",
                "Introduce concept of two model types for this sub-strand",
                "Pair share: 'What is one thing you would now change about your original model?'",
                "Observation check: do students link venation pattern to internal vascular arrangement?"),
        ],
    ),
    _bio_lesson(
        10, "Vascular Arrangement in Stems",
        "How is the vascular tissue arranged differently in monocot and dicot stems?",
        "- Describe and illustrate the arrangement of vascular bundles in monocot stems (scattered) vs. dicot stems (ring) (Obj b)",
        "- Observe, sketch and compare stem cross-sections from a monocot and a dicot",
        "- Appreciate how internal organisation at the cellular level relates to the plant's overall structure",
        "Examines the vascular arrangement in stems at the tissue level, directly fulfilling Objective b.",
        "Pre-cut cross-sections of maize stem (monocot) and bean/sunflower stem (dicot); hand lens; diagrams",
        "Only teacher cuts stems with a blade; students handle pre-cut sections only.",
        [
            _pr("Revisit Lesson 5 stem cross-section diagram; predict: will monocot and dicot stems look the same inside?",
                "Lesson 5 diagram; blank prediction sketch",
                "Elicit predictions; do not reveal answer",
                "Pair discussion; justify prediction",
                "Record predictions — to be verified during observation"),
            _pr("Observe pre-cut monocot (maize) and dicot (bean) stem cross-sections with a hand lens",
                "Pre-cut stem cross-sections; hand lens",
                "Guide observation: 'Where are the vascular bundles? Are they in a ring or scattered?'",
                "Compare: 'Point to the main difference between the two stems'",
                "Annotated sketch of both cross-sections"),
            _pr("Compare observations with labelled diagrams; confirm: scattered (monocot) vs. ring (dicot)",
                "Labelled comparative diagram of monocot vs. dicot stem cross-sections",
                "Guided labelling; emphasise epidermis, cortex, vascular bundle positions, pith",
                "Question: 'Why might a scattered arrangement give extra strength to a grass stem?'",
                "Exit ticket: draw and label monocot and dicot stem cross-sections from memory"),
            _pr("Update DQB: confirm earlier prediction about monocot vs. dicot differences",
                "Driving Question Board",
                "Tick off any confirmed answers; update any wrong predictions",
                "Connect to prior lessons: 'Same xylem and phloem — just differently arranged'",
                "Students write a 'Prediction vs. Finding' statement"),
            _pr("Update plant models: draw TWO versions — monocot stem and dicot stem cross-sections with labels",
                "Running plant model (now branches into monocot and dicot versions)",
                "Students create a side-by-side comparison in their models",
                "Compare with partner: identify at least one correct and one to-improve feature",
                "Model check: correct positioning of scattered vs. ring bundles"),
        ],
    ),
    _bio_lesson(
        11, "Vascular Arrangement in Roots",
        "How does the arrangement of vascular tissue in roots compare between monocots and dicots?",
        "- Describe the arrangement of xylem and phloem in monocot roots (alternating) vs. dicot roots (central xylem star with phloem between arms) (Obj b)",
        "- Compare root and stem cross-sections; observe the continuity of vascular tissue",
        "- Appreciate that structural patterns at the microscopic level reflect evolutionary adaptation",
        "Completes the vascular tissue arrangement study by examining roots, reinforcing Objective b.",
        "Prepared diagrams of monocot and dicot root cross-sections; comparison with stem diagrams from Lesson 10",
        "No specific hazards for diagram-based lesson.",
        [
            _pr("Recall stem vascular arrangements; predict: will root arrangements be similar or different?",
                "Stem cross-section diagrams from Lesson 10",
                "Elicit predictions; accept all hypotheses",
                "Focus question: 'In a root, which tissue is at the centre and why?'",
                "Written prediction before examining root diagrams"),
            _pr("Examine labelled diagrams of monocot and dicot root cross-sections; compare to stem",
                "Labelled monocot and dicot root cross-section diagrams",
                "Guided observation: point out central xylem star (dicot) vs. alternating xylem/phloem (monocot)",
                "Comparison table: root vs. stem arrangement for monocot and dicot",
                "Label an unlabelled root cross-section diagram"),
            _pr("Discuss: why might xylem be at the centre of a dicot root? Connect to function (strength for anchorage)",
                "Completed diagrams and comparison table",
                "Guide discussion: 'Roots are pulled and pushed by the soil — where should strength be?'",
                "Connect to xylem's lignified walls from Lesson 7",
                "Exit ticket: state one difference between monocot and dicot root cross-sections"),
            _pr("Update DQB: mark vascular arrangement questions as answered",
                "Driving Question Board",
                "Consolidate: 'We now know how xylem and phloem are arranged in both types of plant'",
                "Connect forward to Lesson 13: water now needs to cross from root to xylem",
                "Students summarise key differences between monocot and dicot in one paragraph"),
            _pr("Update plant models: add root cross-section detail showing xylem and phloem arrangement",
                "Running plant model",
                "Students add a 'zoom-in' box on the root showing tissue arrangement",
                "Compare monocot and dicot model versions",
                "Check: student models now show correct tissue positions in both stem and root"),
        ],
    ),
    _bio_lesson(
        12, "Practical Classification of Local Plants",
        "Can we identify whether a local plant is a monocot or a dicot by examining its observable features?",
        "- Apply knowledge of monocot/dicot differences to classify unfamiliar local plants\n- Justify classification using observable evidence (Obj b)",
        "- Classify at least 5 local plant specimens; record evidence for each classification",
        "- Appreciate the relevance of biology to identifying familiar plants in their local environment",
        "Applies the monocot/dicot knowledge to authentic local context through a practical classification activity.",
        "5–8 local plant specimens (grasses, crops, weeds); classification recording sheet; hand lens",
        "Wash hands after handling plants; do not handle plants with thorns without gloves.",
        [
            _pr("Receive 5 labelled plant specimens; predict monocot or dicot for each before closer examination",
                "Local plant specimens (labelled A–E); prediction sheet",
                "Guide: 'Use only external features for your initial prediction'",
                "Individual prediction then pair comparison",
                "Record predictions on sheet before investigation"),
            _pr("Examine each specimen; record evidence: leaf venation, root type (if visible), seed cotyledons",
                "Specimens; hand lens; recording sheet",
                "Circulate; prompt: 'What evidence supports your classification?'",
                "Pair discussion; record in structured evidence table",
                "Evidence table: specimen name, feature observed, conclusion"),
            _pr("Compare findings with class; discuss any disagreements; agree final classifications",
                "Class data on board; completed evidence tables",
                "Facilitate discussion of disagreements; model scientific reasoning",
                "Question: 'When two features give different answers, which should you trust more?'",
                "Exit ticket: classify one additional specimen shown on the board with justification"),
            _pr("Update DQB: confirm understanding of monocot and dicot differences",
                "Driving Question Board",
                "Celebrate: 'You can now classify plants based on scientific evidence — not just guessing'",
                "Connect to farming: 'Knowing plant type helps farmers choose appropriate techniques'",
                "Students add one real-world application to their DQB"),
            _pr("Update plant models: annotate model with local examples of monocots and dicots",
                "Running plant model",
                "Students label: 'Maize = monocot, Bean = dicot' with supporting features",
                "Pair share: compare which local examples each chose",
                "Check: correct classification with at least two pieces of supporting evidence"),
        ],
    ),
    _bio_lesson(
        13, "How Water Enters the Root — Osmosis and Diffusion",
        "How does water cross from the soil into the root hair, and then move toward the xylem?",
        "- Define osmosis and diffusion\n- Explain how water enters root hair cells by osmosis down a water potential gradient (Obj c)",
        "- Set up and interpret an osmosis demonstration\n- Trace the pathway of water from soil to xylem",
        "- Develop patience and precision in setting up experiments and recording observations accurately",
        "Introduces osmosis as the mechanism for water uptake, directly fulfilling Objective c.",
        "Dialysis tubing; sucrose solution; beakers; water; ruler; string; coloured water option",
        "Handle glassware carefully; wipe up any spills promptly.",
        [
            _pr("Predict: if you put a bag of sugar-water into plain water, what will happen to the bag?",
                "Blank prediction sheet; diagram of dialysis tubing setup",
                "Do not show result yet; elicit range of predictions",
                "Pair discussion: justify prediction using particle model of diffusion",
                "Record written predictions for comparison after demonstration"),
            _pr("Observe dialysis tubing osmosis demonstration (pre-set up by teacher); measure change in volume/mass",
                "Dialysis tubing setup in beakers; rulers or balance",
                "Show results; ask: 'What happened? Why did water move into the bag?'",
                "Connect to root hair: 'Root hair cell contents are more concentrated than soil water'",
                "Students record and explain the result"),
            _pr("Apply osmosis to root hair cell: trace water pathway from soil → root hair → cortex → xylem",
                "Diagram of root cross-section showing pathway",
                "Guided tracing: 'At each step, water moves from dilute to concentrated solution'",
                "Question: 'What happens to water potential at each step?'",
                "Exit ticket: draw an arrow showing the direction of water movement by osmosis in a root"),
            _pr("Update DQB: mark 'How does water enter the root?' as answered",
                "Driving Question Board",
                "Celebrate progress: 'You can now explain the first step in the water transport journey'",
                "Connect forward to Lesson 14: now water is in the xylem — how does it get to the leaves?",
                "Students write one sentence explaining water entry using the word 'osmosis'"),
            _pr("Update plant model: add osmosis arrows from soil to root hair and from root hair to xylem",
                "Running plant model",
                "Students add concentration gradient labels to their model",
                "Compare with Lesson 1 model: 'Look how much detail you have now'",
                "Model check: correct direction of osmosis arrows and concentration labels"),
        ],
    ),
    _bio_lesson(
        14, "Movement of Water Up the Plant",
        "What forces drive water from the roots all the way up to the leaves against gravity?",
        "- Describe the cohesion-tension theory (simplified): transpiration pull, cohesion of water molecules, adhesion to xylem walls\n- Explain why water moves upward as a continuous column (Obj c)",
        "- Demonstrate water column continuity using model or analogy\n- Explain the role of each force",
        "- Appreciate that physical forces, not a pump, drive water movement in plants",
        "Explains the mechanism driving water up the xylem, completing the water transport pathway from root to leaf.",
        "Paper towels or capillary tubes for capillarity demo; diagram of transpiration stream",
        "No specific hazards.",
        [
            _pr("Discuss: how can water climb 30 metres up a tree against gravity with no pump?",
                "Image of a tall tree; question on board",
                "Elicit ideas; accept all; emphasise there is no heart/pump in plants",
                "Pair discussion: propose a mechanism",
                "Record proposed mechanisms for class comparison"),
            _pr("Capillary rise demonstration: paper towel in water, or narrow tube vs. wide tube",
                "Paper towels; coloured water; capillary tubes of different widths",
                "Demo: water climbs higher in narrower tube — adhesion and surface tension",
                "Connect: xylem vessels are narrow, favouring capillary rise",
                "Students record observation and propose connection to xylem"),
            _pr("Introduce transpiration pull and cohesion-tension: water lost from leaves pulls the column upward",
                "Diagram of transpiration stream in xylem",
                "Analogy: chain of magnets — pulling the top pulls the whole chain",
                "Explain: cohesion keeps water molecules connected; adhesion prevents column collapse",
                "Exit ticket: state two forces that help water move up the xylem"),
            _pr("Update DQB: mark 'What forces drive water up the plant?' as answered",
                "Driving Question Board",
                "Revisit phenomenon: 'When wind dries the leaf surface, transpiration increases — what happens to the pull?'",
                "Connect: faster transpiration = stronger pull = faster water uptake from roots",
                "Students explain wilting using transpiration pull in one paragraph"),
            _pr("Update plant model: add upward arrows along xylem labelled 'transpiration pull'; add cohesion labels",
                "Running plant model",
                "Students draw a 'chain' analogy inside the xylem column",
                "Compare model with partner",
                "Check: student model shows unbroken column with forces labelled"),
        ],
    ),
    _bio_lesson(
        15, "Demonstration of Water Transport — Coloured Water Experiment",
        "Can we see water moving through xylem? How does the dye experiment show xylem function?",
        "- Describe the coloured water experiment and explain what it demonstrates\n- Interpret results to confirm xylem is the water transport vessel (Obj c)",
        "- Set up and observe a coloured water experiment; record and explain results",
        "- Appreciate that scientific demonstrations provide visible evidence for invisible processes",
        "Provides direct observable evidence that xylem transports water, consolidating Lessons 13–14.",
        "White carnations or celery stalks; red/blue food dye; water; beakers; sharp blade (teacher only); timer",
        "Only teacher uses the blade; students observe and record.",
        [
            _pr("Predict: if we put a white flower in red dye water, what will happen and why?",
                "White flowers or celery stalks; prediction sheet",
                "Do not show result; elicit predictions with reasoning",
                "Students connect to xylem knowledge: 'Where should the dye appear?'",
                "Record predictions with reasons before experiment"),
            _pr("Set up experiment: place flowers/celery in dye; observe after 30 minutes (or examine pre-prepared results)",
                "Dye-soaked specimens prepared in advance; fresh setup for class to observe over lesson",
                "Direct observation: which parts are stained? What does this tell us?",
                "Cross-section cut by teacher: show dye concentrated in xylem",
                "Students record: observation, conclusion, connection to xylem"),
            _pr("Discuss: what would happen if we cut the stem half-way and put two different colours in each half?",
                "Diagram of split-stem experiment (or actual demonstration)",
                "Pose extension question; accept hypotheses",
                "Reinforce: dye only travels in xylem — confirms xylem's role",
                "Students write: 'This experiment shows that … because …'"),
            _pr("Update DQB: mark 'How can we see water moving through the plant?' as answered",
                "Driving Question Board",
                "Celebrate: visible proof of xylem function",
                "Connect forward to Lessons 16–20: what happens to water once it reaches the leaves?",
                "Students suggest how transpiration connects to what they just observed"),
            _pr("Update plant model: add experimental evidence label — 'Confirmed by dye experiment'",
                "Running plant model",
                "Emphasise: models should be supported by evidence, not just theory",
                "Peer review: partner checks model for accuracy and evidence labels",
                "Check: model now includes evidence-based annotations"),
        ],
    ),
    _bio_lesson(
        16, "What is Transpiration?",
        "Why do plants lose water vapour through their leaves, and why is this useful?",
        "- Define transpiration: loss of water vapour from a plant, mainly through stomata\n- Explain the importance of transpiration: cooling, mineral uptake, maintaining turgor (Obj d)",
        "- Identify the route of water loss; explain the importance of transpiration for plant function",
        "- Appreciate that what seems like water 'waste' is actually a vital plant process",
        "Introduces transpiration as the driving force behind water transport, directly beginning Objective d.",
        "Diagrams of leaf with stomata; thermometer for temperature analogy; water-logged soil sample",
        "No specific hazards.",
        [
            _pr("Observe a plastic bag wrapped around a leafy branch (set up the day before); record observations",
                "Pre-prepared plant with plastic bag showing water droplets inside",
                "Ask: 'Where did the water come from? How did it get inside the bag?'",
                "Think-Pair-Share: connect to leaf structure and stomata from Lesson 6",
                "Students draw and annotate observation"),
            _pr("Define transpiration; trace the pathway of water vapour: root → xylem → leaf → stomata → air",
                "Diagram of transpiration pathway",
                "Guided diagram labelling; students add directional arrows",
                "Analogy: sweating in humans keeps us cool; stomata opening drives the water column upward",
                "Exit ticket: define transpiration in one sentence"),
            _pr("Discuss three reasons transpiration matters: cooling, mineral flow, turgor maintenance",
                "Board list of three reasons",
                "For each reason: connect to a real-world plant function",
                "Question: 'What would happen if a plant could not transpire at all?'",
                "Students rank the three reasons by importance and justify their choice"),
            _pr("Update DQB: 'What is transpiration and why does it matter?'",
                "Driving Question Board",
                "Connect to phenomenon: 'When spinach wilts, transpiration still continues but water is not replaced'",
                "Foreshadow Lessons 17–18: factors that affect transpiration rate",
                "Students predict which conditions cause fastest transpiration"),
            _pr("Update plant model: add transpiration labels at leaf stomata; show water vapour exiting",
                "Running plant model",
                "Students complete the water pathway from root to atmosphere in their model",
                "Compare full model with original Lesson 1 model",
                "Reflection: 'How has your model changed? What are you most proud of adding?'"),
        ],
    ),
    _bio_lesson(
        17, "Demonstrating Transpiration — Plastic Bag Experiment",
        "How can we use a simple plastic bag experiment to measure transpiration?",
        "- Describe the plastic bag method for demonstrating transpiration\n- Explain how mass loss provides evidence for transpiration rate (Obj d)",
        "- Set up experiment, record mass before and after, calculate water lost; interpret results",
        "- Develop accuracy in measurement and honest recording of experimental data",
        "Provides hands-on experimental evidence for transpiration, preparing students for the variables investigation in Lesson 19.",
        "Small pot plants or leafy branches; clear plastic bags; elastic bands; balance; timer",
        "Tie bags carefully; do not leave bags on plants for extended periods in full sun.",
        [
            _pr("Predict: if we cover a plant with a plastic bag, what will appear inside and why?",
                "Prediction sheet; diagram of setup",
                "Do not reveal result; elicit predictions with reasoning",
                "Connect to transpiration definition from Lesson 16",
                "Written prediction with reasoning"),
            _pr("Set up experiment: bag a plant, record mass, wait 20 minutes, record mass again; observe condensation",
                "Plant specimens; plastic bags; elastic bands; balance",
                "Students take measurements; record carefully in results table",
                "Circulate to ensure accurate reading of balance",
                "Results table: time, mass before, mass after, change in mass"),
            _pr("Calculate water lost; compare results across groups; explain why results vary",
                "Completed results tables",
                "Guided calculation; class data comparison on board",
                "Question: 'Why did some groups have more water loss than others?'",
                "Students identify possible reasons for variation — foreshadow Lesson 18"),
            _pr("Update DQB: 'How do we know transpiration is happening?'",
                "Driving Question Board",
                "Consolidate: direct evidence from experiment confirms transpiration",
                "Connect to phenomenon: 'In a drought, plants cannot replace the water lost through transpiration'",
                "Students explain wilting using transpiration rate in one sentence"),
            _pr("Update plant model: add quantitative note — 'Up to 98% of absorbed water is lost through transpiration'",
                "Running plant model",
                "Students add this statistic as an annotation",
                "Pair discussion: 'Does this surprise you? What are the implications for farming?'",
                "Check: model now includes experimental evidence and quantitative data"),
        ],
    ),
    _bio_lesson(
        18, "Factors Affecting Transpiration",
        "How do temperature, wind, light, and humidity affect how fast a plant loses water?",
        "- Identify four factors that affect transpiration rate: temperature, wind speed, light intensity, humidity (Obj d)\n- Explain the effect of each factor on stomatal aperture and evaporation",
        "- Predict the effect of changing one factor; justify predictions using scientific reasoning",
        "- Appreciate why transpiration is higher during hot, windy, sunny days — connecting to the phenomenon",
        "Explains why wilting is worse on hot, windy days, directly connecting to the anchoring phenomenon.",
        "Cards describing different weather conditions; transpiration rate data table; textbook or projected data",
        "No specific hazards.",
        [
            _pr("Sort weather condition cards (hot/cold, windy/calm, bright/dark, humid/dry) by predicted effect on transpiration",
                "Weather condition cards; sorting mat; prediction sheet",
                "Individual prediction then pair discussion",
                "Guide: 'Think about what makes water evaporate faster in everyday life'",
                "Record sorted cards as starting predictions"),
            _pr("Examine data table or graph showing transpiration rate under different conditions; identify trends",
                "Data table or graph of transpiration rate vs. each factor",
                "Guided graph reading: identify trends for each factor",
                "Question: 'Which factor has the greatest effect? How can you tell?'",
                "Students summarise each factor's effect in one sentence"),
            _pr("Explain the mechanism: high temperature → more evaporation; high wind → removes humid air; high light → opens stomata wider; low humidity → steeper gradient",
                "Annotated diagram showing stomata and moisture gradient",
                "Explain one factor at a time; connect to driving question",
                "Real-world link: 'This is why spinach wilts faster than mango on a hot windy day'",
                "Exit ticket: choose one factor and explain its effect on transpiration"),
            _pr("Revisit the phenomenon: explain why spinach wilts faster than mango on hot windy days",
                "Phenomenon description from Lesson 1; student initial models",
                "Whole-class discussion: use today's knowledge to construct an explanation",
                "Connect to Lesson 20: how can farmers use this knowledge?",
                "Students write a partial explanation of the phenomenon using transpiration factors"),
            _pr("Update plant model: add labels showing 4 factors and their effect on stomatal aperture",
                "Running plant model",
                "Add a weather condition box linking to the leaf diagram",
                "Compare: 'Your model can now explain the phenomenon!'",
                "Check: model includes all 4 factors with correct effect direction"),
        ],
    ),
    _bio_lesson(
        19, "Investigation Activity — Testing One Factor Affecting Transpiration",
        "Does a plant in direct sunlight lose more water than one in the shade?",
        "- Design and conduct a simple investigation to test one factor affecting transpiration\n- Use variables correctly: independent, dependent, controlled (Obj d)",
        "- Conduct investigation; collect data; interpret results; evaluate method",
        "- Develop scientific inquiry skills: hypothesis, fair test, data collection, evaluation",
        "Applies Lesson 18 knowledge through a student-designed investigation, developing scientific process skills.",
        "Two identical potted plants or leafy branches in bags; balance; light source or sunlit window; ruler",
        "Do not leave plants in extreme heat; ensure bags are securely fastened.",
        [
            _pr("Write a hypothesis: 'I predict that the plant in sunlight will lose MORE/LESS water because …'",
                "Hypothesis writing frame",
                "Guide: hypothesis must state direction of effect and reason",
                "Pair share hypotheses; check that reasoning connects to Lesson 18",
                "Written hypothesis before investigation begins"),
            _pr("Set up experiment: place two identical plants in sun vs. shade; record mass every 10 minutes for 30 minutes",
                "Identical plants in bags; balance; timer; results table",
                "Students take measurements and record; teacher circulates for accuracy",
                "Guided: 'What are you keeping the same? What are you changing? What are you measuring?'",
                "Results table with at least 3 time points"),
            _pr("Interpret results: did the data support the hypothesis? Calculate water loss for each condition",
                "Completed results tables; class data compilation",
                "Guided interpretation: connect data pattern to mechanism from Lesson 18",
                "Question: 'What are the limitations of this investigation?'",
                "Written conclusion: results + explanation + one limitation"),
            _pr("Update DQB: confirm answer to 'What factors affect transpiration rate?'",
                "Driving Question Board",
                "Celebrate: 'You designed and conducted a scientific investigation!'",
                "Connect forward to Lesson 20: how can farmers use this knowledge?",
                "Students identify one improvement to their experimental design"),
            _pr("Update plant model: annotate with 'Confirmed by investigation' for the light factor",
                "Running plant model",
                "Emphasise evidence-based science: model claims should be supported by data",
                "Compare investigation results with data from Lesson 18",
                "Check: student can identify IV, DV, and CVs in their own investigation"),
        ],
    ),
    _bio_lesson(
        20, "Application to Farming — Managing Transpiration",
        "How can farmers use knowledge of transpiration to protect their crops from wilting?",
        "- Describe how mulching, windbreaks, shade nets, and irrigation timing reduce water loss\n- Connect transpiration knowledge to practical farming strategies (Obj d)",
        "- Evaluate farming techniques using transpiration knowledge; suggest improvements for a local farm scenario",
        "- Value the practical application of biology to food security in their local context",
        "Connects the science of transpiration directly to the anchoring phenomenon (wilting crops) and to real farming practice.",
        "Photos of mulched/unmulched crops; diagrams of windbreak and shade net; local farming scenario card",
        "No specific hazards.",
        [
            _pr("Scenario: a local farmer's spinach wilts every afternoon despite watering in the morning — what advice would you give?",
                "Scenario description card; local farm photo if available",
                "Pose the scenario; do not give answers",
                "Pair discussion: use knowledge from Lessons 16–19 to generate advice",
                "Record advice before class discussion"),
            _pr("Examine photos of mulching, windbreaks, shade nets, drip irrigation; identify how each reduces transpiration",
                "Photos or diagrams of each technique",
                "Guided discussion: for each technique, ask 'Which factor does this control?'",
                "Connect to Lesson 18 factors: mulch reduces temperature; windbreak reduces wind speed",
                "Summary table: technique → factor controlled → effect on transpiration"),
            _pr("Evaluate which technique is most suitable for a smallholder farmer in a semi-arid area",
                "Cost and practicality considerations for each technique",
                "Small group discussion; each group presents their recommended technique with justification",
                "Teacher adds: no single solution — context determines best approach",
                "Exit ticket: choose one technique and justify using transpiration science"),
            _pr("Revisit the phenomenon: complete explanation of why spinach wilts faster than mango on hot windy days",
                "Phenomenon from Lesson 1; student explanation attempts from Lesson 18",
                "Class builds a complete explanation on the board",
                "Check: explanation includes stomata, transpiration pull, xylem, factors",
                "Students write their personal complete explanation of the phenomenon"),
            _pr("Update plant model: add 'farming interventions' box showing techniques that reduce transpiration rate",
                "Running plant model",
                "Final model should now connect biology to farming practice",
                "Compare Lesson 1 model with Lesson 20 model side by side",
                "Reflection: 'What is the single most important thing you have learned about plant transport?'"),
        ],
    ),
    _bio_lesson(
        21, "Movement of Food in Plants — Translocation",
        "How does the food made in leaves reach all other parts of the plant?",
        "- Define translocation: movement of dissolved food (sucrose) through phloem from source to sink\n- Identify sources (leaves) and sinks (roots, fruits, growing tips) (Obj e)",
        "- Trace the pathway of sucrose from source to sink; interpret ringing experiment evidence",
        "- Appreciate the complexity of plant physiology and how each system supports the others",
        "Directly fulfils Objective e and completes the transport system overview by adding the food transport pathway.",
        "Diagram of source/sink model; ringing experiment diagram; labelled phloem diagram from Lesson 8",
        "No specific hazards.",
        [
            _pr("Recall from Lesson 8: phloem transports sugar. Pose question: 'Where is sugar made? Where does it need to go?'",
                "Lesson 8 xylem/phloem comparison table",
                "Elicit: sugar made in leaves by photosynthesis; must reach roots, growing tips, fruits",
                "Introduce source (where sugar is made) and sink (where sugar is used)",
                "Students draw a simple source-to-sink diagram for one plant"),
            _pr("Introduce translocation and the ringing experiment: removing a ring of bark (phloem) blocks food transport",
                "Diagram of ringing experiment showing swelling above and withering below",
                "Explain: bark includes phloem; removal blocks sugar but NOT water (xylem is deeper)",
                "Question: 'Why does the tissue ABOVE the ring swell while the root withers?'",
                "Students explain the ringing experiment result using source/sink/phloem knowledge"),
            _pr("Trace translocation pathway: leaf → phloem → stem → root/fruit/growing tip",
                "Whole-plant diagram showing source and sink with phloem pathway",
                "Emphasise: phloem carries sugar in BOTH directions (up to growing tips AND down to roots)",
                "Compare with xylem: 'Xylem goes only up; phloem goes both ways — why?'",
                "Exit ticket: state the definition of translocation and name one source and two sinks"),
            _pr("Update DQB: mark 'How does food move in the plant?' as answered",
                "Driving Question Board",
                "Connect translocation to xylem transport: both needed for a fully functioning plant",
                "Real-world link: sugar translocation to fruits is what makes them edible and sweet",
                "Students write a comparison: 'Xylem transports … Phloem transports … Both are similar in that …'"),
            _pr("Update plant model: add phloem pathway with source and sink labels; show bidirectional sugar flow",
                "Running plant model",
                "Students complete the food transport pathway in their model",
                "Final model should now show BOTH water transport (xylem) and food transport (phloem)",
                "Pair review: check both pathways are correctly drawn and labelled"),
        ],
    ),
    _bio_lesson(
        22, "Unit Summary and Assessment",
        "How do xylem and phloem work together to keep a plant alive — and how does this explain our phenomenon?",
        "- Summarise all five learning objectives: structures (Obj a), monocot/dicot (Obj b), water uptake (Obj c), transpiration (Obj d), translocation (Obj e)\n- Construct a complete explanation of the anchoring phenomenon",
        "- Present or write a complete explanation linking the whole transport system to the phenomenon\n- Self-assess against unit objectives",
        "- Reflect on how scientific understanding developed through the unit; appreciate the integration of all systems",
        "Summative lesson: students return to the phenomenon and construct a full evidence-based explanation using all unit knowledge.",
        "Student running plant models; Summary Table; unit notes; assessment sheet",
        "No specific hazards.",
        [
            _pr("Return to Lesson 1 initial model: compare with current model; identify growth in understanding",
                "Lesson 1 model and current model side by side",
                "Students reflect: 'What is in my current model that was not in my Lesson 1 model?'",
                "Class discussion: celebrate how far understanding has developed",
                "Quick inventory: can students name all 5 objectives?"),
            _pr("Write a complete explanation of the phenomenon using all five objectives",
                "Phenomenon description from Lesson 1; writing frame with 5 sentence starters",
                "Students write independently for 15 minutes",
                "Circulate: check for use of key vocabulary (osmosis, transpiration, xylem, phloem, translocation)",
                "Collect as summative assessment piece"),
            _pr("Present explanation to a partner; partner gives feedback using a checklist",
                "Peer feedback checklist linked to 5 objectives",
                "Structured peer review: 'I can see … I think you should add …'",
                "Students revise explanation based on feedback",
                "Final self-assessment: rate confidence on each objective (1–3)"),
            _pr("Update DQB: mark all questions as answered; identify any remaining uncertainties",
                "Driving Question Board",
                "Class celebration: 'We answered the driving question!'",
                "Acknowledge remaining questions as areas for future study",
                "Students add one new question that the unit raised for them"),
            _pr("Submit final updated plant model with all pathways, labels and evidence annotations",
                "Running plant model (final version)",
                "Final model should show: root hair absorption, xylem pathway, leaf stomata, transpiration factors, phloem pathway, source/sink",
                "Teacher collects models for portfolio assessment",
                "Reflection: 'Which lesson most changed your thinking and why?'"),
        ],
    ),
]

# Biology summary table rows (source: lesson titles; detail: generated)
_BIO_SUMMARY_ROWS = [
    ("Lesson 1: Introduction to the Phenomenon",
     "Observed wilting spinach; drew initial models of water movement inside plants; posted questions on Driving Question Board",
     "Plants need water for survival; when water is lost faster than absorbed, cells lose turgor and the plant wilts",
     "The phenomenon shows what happens when the transport system cannot keep up with water loss",
     "DQB Started: Why do some crops wilt quickly? Initial Model: Water moves from roots to leaves — unclear how"),
    ("Lesson 2: What Do Plants Need to Survive?",
     "Sorted plant needs into categories; linked each need to a transport process; updated initial model",
     "Plants need water, mineral salts, CO₂ and light; water and minerals must be transported from soil to all cells",
     "Transport is essential because plant cells everywhere need water, not just roots",
     "DQB Updated: 'How do plants get water from soil?' Model Revised: Added entry points for water and minerals"),
    ("Lesson 3: Overview of Plant Transport System",
     "Colour-coded whole-plant diagram with xylem (blue) and phloem (green) pathways",
     "Xylem transports water and minerals upward; phloem transports sugars in both directions",
     "The plant has two separate transport vessels, each with a different role in keeping all cells supplied",
     "DQB Updated: 'What is the difference between xylem and phloem?' Model Revised: Added two distinct pathways"),
    ("Lesson 4: Root Structure and Function",
     "Observed germinated bean roots; drew and labelled root hair cell; calculated surface area advantage",
     "Root hair cells have a large surface area, thin walls and large vacuole — all adaptations for rapid water absorption",
     "The root hair design maximises the rate of water uptake, which is critical when plants are losing water through leaves",
     "DQB Updated: 'How does water enter the root?' Model Revised: Added root hair cells with absorption arrows"),
    ("Lesson 5: Stem Structure",
     "Observed stem cross-sections; identified vascular bundles; watched dye travel through xylem",
     "The stem contains vascular bundles with xylem (innermost) and phloem; xylem is continuous from root to leaf",
     "The stem's internal transport highway connects root absorption to leaf use — disrupting it would kill the plant",
     "DQB Updated: 'Where exactly is xylem in the stem?' Model Revised: Added stem cross-section with vascular bundle ring"),
    ("Lesson 6: Leaf Structure and Transport Role",
     "Examined leaf cross-sections; identified veins, stomata and guard cells; linked to transpiration",
     "Leaves have veins (extensions of vascular bundles) and stomata (controlled pores) where water vapour exits",
     "The leaf is where water exits the plant — stomata are the gateway between the plant's transport system and the atmosphere",
     "DQB Updated: 'How does water leave the plant through the leaf?' Model Revised: Added stomata on leaf underside"),
    ("Lesson 7: Xylem — Structure and Function",
     "Drew labelled xylem vessel; compared with student-designed ideal water tube; discussed lignification",
     "Xylem vessels are dead, hollow, lignified cells with no end walls — perfectly adapted for continuous water flow",
     "Xylem structure explains why water can travel continuously from root to leaf without leaking or collapsing",
     "DQB Updated: 'How does water move up against gravity?' Model Revised: Added lignified walls and hollow lumen labels"),
    ("Lesson 8: Phloem — Structure and Function",
     "Drew phloem diagram; completed xylem vs. phloem comparison table",
     "Phloem has living sieve tube cells and companion cells; sugar transport requires living cells with active processes",
     "Unlike xylem, phloem must be alive because sugar loading and unloading requires energy — this is active transport",
     "DQB Updated: 'Why do plants need to move sugars too?' Model Revised: Added phloem with bidirectional arrows"),
    ("Lesson 9: Introduction to Monocots and Dicots",
     "Compared maize and bean specimens; identified differences in venation, root type and cotyledon number",
     "Monocots have one cotyledon, parallel leaf venation; dicots have two cotyledons, net-like venation",
     "Both monocots and dicots have xylem and phloem — but their arrangement differs, which will affect transport efficiency",
     "DQB Updated: 'Are xylem and phloem arranged differently in monocots and dicots?' Model Revised: Two model types started"),
    ("Lesson 10: Vascular Arrangement in Stems",
     "Observed pre-cut monocot (maize) and dicot (bean) stem cross-sections; compared and sketched both",
     "Monocot stems have vascular bundles scattered throughout; dicot stems have vascular bundles in a ring",
     "The different arrangements reflect different evolutionary strategies — both still successfully transport water and food",
     "DQB Updated: Confirmed monocot/dicot stem differences. Model Revised: Added correct vascular arrangement to both stem types"),
    ("Lesson 11: Vascular Arrangement in Roots",
     "Compared monocot and dicot root cross-sections using diagrams; linked position of xylem to function",
     "Dicot roots have a central xylem star with phloem between the arms; monocot roots have alternating xylem and phloem",
     "Central xylem in dicot roots provides structural strength for anchorage — form follows function",
     "DQB Updated: Confirmed monocot/dicot root differences. Model Revised: Added root cross-section with correct tissue positions"),
    ("Lesson 12: Practical Classification of Local Plants",
     "Examined 5+ local plant specimens; classified each as monocot or dicot using observable evidence",
     "Local plants can be classified as monocots or dicots using venation, root type and cotyledon evidence",
     "Classification skills allow farmers and scientists to predict plant behaviour and apply appropriate management techniques",
     "DQB Updated: Added application question. Model Revised: Annotated with local plant examples"),
    ("Lesson 13: How Water Enters the Root",
     "Observed osmosis demonstration; traced water pathway from soil to xylem; applied osmosis to root hair cells",
     "Water enters root hair cells by osmosis down a water potential gradient; it moves from dilute soil water to concentrated cell contents",
     "Osmosis is the mechanism that drives the first step in water transport — without it, the entire system stops",
     "DQB Updated: Marked 'How does water enter the root?' as answered. Model Revised: Added osmosis arrows and concentration gradients"),
    ("Lesson 14: Movement of Water Up the Plant",
     "Observed capillary rise; explained transpiration pull and cohesion-tension theory",
     "Transpiration pull, cohesion between water molecules and adhesion to xylem walls drive the continuous upward flow of water",
     "The driving force for water transport is at the TOP of the plant (leaves), not the bottom — this explains how trees can be so tall",
     "DQB Updated: Marked 'What forces drive water upward?' as answered. Model Revised: Added force labels along xylem column"),
    ("Lesson 15: Demonstration of Water Transport",
     "Set up and observed coloured water experiment; cross-sectioned dye-stained stems to confirm xylem transport",
     "The coloured water experiment provides direct evidence that xylem (and only xylem) transports water through the stem",
     "Visible evidence reinforces the model: xylem is the water highway, and it runs continuously from root to leaf",
     "DQB Updated: 'How can we see water moving through the plant?' answered. Model Revised: Added 'Confirmed by dye experiment' label"),
    ("Lesson 16: What is Transpiration?",
     "Observed condensation in plastic bag experiment; defined transpiration; explained its importance",
     "Transpiration is the loss of water vapour from a plant mainly through stomata; it drives water uptake and cools the plant",
     "Transpiration explains why plants continuously need water — it is the engine that pulls water upward through the whole plant",
     "DQB Updated: 'What is transpiration?' answered. Model Revised: Completed the water pathway from root to atmosphere"),
    ("Lesson 17: Demonstrating Transpiration",
     "Set up plastic bag experiment; measured and calculated water loss; compared results across groups",
     "Transpiration can be measured by recording mass loss; rate varies between plants and conditions",
     "Quantitative evidence shows that plants lose significant amounts of water through transpiration — critical in drought conditions",
     "DQB Updated: Added quantitative data on transpiration rate. Model Revised: Added 'Up to 98% of water is lost through transpiration'"),
    ("Lesson 18: Factors Affecting Transpiration",
     "Examined data on transpiration rate under different conditions; explained mechanism of each factor",
     "Higher temperature, increased light, stronger wind and lower humidity all increase the transpiration rate",
     "These factors explain exactly why spinach wilts faster on hot, windy days — the conditions accelerate water loss beyond what roots can replace",
     "DQB Updated: 'What conditions increase transpiration?' answered — directly explains the phenomenon. Model Revised: Added 4-factor weather box"),
    ("Lesson 19: Investigation Activity",
     "Designed and conducted experiment testing effect of light on transpiration; collected and interpreted data",
     "Scientific investigation skills: hypothesis, variable control, data collection, interpretation and evaluation",
     "Students tested the connection between the transpiration factor (light) and the phenomenon (wilting in sun) experimentally",
     "DQB Updated: Light factor confirmed by student data. Model Revised: Added 'Confirmed by investigation' for light factor"),
    ("Lesson 20: Application to Farming",
     "Analysed farming techniques (mulching, windbreaks, shade nets); linked each to a transpiration factor",
     "Mulching, windbreaks and shade nets reduce transpiration by controlling temperature, wind speed and light intensity",
     "Knowledge of transpiration allows farmers to protect crops from wilting — directly solving the problem in the phenomenon",
     "DQB Updated: Phenomenon fully explained. Model Revised: Added farming interventions box connecting biology to practice"),
    ("Lesson 21: Movement of Food in Plants",
     "Introduced translocation; interpreted ringing experiment; traced sucrose from source to sink",
     "Translocation is the movement of dissolved sugars through phloem from source (leaves) to sink (roots, fruits, growing tips)",
     "While xylem transports water for survival, phloem transports food for growth — both systems are essential for a complete plant",
     "DQB Updated: 'How does food move in the plant?' answered. Model Revised: Added phloem source/sink labels and bidirectional arrows"),
    ("Lesson 22: Unit Summary and Assessment",
     "Compared initial and final plant models; wrote complete explanation of phenomenon; presented to peers",
     "The complete plant transport system: osmosis (water entry) + xylem (water pathway) + transpiration (water exit) + phloem (food transport)",
     "The phenomenon is fully explained: spinach wilts because transpiration in hot, windy conditions exceeds the rate of water uptake via xylem",
     "DQB Complete: All questions answered. Final Model: Shows complete dual transport system from soil to atmosphere with all mechanisms labelled"),
]

# Biology final explanation sections
_BIO_FE_SECTIONS = [
    (
        "SECTION 1: HOW DO PLANTS OBTAIN AND TRANSPORT WATER?",
        "How does water enter a plant from the soil?\n"
        "What is osmosis and how does it apply to root hair cells?\n"
        "What forces drive water up the xylem from root to leaf?\n"
        "How is the structure of xylem adapted for water transport?",
        "Water enters the plant through root hair cells by osmosis. Root hair cells have a large surface area, thin walls and a vacuole containing a more concentrated solution than the surrounding soil water. Water moves by osmosis from the dilute soil solution into the concentrated cell contents, across the semi-permeable membrane.\n\n"
        "Once inside the root, water moves toward the xylem through the cortex cells. It then enters the xylem and travels upward in a continuous column. The driving force is transpiration pull: as water evaporates from leaf cells through stomata, it creates tension that pulls the water column upward. Water molecules remain connected (cohesion) and adhere to xylem walls (adhesion), preventing the column from breaking.\n\n"
        "Xylem vessels are perfectly adapted for this role: dead cells form a hollow, continuous tube; lignified walls provide strength; absence of end walls allows unobstructed flow."
    ),
    (
        "SECTION 2: HOW IS VASCULAR TISSUE ARRANGED IN MONOCOTS AND DICOTS?",
        "What is the difference between monocotyledonous and dicotyledonous plants?\n"
        "How are vascular bundles arranged in a monocot stem compared to a dicot stem?\n"
        "How is the arrangement in roots different from stems?\n"
        "Draw and label cross-sections of monocot and dicot stems and roots.",
        "Monocots (e.g. maize) have one cotyledon, parallel leaf venation and fibrous root systems. Dicots (e.g. bean, sunflower) have two cotyledons, net-like venation and a tap root system.\n\n"
        "In monocot stems, vascular bundles are scattered throughout the ground tissue. In dicot stems, vascular bundles are arranged in a ring near the outside. In both cases, each vascular bundle contains xylem (inner) and phloem (outer).\n\n"
        "In dicot roots, xylem forms a central star shape with phloem positioned between the arms. In monocot roots, xylem and phloem alternate around the centre. The central xylem in dicot roots provides structural strength for anchorage."
    ),
    (
        "SECTION 3: WHAT FACTORS AFFECT TRANSPIRATION AND WHY DOES IT MATTER?",
        "What is transpiration and where does it occur?\n"
        "What four environmental factors affect the rate of transpiration?\n"
        "Explain the mechanism by which each factor increases or decreases transpiration.\n"
        "How can farmers use knowledge of transpiration to protect crops?",
        "Transpiration is the loss of water vapour from a plant, mainly through stomata in the leaf surface. It drives the upward movement of water through xylem and helps cool the plant.\n\n"
        "Four factors affect transpiration rate:\n"
        "• Temperature: higher temperature increases the kinetic energy of water molecules, increasing evaporation rate.\n"
        "• Wind speed: moving air removes the humid layer around stomata, maintaining a steep water potential gradient and increasing evaporation.\n"
        "• Light intensity: light causes stomata to open wider, allowing more water vapour to escape.\n"
        "• Humidity: lower humidity outside the leaf creates a steeper water potential gradient, increasing the rate of water vapour loss.\n\n"
        "Farmers can reduce transpiration by: mulching (reduces soil temperature), windbreaks (reduce wind speed), shade nets (reduce light intensity and temperature), and timing irrigation to early morning or evening when transpiration is lowest."
    ),
    (
        "SECTION 4: HOW IS FOOD TRANSPORTED IN PLANTS?",
        "What is translocation and which vessel carries it out?\n"
        "Define source and sink in the context of translocation.\n"
        "How does the ringing experiment provide evidence for phloem transport?\n"
        "How does phloem structure support translocation?",
        "Translocation is the movement of dissolved organic substances (mainly sucrose) through phloem from sources to sinks. Sources are regions of sugar production or release (mainly leaves carrying out photosynthesis). Sinks are regions of sugar use or storage (roots, growing tips, developing fruits and seeds).\n\n"
        "The ringing experiment provides evidence: removing a ring of bark (which includes phloem) from a stem causes swelling above the ring (sugars accumulate) and withering below (sugars cannot reach the roots). The xylem below the ring is unaffected because xylem lies deeper in the stem.\n\n"
        "Phloem sieve tube cells are living (unlike xylem), with sieve plates at their ends allowing sugar solution to flow. Companion cells supply the energy needed to actively load sugar into the phloem. Unlike xylem, phloem transports in both directions — up to growing tips and down to roots."
    ),
    (
        "SECTION 5: HOW DOES THE PLANT TRANSPORT SYSTEM EXPLAIN WILTING?",
        "Explain the complete pathway of water from soil to atmosphere.\n"
        "Why do some plants wilt faster than others under the same conditions?\n"
        "Connect all five learning objectives to explain the anchoring phenomenon.\n"
        "How does the transport system enable a plant to survive in varying conditions?",
        "The complete water pathway: Soil → (osmosis) → Root hair cells → (osmosis/diffusion) → Cortex → Xylem → (transpiration pull/cohesion-tension) → Stem xylem → Leaf xylem → Mesophyll cells → (evaporation) → Stomata → Atmosphere.\n\n"
        "Spinach wilts faster than mango because: spinach has large, thin leaves with many stomata exposed to wind and sun (high transpiration rate); its shallow root system reaches less soil water; mango has a deep tap root and waxy leaves with fewer stomata (lower transpiration rate).\n\n"
        "When transpiration rate exceeds the rate of water uptake through roots, cells lose turgor pressure and the plant wilts. In severe cases, stomata close to reduce water loss, but this also stops photosynthesis.\n\n"
        "The complete transport system (xylem for water, phloem for food) enables every cell to receive water and nutrients regardless of its distance from the soil or leaves — making the plant an integrated, coordinated organism."
    ),
]

# Biology META
_BIO_META = {
    "subject":    "Biology",
    "grade":      "Grade 10",
    "strand":     "2.0 Anatomy and Physiology of Plants",
    "substrand":  "2.2 Transport System in Plants",
    "lessons":    22,
    "author":     "Jackline Mwambere",
    "driving_question": (
        "Why do some crops like spinach wilt quickly during dry, windy days, "
        "while others like mango trees remain fresh and green?"
    ),
    "phenomenon": (
        "Spinach and sukuma wiki plants in a village garden wilt quickly on hot, "
        "windy days — even after being watered. Nearby mango and cassava plants "
        "remain green and upright under the same conditions."
    ),
    "substrand_overview_rows": [
        ("Subject",            "Biology"),
        ("Grade",              "Grade 10"),
        ("Strand",             "2.0 Anatomy and Physiology of Plants"),
        ("Sub-Strand",         "2.2 Transport System in Plants"),
        ("Author",             "Jackline Mwambere"),
        ("Number of Lessons",  "22"),
        ("Number of Periods",  "44"),
        ("Anchoring Phenomenon",
         "Spinach and sukuma wiki plants wilt quickly on hot, windy days even after watering, "
         "while nearby mango and cassava plants remain fresh. Why?"),
        ("Driving Question",
         "Why do some crops like spinach wilt quickly during dry, windy days, "
         "while others like mango trees remain fresh and green?"),
        ("Learning Outcomes: Knowledge",
         "a) Relate structures of the plant transport system to their functions\n"
         "b) Illustrate vascular tissue arrangement in monocots and dicots\n"
         "c) Demonstrate water and mineral uptake from roots to leaves\n"
         "d) Demonstrate factors affecting transpiration rate\n"
         "e) Describe translocation of manufactured food"),
        ("Learning Outcomes: Skills",
         "- Conduct investigations on transport, transpiration and translocation\n"
         "- Draw and label diagrams of plant transport structures\n"
         "- Collect, record and interpret experimental data\n"
         "- Apply knowledge to real farming challenges"),
        ("Learning Outcomes: Attitudes",
         "- Appreciate the importance of plant transport to food security\n"
         "- Show curiosity and care with living materials\n"
         "- Value teamwork and honest recording in investigations"),
        ("Core Competencies",
         "Digital literacy · Collaboration/Communication · Critical thinking · "
         "Self-efficacy · Learning to learn"),
        ("Core Values",
         "Curiosity · Responsibility · Respect · Care and Compassion · "
         "Excellence"),
        ("Assessment",
         "Formative: observation, discussion, diagrams, practical reports, summary table\n"
         "Summative: Final Explanation document"),
        ("Resources",
         "Live plant specimens (spinach, maize, bean, mango leaf); microscopes/hand lens; "
         "dialysis tubing; food dye; Khan Academy videos; PhET simulations"),
    ],
    "summary_table_headers": [
        "Lesson", "Learner Activities", "Key Understanding Built",
        "Connection to Driving Question", "DQB / Model Evolution"
    ],
    "summary_table_rows": _BIO_SUMMARY_ROWS,
    "summary_instructions": (
        "Use this table to track your learning journey across all 22 lessons. "
        "For each lesson, record: what you did, what you understood, how it connects "
        "to the driving question, and how your Driving Question Board and plant model evolved. "
        "Your completed Summary Table is evidence of your scientific thinking."
    ),
    "final_explanation_sections": _BIO_FE_SECTIONS,
    "final_explanation_rubric": {
        "headers": ["Criterion", "Beginning (1)", "Developing (2)", "Proficient (3)", "Advanced (4)"],
        "rows": [
            ("1. Water Transport",
             "Names xylem but cannot explain how water moves",
             "Describes osmosis and xylem but links are incomplete",
             "Explains osmosis, transpiration pull and cohesion-tension correctly",
             "Connects all forces with clear diagrams and accurate scientific vocabulary"),
            ("2. Vascular Structure",
             "Cannot distinguish monocot from dicot",
             "Identifies one difference between monocot and dicot",
             "Accurately describes vascular arrangement in stems and roots for both types",
             "Draws labelled cross-sections with correct tissue positions and explains why arrangements differ"),
            ("3. Transpiration",
             "Defines transpiration but cannot identify factors",
             "Lists factors but explanations are partial",
             "Explains mechanism of each factor correctly; links to phenomenon",
             "Includes quantitative reasoning; connects to farming applications and phenomenon"),
            ("4. Translocation",
             "Confuses xylem and phloem functions",
             "States that phloem carries sugar but cannot explain source/sink",
             "Explains source, sink, and ringing experiment evidence correctly",
             "Integrates translocation with water transport to show the complete dual-system model"),
            ("5. Phenomenon Explanation",
             "Cannot connect transport system to wilting",
             "Makes partial connection to wilting but misses key factors",
             "Explains why spinach wilts faster using transpiration and xylem concepts",
             "Complete, evidence-based explanation using all 5 objectives; proposes farming solutions"),
        ],
    },
    "lessons_data": _BIO_LESSONS,
}

# ═══════════════════════════════════════════════════════════════════════════════
# META — Physics: Moments and Equilibrium (6 lessons)
# ═══════════════════════════════════════════════════════════════════════════════

_PHY_SLO_K = (
    "- Define centre of gravity for regular and irregular objects\n"
    "- Identify the states of equilibrium (stable, unstable, neutral)\n"
    "- Explain the principle of moments\n"
    "- Describe moment of a force, torque, and couple\n"
    "- Explain moment about one and two points of support\n"
    "- State how forces are resolved in equilibrium situations\n"
    "- Outline applications of moments and stability in everyday life"
)
_PHY_SLO_S = (
    "- Determine the centre of gravity of regular and irregular objects through practical activities\n"
    "- Calculate the moment of a force about a point\n"
    "- Apply the principle of moments to solve numerical and practical problems\n"
    "- Demonstrate the turning effect of forces using simple apparatus\n"
    "- Carry out investigations to determine factors affecting the stability of objects"
)
_PHY_SLO_A = (
    "- Appreciate the importance of moments and stability in daily life (e.g., transport, structures)\n"
    "- Demonstrate curiosity and creativity during practical investigations\n"
    "- Show teamwork and collaboration when carrying out experiments\n"
    "- Develop honesty and integrity in recording and presenting results"
)

def _phy_lesson(num, title, inquiry, know, skills, attit, purpose, materials, safety, p1):
    return {
        "number": num, "title": title,
        "inquiry_question": inquiry,
        "slo_knowledge": know, "slo_skills": skills, "slo_attitudes": attit,
        "overview_purpose": purpose,
        "materials": materials, "safety": safety,
        "period1_heading": "Period 1 (40 minutes)",
        "period1_table": p1,
        "period2_heading": "Period 2 (40 minutes)",
        "period2_table": [],
        "reflections": [],
    }

_PHY_LESSONS = [
    _phy_lesson(
        1, "Introduction to the Phenomenon",
        "Why do some objects balance easily while others topple, even when they seem similar?",
        "- Recall that forces can cause objects to rotate or balance\n- Identify the driving question for the unit",
        "- Observe and record balance/toppling behaviour; draw an initial model showing where balance occurs",
        "- Show curiosity about balance and stability in everyday objects",
        "Launches the unit phenomenon. Students observe balance and toppling in real objects, generate questions, and produce initial models.",
        "Ruler; stones or small weights; a wheelbarrow or photograph of one; Driving Question Board; blank paper",
        "Ensure stones/weights are handled safely; no sharp edges on improvised pivots.",
        [
            _pr("Observe a ruler balanced at different points with a stone; record what happens when the stone moves",
                "Ruler and small stone or eraser as pivot; demonstration setup",
                "Pose the phenomenon: 'Why do some objects topple while others stay stable?' — show without explaining",
                "Think–Pair–Share: individual observation → partner discussion → whole-class share",
                "Exit ticket: 'One thing I observed about balance today and one question I still have'"),
            _pr("Draw initial model: where does balance occur and why does an object topple?",
                "Blank A4 paper; large class diagram for reference",
                "Prompt: 'Draw what you think is happening when the ruler balances and when it tips'",
                "Pair-share diagrams; identify agreements and disagreements about where balance occurs",
                "Collect initial model diagrams for formative comparison at end of unit"),
            _pr("Contribute questions to the Driving Question Board",
                "Driving Question Board (poster or whiteboard section); sticky notes",
                "Facilitate DQB creation; group questions by theme (balance, forces, everyday objects)",
                "Model 'I notice… I wonder…' sentence stems to scaffold question generation",
                "Count and categorise student questions to inform lesson sequence"),
            _pr("Discuss: what real-life situations involve balance? (wheelbarrows, ladders, vehicles)",
                "Photographs or sketches of wheelbarrow, ladder, tall vehicle, see-saw",
                "Elicit prior knowledge without correcting; accept all ideas at this stage",
                "Record on KWL chart; connect to driving question",
                "Oral probe: 'Name one situation where balance matters and guess why'"),
            _pr("Update initial model: add arrows showing forces and where weight seems to act",
                "Initial model from earlier in the lesson",
                "Explain models will be revised as evidence accumulates each lesson",
                "Students share updated models with partners; note what is uncertain",
                "Collect models; note common misconceptions for future lessons"),
        ],
    ),
    _phy_lesson(
        2, "Centre of Gravity",
        "Where is the balancing point of an object, and what determines it?",
        "- Define centre of gravity as the point where an object's entire weight appears to act\n- Locate the centre of gravity of regular shapes by symmetry\n- Determine the centre of gravity of irregular shapes by experiment",
        "- Balance objects on a pivot to find the centre of gravity\n- Use the plumb-line method for irregular cut-out shapes\n- Record and explain experimental results",
        "- Show patience and precision when locating balance points\n- Appreciate that the centre of gravity concept applies to all objects including the human body",
        "Students investigate the concept of centre of gravity through hands-on balancing activities with regular and irregular shapes.",
        "Ruler; pencil as pivot; cardboard cut-outs (regular and irregular); string; small weight (plumb bob); pins or needles; retort stand",
        "Take care with pins/needles; ensure retort stands are stable.",
        [
            _pr("Balance a regular cardboard rectangle on a pencil; mark the balance line; repeat from different direction to find COG",
                "Cardboard rectangles; pencil pivot; marker pen",
                "Demonstrate first; then guide students: 'Balance the shape — the point where all balance lines cross is the COG'",
                "Groups share COG locations; discuss: 'Is the COG always in the centre for regular shapes?'",
                "Quick check: can students predict COG of a square before testing?"),
            _pr("Use plumb-line method to find COG of irregular cardboard shapes",
                "Irregular cardboard cut-outs; string and small weight; pins; retort stand or clip",
                "Model the procedure: hang shape, draw plumb line; repeat from second point; intersection = COG",
                "Groups compare irregular COG locations; discuss: 'Is the COG always inside the object?'",
                "Observation: accuracy of plumb lines; can students explain the method in their own words?"),
            _pr("Investigate: can the COG be outside the object? (e.g., a ring shape or L-shape)",
                "L-shaped or ring-shaped cardboard cut-outs",
                "Challenge question: 'Draw where you think the COG is before testing — then test and compare'",
                "Discuss findings as a class; link to everyday objects (boomerangs, rings, hollow objects)",
                "Exit ticket: sketch a shape and mark its approximate COG, explain your reasoning"),
            _pr("Update model: mark COG on the initial balance model from Lesson 1",
                "Initial model diagrams from Lesson 1",
                "Prompt: 'Now that you know about COG, update your model to show where weight acts'",
                "Pair discussion: 'What changed in your model? What do you now understand that you didn't before?'",
                "Observe model updates; check that students connect COG to balance point"),
            _pr("Solve 2–3 practice problems: predict COG location for given shapes; verify by experiment",
                "Worksheet with shape diagrams; cardboard cut-outs to test predictions",
                "Circulate and ask probing questions; correct misconceptions individually",
                "Peer-check answers; discuss any disagreements as a class",
                "Written responses: 'Define COG in your own words and describe how to find it'"),
        ],
    ),
    _phy_lesson(
        3, "States of Equilibrium",
        "Why do some objects return to their original position after being tilted, while others fall?",
        "- Define stable, unstable, and neutral equilibrium\n- Explain why the position of the centre of gravity and the size of the base determine stability\n- Classify everyday objects by their state of equilibrium",
        "- Demonstrate all three states of equilibrium using simple objects\n- Relate stability to COG height and base area\n- Apply criteria to classify objects as stable, unstable, or neutral",
        "- Appreciate safety implications of stability in buildings, vehicles, and furniture\n- Show systematic thinking when classifying objects",
        "Students investigate the three states of equilibrium and connect stability rules to real-life objects.",
        "Wooden blocks of different base sizes; cone; ball; cylinder; ruler; pivot; photographs of vehicles and buildings",
        "No significant hazards; ensure blocks do not topple onto hands.",
        [
            _pr("Tilt a wide-based block slightly and release; then tilt a narrow-based block — observe and record what happens",
                "Wide-based and narrow-based wooden blocks; flat surface",
                "Ask: 'What will happen when I release this?' before each demonstration; elicit predictions",
                "Discuss: 'What is different about the two blocks? What causes one to return and the other to fall?'",
                "Oral questioning: can students describe what they observed using 'COG' and 'base'?"),
            _pr("Demonstrate neutral equilibrium using a ball on a flat surface; compare with stable and unstable examples",
                "Ball; cone (balanced on tip); cylinder on its side",
                "Guide students to see: stable returns, unstable topples, neutral stays in new position",
                "Three-way sort: students categorise given objects into stable/unstable/neutral with justification",
                "Exit ticket: sketch three objects and label each state; explain using COG and base area"),
            _pr("Investigate: how does base width affect stability? Change base area while keeping height constant",
                "Cardboard boxes with variable base inserts; weights to add to tops",
                "Challenge: 'Find the minimum base width needed to keep this object stable when tilted 20°'",
                "Groups share findings; build rule: wider base + lower COG = more stable",
                "Record rule in notebooks; teacher checks phrasing for scientific accuracy"),
            _pr("Examine photographs: identify states of equilibrium in vehicles, furniture, cranes, and human posture",
                "Printed photographs or projected images of wide-based trucks, tall cranes, gymnasts",
                "Prompt students to apply stability rules to each real-life case",
                "Class discussion: 'Why do double-decker buses not topple? Why do gymnasts lower their bodies?'",
                "Peer assessment: students swap notebooks and check each other's classifications"),
            _pr("Update DQB and model: add stability conditions to the initial balance diagram",
                "DQB poster; initial model diagrams",
                "Prompt: 'How does what we learned today help answer the driving question about toppling?'",
                "Students add sticky notes to DQB marking evidence gathered",
                "Review DQB as class; teacher notes which sub-questions have been answered"),
        ],
    ),
    _phy_lesson(
        4, "Moment of a Force",
        "What determines how easily a force can rotate an object about a pivot?",
        "- Define moment of a force (torque) as force × perpendicular distance from pivot\n- State the unit of moment (N·m)\n- Explain that increasing force or distance increases the turning effect",
        "- Calculate the moment of a force given force and perpendicular distance\n- Design investigations to show how distance affects turning effect\n- Construct and interpret moment equations",
        "- Appreciate the practical value of moments in tools (spanners, levers, door handles)\n- Show accuracy and care when measuring distances",
        "Students discover that turning effect depends on both force magnitude and distance from pivot through hands-on ruler investigations.",
        "Metre rule; pivot (knife edge or pencil); known weights (100 g, 200 g, 500 g); ruler for measuring distances; graph paper",
        "Ensure pivot is stable; weights should not be dropped.",
        [
            _pr("Attach a 200 g weight at different distances from pivot; observe and record turning effect",
                "Metre rule; knife-edge pivot; 200 g weight; ruler",
                "Demonstrate setup; ask: 'Where does the weight have the greatest turning effect?'",
                "Students record force, distance, and qualitative turning effect in results table",
                "Can students identify the trend: further = greater turning? Ask for explanation."),
            _pr("Calculate moment = force × distance for each position; record in table; identify pattern",
                "Results table; calculator",
                "Introduce formula: Moment (N·m) = F (N) × d (m); work through first example together",
                "Students complete calculations; compare values across positions",
                "Quick quiz: given F and d, calculate moment; given moment and d, find F"),
            _pr("Investigate: does the same moment result from different F × d combinations? (e.g., 2 N at 0.5 m = 4 N at 0.25 m)",
                "Various weights; metre rule and pivot",
                "Students design their own F × d combinations to achieve a target moment of 0.4 N·m",
                "Groups share their combinations; confirm that F × d = constant for same turning effect",
                "Written explanation: 'Why can a long spanner remove a tight bolt more easily than a short one?'"),
            _pr("Apply moments to door handles and lever tools: calculate force needed at various distances",
                "Photographs or diagrams of spanner, crowbar, door handle at different positions",
                "Pose problem: 'A bolt requires 12 N·m to loosen. What force is needed with a 0.3 m spanner?'",
                "Students solve 2–3 application problems; peer-check answers",
                "Exit ticket: word problem requiring moment calculation"),
            _pr("Update model: add moment arrows to the balance diagram; label F and d",
                "Initial model diagrams; ruler; pencil",
                "Prompt students to show on their diagrams where turning occurs and how it is calculated",
                "Pair-share updated models; identify improvements",
                "Collect updated models; assess whether moment concept is correctly applied"),
        ],
    ),
    _phy_lesson(
        5, "Principle of Moments",
        "What conditions must be met for an object to be perfectly balanced by unequal forces?",
        "- State the principle of moments: sum of clockwise moments = sum of anticlockwise moments for equilibrium\n- Apply the principle to solve problems involving one and two pivots",
        "- Set up a balanced beam with unequal weights at different distances\n- Calculate unknown forces or distances using the principle of moments\n- Verify the principle experimentally",
        "- Value precision in measurement as the key to verifying a scientific principle\n- Appreciate how the principle underpins weighing scales and balance beams",
        "Students discover the principle of moments by balancing unequal loads on a metre rule and verify it mathematically.",
        "Metre rule; knife-edge pivot; weights (100 g, 200 g, 500 g); ruler; string for suspending weights",
        "Ensure pivot is stable; do not lean over the balanced beam.",
        [
            _pr("Balance a 200 g weight on one side with a 400 g weight on the other; measure required distances",
                "Metre rule; knife-edge pivot; 200 g and 400 g weights",
                "Ask students to predict distances before experimenting; then test predictions",
                "Record clockwise moment and anticlockwise moment for each trial; compare",
                "Can students articulate why the heavier weight must be closer to the pivot?"),
            _pr("Tabulate results for 5 different weight combinations; calculate both moments; compare",
                "Results table; calculator",
                "Guide discovery: 'What do you notice about the two moment values each time?'",
                "Students state the principle of moments in their own words",
                "Quick written check: state the principle of moments using correct terminology"),
            _pr("Apply principle to solve: unknown weight problems and unknown distance problems",
                "Worksheet with 4 structured problems",
                "Model one example of each type; students complete remaining problems independently",
                "Peer-check solutions; discuss common errors",
                "Formative quiz: 3 problems covering F×d = F×d in varied contexts"),
            _pr("Investigate moments about two supports: beam resting on two pivots with central load",
                "Metre rule; two supports at different positions; central weight",
                "Challenge: 'If the beam has supports at 20 cm and 80 cm, and a 5 N weight is at 50 cm, find reaction forces'",
                "Groups work through guided steps; share solutions",
                "Exit ticket: sketch a balanced beam with two supports; label forces and distances"),
            _pr("Update DQB and model: show that balanced = equal moments; link to weighing scales",
                "DQB poster; model diagrams; photographs of a beam balance",
                "Prompt: 'How does a beam balance use the principle of moments to measure mass?'",
                "Students add new evidence to DQB; revise models to show moment arrows and balance condition",
                "Teacher reviews DQB with class; confirm which questions are now answered"),
        ],
    ),
    _phy_lesson(
        6, "Torque, Couple, and Applications of Stability",
        "How do engineers apply knowledge of moments and stability to design safe structures and tools?",
        "- Define torque as the turning effect of a force and couple as two equal, opposite, non-collinear forces\n- Explain how stability principles are applied in vehicles, buildings, and everyday tools\n- Construct a complete final model explaining the phenomenon",
        "- Identify couples in everyday objects (steering wheel, bottle cap, spanner)\n- Analyse stability of objects using COG and base area\n- Construct a complete written explanation linking all unit concepts",
        "- Appreciate the role of physics in engineering and safety design\n- Show confidence in applying unit concepts to unfamiliar real-life situations",
        "Consolidation lesson: students connect torque, couple, moments, and stability to solve real engineering problems and construct their final explanation.",
        "Photographs/models of wheelbarrow, vehicle, crane, steering wheel; student Summary Tables; DQB",
        "No specific hazards.",
        [
            _pr("Examine a steering wheel diagram and a bottle cap: identify the two forces forming a couple",
                "Diagram of steering wheel; bottle cap; pair of equal and opposite force arrows",
                "Introduce 'couple': two equal forces in opposite directions with different lines of action; no net force but net torque",
                "Students sketch their own couple example from daily life and label forces",
                "Can students distinguish between a single moment and a couple? Oral probe."),
            _pr("Analyse stability of given objects: double-decker bus, crane, wheelbarrow, standing human",
                "Photographs of each; worksheets with COG diagrams",
                "Guide analysis: 'Locate the COG, identify the base, determine if the object is stable under stated conditions'",
                "Groups present analysis to class; class gives feedback using stability criteria",
                "Written assessment: analyse a new object not previously discussed"),
            _pr("Revisit the driving question: 'Why do some objects balance easily while others topple?' — write a class explanation",
                "DQB poster; all model diagrams; Summary Table drafts",
                "Facilitate class construction of a complete explanation citing unit evidence",
                "Students evaluate the class explanation against the unit SLOs",
                "Individual written task: draft personal Final Explanation paragraph connecting COG, stability, moments, and couples"),
            _pr("Review Summary Table: ensure all 6 lessons are recorded with evidence and connections",
                "Individual Summary Tables",
                "Students complete any missing rows; teacher circulates to check accuracy",
                "Peer review Summary Tables for completeness and scientific accuracy",
                "Teacher reviews sample Summary Tables; provide written feedback"),
            _pr("Reflect on unit: what new understanding have you built? How will you approach the Final Explanation?",
                "Final Explanation task description; rubric",
                "Introduce the Final Explanation document and rubric; students self-assess readiness",
                "Class discussion: which section of the Final Explanation will be most challenging and why?",
                "Exit ticket: 'The most important thing I learned in this unit is… because…'"),
        ],
    ),
]

_PHY_SUMMARY_ROWS = [
    (1, "Observed balance and toppling; drew initial model; posted questions on DQB",
        "All objects have a balance point; moving a load changes balance",
        "Introduces the phenomenon: why do objects topple?",
        "DQB created; initial model drawn showing balance point"),
    (2, "Located COG of regular and irregular shapes by balancing and plumb-line method",
        "COG is where all weight appears to act; can be outside the object for irregular shapes",
        "COG determines where an object balances — key to understanding toppling",
        "Model updated: COG marked on balance diagram"),
    (3, "Demonstrated stable, unstable, and neutral equilibrium; applied criteria to real objects",
        "Stability depends on COG height and base area; wider base and lower COG = more stable",
        "Explains why wide-based objects are stable and tall narrow ones topple",
        "DQB: evidence added; model updated with stability conditions"),
    (4, "Investigated turning effect at different distances; calculated moment = F × d",
        "Moment depends on force AND distance; longer lever = greater turning effect",
        "Turning effect of forces drives the physics of balance and toppling",
        "Model updated: moment arrows added with F and d labels"),
    (5, "Balanced unequal weights; verified principle of moments experimentally",
        "For equilibrium: total clockwise moments = total anticlockwise moments",
        "Principle of moments explains how unequal forces can maintain balance",
        "Model updated: moment balance condition shown; DQB sub-questions answered"),
    (6, "Identified couples; analysed stability of real objects; drafted Final Explanation",
        "Torque, couple, COG, and stability principles all link to explain the phenomenon",
        "Complete understanding: objects balance or topple depending on COG, base, and moments",
        "Final model complete; DQB fully answered; Summary Table reviewed"),
]

_PHY_FE_SECTIONS = [
    ("1. Centre of Gravity and Stability",
     "Explain what centre of gravity is, how it is determined, and how it relates to stability. "
     "Include the three states of equilibrium with diagrams.\n\n"
     "Guiding Questions:\n- How do you find the COG of a regular shape? An irregular shape?\n"
     "- What is the difference between stable, unstable, and neutral equilibrium?\n"
     "- How do COG position and base area determine stability?",
     "Every object has a centre of gravity — the point where its entire weight appears to act. "
     "For regular shapes this lies at the centre of symmetry; for irregular shapes it is found "
     "experimentally using the plumb-line method. An object is stable if, when slightly tilted, "
     "the COG falls back inside the base and the object returns to its original position. "
     "Unstable equilibrium occurs when any tilt moves the COG outside the base, causing toppling. "
     "Neutral equilibrium occurs when the COG stays at the same height regardless of position. "
     "Wider base and lower COG both increase stability."),
    ("2. Moment of a Force",
     "Define moment of a force and explain what factors affect its magnitude. Include worked calculations.\n\n"
     "Guiding Questions:\n- What is the formula for moment of a force?\n"
     "- How does increasing distance from the pivot affect the turning effect?\n"
     "- Give two real-life examples where moments are used to make work easier.",
     "The moment of a force is the turning effect produced by a force about a pivot. "
     "It is calculated as: Moment (N·m) = Force (N) × Perpendicular distance from pivot (m). "
     "Increasing either the force or the distance increases the moment. "
     "This is why a long spanner can loosen a tight bolt more easily than a short one — "
     "the same force applied at greater distance produces a larger moment. "
     "Door handles are placed far from the hinges for the same reason."),
    ("3. Principle of Moments",
     "State the principle of moments and apply it to solve a balance problem. Show your working.\n\n"
     "Guiding Questions:\n- State the principle of moments in words.\n"
     "- Show how the principle explains why a lighter person can balance a heavier person on a see-saw.\n"
     "- Solve: a 300 N force acts 0.4 m from a pivot. What force at 0.6 m balances it?",
     "For an object in equilibrium, the sum of all clockwise moments about any pivot equals "
     "the sum of all anticlockwise moments. This is the principle of moments. "
     "On a see-saw, a lighter person can balance a heavier one by sitting further from the pivot, "
     "increasing their moment. In beam balances, the principle allows unknown masses to be "
     "determined by adjusting distances. Solution: 300 × 0.4 = F × 0.6 → F = 120 ÷ 0.6 = 200 N."),
    ("4. Torque, Couple, and Engineering Applications",
     "Explain torque and couple. Describe how engineers use moments and stability principles in design.\n\n"
     "Guiding Questions:\n- What is a couple and how does it differ from a single moment?\n"
     "- Why are double-decker buses designed with heavy engines at the base?\n"
     "- Describe one engineering application of the principle of moments.",
     "Torque is the rotational effect of a force about an axis. A couple consists of two equal, "
     "opposite, non-collinear forces that produce rotation without translation. Steering wheels "
     "and bottle caps are opened using couples. Engineers lower the centre of gravity of vehicles "
     "by placing heavy components (engines, batteries) near the ground, widening the wheelbase, "
     "and designing low-profile bodies — all to increase the range of tilt before the COG falls "
     "outside the base. Cranes counterbalance heavy loads using the principle of moments."),
    ("5. Phenomenon Explanation",
     "Answer the driving question using all evidence gathered in this unit.\n\n"
     "Driving Question: Why do some objects balance easily while others topple, even when they seem similar?\n\n"
     "Use concepts: centre of gravity, states of equilibrium, moment of a force, principle of moments, couple.",
     "Objects balance or topple depending on the position of their centre of gravity relative to "
     "their base of support. When an object is tilted, if the vertical line through the COG "
     "falls within the base, the resulting moment restores the object to equilibrium. If the COG "
     "falls outside the base, the unbalanced moment causes toppling. Objects with a low COG and "
     "wide base are more stable. Two objects can look similar but have different COG heights due "
     "to the distribution of their mass — this is why a hollow cylinder tips more easily than a "
     "solid one of the same dimensions. Engineers exploit these principles in every structure "
     "they design: from wide-based stools and vehicles with lowered chassis to cranes whose "
     "counterweights apply equal and opposite moments to maintain equilibrium."),
]

_PHY_META = {
    "subject": "Physics",
    "grade": "Grade 10",
    "substrand": "Sub-Strand 1.5: Moments and Equilibrium",
    "lessons": 6,
    "driving_question": "Why do some objects balance easily while others topple, even when they seem similar?",
    "substrand_overview_rows": [
        ("Phenomenon", "A wheelbarrow, bicycle, or person carrying a heavy load unevenly — why do some objects balance while others topple?"),
        ("Driving Question", "Why do some objects balance easily while others topple, even when they seem similar?"),
        ("Sub-Strand SLOs — Knowledge", _PHY_SLO_K),
        ("Sub-Strand SLOs — Skills", _PHY_SLO_S),
        ("Sub-Strand SLOs — Attitudes/Values", _PHY_SLO_A),
        ("Number of Lessons", "6 lessons (12 periods of 40 minutes each)"),
        ("Core Competencies", "Communication and Collaboration · Critical Thinking and Problem Solving · Self-efficacy · Digital Literacy · Learning to Learn"),
        ("Core Values", "Curiosity · Responsibility · Integrity · Excellence · Teamwork"),
        ("Assessment", "Formative: observation, oral questioning, calculations, practical reports, summary table\nSummative: Final Explanation document"),
        ("Resources", "Metre rule; knife-edge pivot; weights (100–500 g); cardboard cut-outs; plumb bob; Khan Academy — Torque and Balance; PhET Balancing Act simulation"),
    ],
    "summary_table_headers": [
        "Lesson", "What I Did", "Key Understanding Built",
        "Connection to Driving Question", "Model / DQB Evolution"
    ],
    "summary_table_rows": _PHY_SUMMARY_ROWS,
    "summary_instructions": (
        "Use this table to track your learning across all 6 lessons. "
        "For each lesson record: what you did, what you understood, how it connects to the driving question, "
        "and how your model and DQB evolved. Your completed Summary Table is evidence of your scientific thinking."
    ),
    "final_explanation_sections": _PHY_FE_SECTIONS,
    "final_explanation_rubric": {
        "headers": ["Criterion", "Beginning (1)", "Developing (2)", "Proficient (3)", "Advanced (4)"],
        "rows": [
            ("1. Centre of Gravity",
             "Cannot define COG or locate it",
             "Defines COG but cannot explain how to find it experimentally",
             "Correctly defines COG, describes both methods, and links to stability",
             "Explains COG with labelled diagrams; connects to phenomenon with specific examples"),
            ("2. States of Equilibrium",
             "Cannot distinguish stable from unstable",
             "Names the three states but explanations are incomplete",
             "Correctly explains all three states with COG and base area criteria",
             "Applies all three states to novel real-life examples with full reasoning"),
            ("3. Moments and Principle",
             "Cannot define moment or apply formula",
             "States formula but makes errors in calculation or application",
             "Correctly calculates moments and applies the principle of moments to problems",
             "Solves multi-step problems; explains principle with worked examples and diagrams"),
            ("4. Torque, Couple and Applications",
             "Cannot describe torque or couple",
             "Describes torque but confuses couple with a single moment",
             "Correctly defines couple; gives two real-life engineering applications",
             "Analyses engineering designs using torque, couple, COG, and stability with quantitative reasoning"),
            ("5. Phenomenon Explanation",
             "Cannot connect unit concepts to the driving question",
             "Makes partial connection but misses key factors",
             "Explains the phenomenon using COG, stability, and moments correctly",
             "Complete evidence-based explanation covering all 5 objectives; proposes novel engineering solution"),
        ],
    },
    "lessons_data": _PHY_LESSONS,
}

# ═══════════════════════════════════════════════════════════════════════════════
# META — Chemistry: The Periodicity (7 lessons)
# ═══════════════════════════════════════════════════════════════════════════════

_CHEM_SLO_K = (
    "a) Describe the trend in physical properties of elements of the periodic table\n"
    "b) Investigate the chemical properties of elements in groups of the periodic table\n"
    "c) Describe the trend in properties across a period\n"
    "d) Outline applications of elements of the periodic table\n"
    "e) Appreciate the application of various elements of the periodic table"
)
_CHEM_SLO_S = (
    "- Draw and interpret models of atomic structure for elements in Groups I, II, VII, and VIII\n"
    "- Investigate chemical properties of Group I and Group II elements through demonstration and video\n"
    "- Construct summary tables showing trends in physical and chemical properties\n"
    "- Apply knowledge of periodicity to predict behaviour of unfamiliar elements"
)
_CHEM_SLO_A = (
    "- Appreciate the elegance of patterns in the periodic table\n"
    "- Show curiosity when observing chemical reactions and comparing element behaviour\n"
    "- Value safe, careful handling of reactive substances\n"
    "- Recognise real-world applications of periodic table knowledge (medicine, environment, industry)"
)

def _chem_lesson(num, title, inquiry, know, skills, attit, purpose, materials, safety, p1):
    return {
        "number": num, "title": title,
        "inquiry_question": inquiry,
        "slo_knowledge": know, "slo_skills": skills, "slo_attitudes": attit,
        "overview_purpose": purpose,
        "materials": materials, "safety": safety,
        "period1_heading": "Period 1 (40 minutes)",
        "period1_table": p1,
        "period2_heading": "Period 2 (40 minutes)",
        "period2_table": [],
        "reflections": [],
    }

_CHEM_LESSONS = [
    _chem_lesson(
        1, "Introduction to Periodicity",
        "Why do elements show patterns in their properties across groups and periods, and how can we use these patterns to predict the behaviour of unfamiliar elements?",
        "- Recall the structure of the periodic table: groups, periods, metals, non-metals\n- Identify that elements in the same group have the same number of outer electrons\n- State the driving question for the unit",
        "- Locate elements on the periodic table and identify their group and period\n- Draw a discussion diamond recording initial ideas about why elements show patterns",
        "- Show curiosity about the order underlying the periodic table\n- Appreciate that patterns in science allow us to make predictions",
        "Launches the unit phenomenon. Students observe reactions of sodium in water and burning magnesium, generate questions, and record initial ideas about why elements behave differently.",
        "Periodic table poster or printed copies; video clip of sodium reacting with water; video clip of potassium in water; magnesium ribbon; Bunsen burner; Driving Question Board; discussion diamond worksheets",
        "Teacher demonstration only for sodium/potassium in water and burning magnesium. Students observe from a safe distance. Safety goggles required.",
        [
            _pr("Watch video: sodium reacting with water — record observations; predict what potassium will do",
                "Video of sodium in water (Khan Academy or downloaded); potassium in water video",
                "Play sodium video without introduction; ask 'What do you observe?' before revealing what the substance is",
                "Discussion diamond: students write individual observations, then share and reach a group consensus",
                "Exit ticket: 'What do you predict will happen when potassium reacts with water? Why?'"),
            _pr("Watch teacher demonstrate burning magnesium ribbon; observe flame colour and product",
                "Magnesium ribbon; Bunsen burner; tongs; dark card for observing flame",
                "Safety briefing first; demonstrate; students record observations of flame colour, brightness, and ash produced",
                "Class discussion: 'What product formed? Write the word equation.' Introduce the idea of metals forming oxides",
                "Oral check: can students write the word equation for magnesium burning in air?"),
            _pr("Place sodium, potassium, and magnesium on the periodic table; identify their groups and electron configurations",
                "Periodic table; atomic structure diagrams for Na, K, Mg",
                "Guide students to draw electron shells for Na (2,8,1) and K (2,8,8,1); ask 'What do they have in common?'",
                "Class discussion: same number of outer electrons → same group → similar chemical behaviour",
                "Quick quiz: given atomic number, draw electron configuration and identify group"),
            _pr("Create the Driving Question Board: post initial questions about element patterns and uses",
                "DQB poster; sticky notes; markers",
                "Model 'I wonder why…' and 'I notice that…' sentence starters; groups post questions by theme",
                "Read all questions aloud; categorise (physical properties, chemical reactions, applications)",
                "Count questions per category to inform lesson sequence; photograph DQB for record"),
            _pr("Draw 'My Initial Model': how do you think the properties of elements change across the periodic table?",
                "Blank A4 paper; periodic table reference",
                "Prompt: 'Draw or write your current thinking about what changes and what stays the same across a period or group'",
                "Pair-share models; note common ideas and disagreements",
                "Collect initial models; note misconceptions to address in future lessons"),
        ],
    ),
    _chem_lesson(
        2, "Gradation in Physical Properties",
        "How do physical properties such as atomic size, melting point, and ionisation energy change as you move down a group or across a period?",
        "- Define atomic radius, ionisation energy, and electron affinity\n- State the trend in atomic radius down a group and across a period\n- Explain why atomic radius increases down a group (more electron shells) and decreases across a period (increasing nuclear charge)",
        "- Draw atomic models for elements in Group I (Li, Na, K) showing increasing shell number\n- Construct a comparison table of physical properties for Group I and Period 3 elements\n- Interpret graphs of melting point and ionisation energy across Period 3",
        "- Show appreciation for how microscopic atomic structure determines macroscopic properties\n- Value careful observation when comparing trends in data",
        "Students investigate how physical properties change systematically across the periodic table by drawing atomic models and analysing data.",
        "Periodic table; atomic radius data table; melting point and ionisation energy graph for Period 3; coloured pencils; graph paper or projected graphs",
        "No hazardous materials in this lesson.",
        [
            _pr("Draw atomic models for Li, Na, K — count electron shells; observe increasing size",
                "Atomic model diagrams (Bohr model); coloured pencils; periodic table",
                "Guide students to draw Li (2,1), Na (2,8,1), K (2,8,8,1); ask 'What changes as you go down the group?'",
                "Class discussion: more shells → larger atom → lower ionisation energy (outer electron further from nucleus)",
                "Quick sketch check: can students draw the model for Rb (2,8,18,8,1) without being told?"),
            _pr("Compare atomic radii across Period 3 (Na to Ar) using data table; plot graph",
                "Atomic radius data for Period 3; graph paper",
                "Students plot atomic radius vs. atomic number across Period 3; identify the trend",
                "Discussion: nuclear charge increases across period → pulls electrons closer → radius decreases",
                "Exit ticket: explain in one sentence why atomic radius decreases across a period"),
            _pr("Examine ionisation energy graph for Period 3 — identify the trend and explain peaks for Mg and P",
                "Projected or printed graph of first ionisation energy vs. atomic number for Period 3",
                "Guided questions: 'Where are the peaks? Why does Al have lower IE than Mg?' (sub-shell explanation for context)",
                "Students annotate the graph with explanations for key points",
                "Discussion: connect ionisation energy trend to reactivity — low IE = more reactive metal"),
            _pr("Construct a summary comparison table: Group I properties (Li, Na, K, Rb) vs. Group II (Be, Mg, Ca)",
                "Data cards with melting point, density, ionisation energy for each element",
                "Students complete the table; identify trends down each group",
                "Pair discussion: 'What pattern do you notice? Can you predict the properties of caesium?'",
                "Formative assessment: predict two properties of strontium (between Ca and Ba) using trends"),
            _pr("Modify initial model: update atomic diagrams to show trends in size and ionisation energy",
                "Initial model from Lesson 1",
                "Prompt: 'What can you add or change in your model now that you know about atomic radius trends?'",
                "Share updated models; check that students show increasing size down a group",
                "Collect updated models; assess quality of trend representation"),
        ],
    ),
    _chem_lesson(
        3, "Chemical Properties of Group I Elements",
        "How do the chemical properties of alkali metals change as you move down Group I, and what explains this pattern?",
        "- Describe the reactions of Group I metals with oxygen, cold water, and dilute acids\n- Explain that reactivity increases down Group I due to decreasing ionisation energy\n- Write word and balanced symbol equations for Group I reactions",
        "- Write word equations for reactions of Li, Na, and K with water\n- Balance symbol equations for Group I reactions\n- Predict the behaviour of rubidium based on Group I trends",
        "- Show safe curiosity when observing reactive metal demonstrations\n- Appreciate that reactivity trends allow chemists to handle substances safely",
        "Students investigate Group I reactivity trends through teacher demonstrations, videos, and equation writing.",
        "Video clips of Li, Na, K reacting with water; phenolphthalein indicator; trough of water; sodium (teacher demonstration only); safety screen; periodic table",
        "All Group I metal reactions must be teacher demonstration only. Safety screen required. Students must not handle sodium or potassium. Goggles must be worn.",
        [
            _pr("Watch videos of Li, Na, and K reacting with water in sequence; record observations for each",
                "Videos of Li, Na, K in water (Khan Academy or downloaded); observation table",
                "Pause after each video; prompt: 'What did you observe? How was this different from the previous reaction?'",
                "Discussion: violence of reaction increases Li → Na → K; link to decreasing ionisation energy",
                "Exit ticket: rank Li, Na, K in order of reactivity and give one piece of evidence"),
            _pr("Write word equations for each Group I metal reacting with water; identify the products",
                "Whiteboard; equation writing guide",
                "Model: sodium + water → sodium hydroxide + hydrogen; students write Li and K equations",
                "Introduce phenolphthalein: demonstrate that the product is alkaline (pink colour change)",
                "Oral probe: 'What two products always form when a Group I metal reacts with water?'"),
            _pr("Balance symbol equations for Group I reactions with water and with oxygen",
                "Symbol equation worksheets; periodic table; PhET balancing equations simulation",
                "Work through balancing 2Na + 2H₂O → 2NaOH + H₂ step by step; students balance remaining equations",
                "Peer-check balanced equations; discuss common errors (forgetting to balance H₂)",
                "Formative quiz: balance 3 Group I equations independently"),
            _pr("Predict the behaviour of rubidium and caesium using Group I trends",
                "Periodic table; Group I data table from Lesson 2",
                "Ask: 'Based on what you know about Li, Na, K — what will Rb do with water? Why?'",
                "Watch short video of Rb in water; compare prediction with observation",
                "Written prediction: 'I predict caesium will… because…' — assess reasoning quality"),
            _pr("Update DQB and model: add evidence about Group I reactivity trend",
                "DQB poster; student models",
                "Prompt: 'Which DQB questions can we now answer? Add your evidence.'",
                "Students annotate models with Group I reactivity arrows and equations",
                "Review DQB with class; mark answered questions"),
        ],
    ),
    _chem_lesson(
        4, "Chemical Properties of Group II Elements",
        "How do Group II metals compare with Group I in reactivity, and what explains the differences?",
        "- Describe the reactions of Group II metals with oxygen, water, and dilute acids\n- Compare reactivity of Group I and Group II metals\n- Explain why Group II metals are less reactive than Group I (higher ionisation energy, two electrons to lose)",
        "- Write and balance equations for Group II reactions with water and acids\n- Construct a comparison table of Group I vs. Group II reactivity\n- Apply trends to predict behaviour of barium",
        "- Appreciate that small differences in electron configuration produce significant differences in chemical behaviour\n- Show precision when writing and balancing chemical equations",
        "Students investigate Group II reactivity by comparing it with Group I data and through demonstrations of magnesium reactions.",
        "Magnesium ribbon; dilute hydrochloric acid; test tube; trough of water; Bunsen burner; periodic table; Group I and II comparison table",
        "Dilute acid use: teacher demonstration for magnesium in acid. Students may observe burning magnesium from a safe distance. Goggles required.",
        [
            _pr("Teacher demonstration: magnesium ribbon reacting with steam (hot water) — observe and compare to sodium in cold water",
                "Magnesium ribbon; hot water or steam generator; retort stand; Bunsen burner",
                "Ask: 'Why do you think Mg reacts with steam but not cold water, while Na reacts vigorously with cold water?'",
                "Discussion: Group II has two outer electrons; higher ionisation energy → less reactive",
                "Oral check: students explain the difference between Group I and Group II reactivity in their own words"),
            _pr("Teacher demonstration: magnesium with dilute HCl — observe hydrogen gas production",
                "Magnesium ribbon; dilute hydrochloric acid; test tube; wooden splint",
                "Demonstrate the squeaky-pop test for hydrogen; students record word equation",
                "Students write: magnesium + hydrochloric acid → magnesium chloride + hydrogen",
                "Quick calculation: if 0.24 g Mg reacts, how many moles of H₂ are produced? (extension)"),
            _pr("Construct comparison table: Group I vs. Group II — reaction with oxygen, water, acids",
                "Data cards; comparison table template; textbook or reference sheet",
                "Students complete the table; identify similarities (both form basic oxides) and differences (reactivity level)",
                "Pair discussion: 'If Ca is in Group II, predict how it reacts with water compared to Mg and Na'",
                "Exit ticket: give two differences and one similarity between Group I and Group II chemical properties"),
            _pr("Balance symbol equations for Group II reactions: Mg + O₂, Ca + H₂O, Mg + 2HCl",
                "Equation balancing worksheet",
                "Guided practice: work through Mg + O₂ → MgO; students complete Ca and Mg equations",
                "Peer-check; discuss why coefficients must be whole numbers",
                "Formative quiz: balance 3 Group II equations independently"),
            _pr("Update summary table and model: add Group II reactivity data alongside Group I",
                "Summary comparison table; student models",
                "Prompt: 'How has your understanding of 'why elements in the same group behave similarly' grown?'",
                "Students revise their models to show Group I and II side by side",
                "Class discussion: which DQB questions are now answered?"),
        ],
    ),
    _chem_lesson(
        5, "Properties of Halogens (Group VII)",
        "Why do the physical and chemical properties of halogens change down the group, and how does this explain their reactivity and uses?",
        "- Describe the physical states and colours of fluorine, chlorine, bromine, and iodine at room temperature\n- Explain that halogen reactivity decreases down Group VII due to increasing atomic size\n- Describe displacement reactions of halogens",
        "- Observe and record the colours and states of halogens from video/images\n- Write equations for halogen displacement reactions\n- Apply reactivity trend to predict displacement outcomes",
        "- Appreciate the importance of halogens in water purification, medicine, and the food industry\n- Show awareness of the hazardous nature of halogens and the need for safety precautions",
        "Students investigate the Group VII trend in physical and chemical properties, focusing on displacement reactions as evidence of reactivity order.",
        "Video clips of fluorine, chlorine, bromine, iodine reactions; halogen colour chart; displacement reaction diagrams; periodic table; safety data for halogens",
        "Halogens are toxic — all demonstrations are teacher-led or video-based. No student handling of halogen solutions without close supervision. Goggles and ventilation required for any live demonstration.",
        [
            _pr("Observe images/video of F₂, Cl₂, Br₂, I₂ — record colour, state, and smell (described, not experienced)",
                "Projected images of halogen colours; video showing chlorine as yellow-green gas, bromine as red-brown liquid, iodine as dark solid",
                "Present each halogen one at a time; ask students to record observations in a table",
                "Discussion: physical properties change systematically — state: gas → gas → liquid → solid down the group",
                "Exit ticket: complete a table of halogen states and colours from memory"),
            _pr("Compare halogen reactivity: watch video of Cl₂ displacing Br⁻ from solution; predict I₂ displacement",
                "Video of chlorine water displacing bromide ions (solution turns orange); displacement reaction diagrams",
                "Ask: 'If Cl₂ displaces Br⁻, can Br₂ displace Cl⁻? Why or why not?' — elicit prior reasoning",
                "Discussion: more reactive halogen displaces less reactive one; reactivity decreases down the group",
                "Oral probe: 'Will iodine displace chloride from sodium chloride solution? Explain your answer.'"),
            _pr("Write displacement reaction equations: Cl₂ + 2KBr → 2KCl + Br₂; and others",
                "Displacement reaction worksheet",
                "Model the first equation; students write Cl₂ + KI and Br₂ + KI equations",
                "Peer-check equations; discuss which halogens can and cannot displace which halides",
                "Formative quiz: predict and write equation for Br₂ + NaI"),
            _pr("Research and present: one real-life application of each halogen (chlorine in water, iodine in medicine, fluorine in toothpaste)",
                "Brief reading cards on halogen applications; textbook section on uses",
                "Groups each take one halogen; prepare a 2-minute explanation of its application and why its properties make it suitable",
                "Groups present; class gives one piece of feedback per presentation",
                "Written record: each student notes all four applications in their notebooks"),
            _pr("Update model and DQB: how does electron configuration explain decreasing halogen reactivity?",
                "Atomic model diagrams for F, Cl, Br, I; DQB poster",
                "Guide students to connect: larger atom → outer shell further from nucleus → less attraction for electrons → less reactive",
                "Students add halogen reactivity arrows to their models",
                "Review DQB; add new questions raised by the halogen lesson"),
        ],
    ),
    _chem_lesson(
        6, "Noble Gases and Period 3 Elements — Part 1",
        "What makes noble gases unreactive, and how do the properties of Period 3 elements change from left to right?",
        "- Explain why noble gases are unreactive (full outer shell of electrons)\n- Describe the trend in metallic character across Period 3 (Na to Cl)\n- State the physical properties of Na, Mg, Al, Si, P, S, Cl across Period 3",
        "- Locate Period 3 elements on the periodic table and identify their electron configurations\n- Construct a table comparing physical properties across Period 3\n- Distinguish between metals, metalloids, and non-metals in Period 3",
        "- Appreciate that stability of full electron shells is a key principle in chemistry\n- Show systematic thinking when comparing a sequence of elements",
        "Students explore why noble gases are unique and begin investigating the systematic changes in properties across Period 3.",
        "Periodic table; atomic model diagrams for Period 3 elements; physical properties data table; coloured pencils",
        "No hazardous materials in this lesson.",
        [
            _pr("Draw electron configurations for He (2), Ne (2,8), Ar (2,8,8) — identify the full outer shell pattern",
                "Bohr model diagrams; periodic table",
                "Ask: 'What do all noble gases have in common in their electron configurations? Why does this make them stable?'",
                "Discussion: full outer shell = maximum stability = no tendency to gain, lose, or share electrons",
                "Quick quiz: draw electron configurations for Kr and Xe and explain why they are also unreactive"),
            _pr("Survey Period 3 elements: Na to Ar — record state at room temperature, conductivity, appearance",
                "Projected images of Period 3 elements; data table",
                "Students complete the data table from images and reference cards; identify the transition from metal to non-metal",
                "Discussion: Period 3 goes from highly reactive metal (Na) to reactive non-metal (Cl) to inert gas (Ar)",
                "Exit ticket: where in Period 3 is the transition from metal to non-metal? What element is the metalloid?"),
            _pr("Plot melting points of Period 3 elements vs. atomic number; explain the pattern",
                "Melting point data; graph paper or projected graph",
                "Students plot and identify: high melting points for metals (Na, Mg, Al), peak at Si (giant covalent), low for non-metals",
                "Guided discussion: type of bonding determines melting point; metallic bonding → moderate; giant covalent → very high; simple molecular → low",
                "Peer discussion: 'Why does sulphur have a higher melting point than phosphorus?'"),
            _pr("Compare Period 3 oxides: Na₂O and MgO are basic; Al₂O₃ is amphoteric; P₄O₁₀, SO₃, Cl₂O₇ are acidic",
                "Summary table of Period 3 oxides and their acid/base character",
                "Teacher-led introduction; students complete the table from reference information",
                "Discussion: trend from basic → amphoteric → acidic oxides reflects increasing non-metallic character across the period",
                "Written record: students write a sentence summarising the trend in Period 3 oxide character"),
            _pr("Update model: show the transition from metal to non-metal across Period 3 with electron configuration evidence",
                "Student model diagrams; periodic table",
                "Students draw Period 3 electron configs side by side and annotate with metallic/non-metallic character",
                "Pair discussion: 'How does this connect to the trends we saw in Group I and Group II?'",
                "Collect models; assess quality of cross-period comparison"),
        ],
    ),
    _chem_lesson(
        7, "Period 3 Elements and Compounds — Part 2 / Unit Consolidation",
        "How does knowledge of periodic trends allow us to predict the behaviour of unfamiliar elements and explain the phenomenon?",
        "- Summarise the trends in physical and chemical properties across Period 3 and down Groups I, II, and VII\n- Apply periodic trends to predict properties of unfamiliar elements\n- Construct a complete explanation of the unit driving question",
        "- Solve prediction problems using periodic trend knowledge\n- Complete the Summary Table for all 7 lessons\n- Draft the Final Explanation connecting all unit concepts to the driving question",
        "- Appreciate the predictive power of the periodic table\n- Show confidence in applying learned patterns to unfamiliar situations",
        "Consolidation lesson connecting all periodicity trends; students apply their understanding to new elements and construct their complete Final Explanation.",
        "Summary Table (student copies); DQB poster; periodic table; Final Explanation task description; rubric",
        "No hazardous materials.",
        [
            _pr("Prediction challenge: given element X (atomic number 19), predict its group, period, reactivity, reaction with water",
                "Periodic table; prediction worksheet",
                "Students work independently; then compare with partners",
                "Reveal element X is potassium; students check their predictions against known properties",
                "Oral debrief: which students used electron configuration? Which used period/group position?"),
            _pr("Revisit the driving question: 'Why do elements show patterns in their properties?' — construct class answer",
                "DQB poster; all evidence gathered across 7 lessons",
                "Facilitate class construction of a complete answer citing evidence from each lesson",
                "Students individually write a one-paragraph answer to the driving question",
                "Peer review paragraphs; give written feedback using the unit rubric criteria"),
            _pr("Complete the Summary Table: check all 7 lessons are recorded with evidence and connections",
                "Individual Summary Tables",
                "Students complete any missing rows; teacher circulates and checks accuracy",
                "Peer review for completeness and scientific accuracy",
                "Teacher reviews sample Summary Tables; provide written feedback"),
            _pr("Discuss real-world applications: link periodic trends to careers (medicine, water treatment, materials science)",
                "Brief reading cards on careers in chemistry using periodic table knowledge",
                "Students discuss: 'Which application is most important to Kenya? Why?'",
                "Groups each present one application to the class",
                "Written reflection: 'How has understanding periodicity changed how you see everyday materials?'"),
            _pr("Self-assess readiness for Final Explanation; review rubric criteria; identify target areas",
                "Final Explanation task description; rubric; student models and Summary Tables",
                "Students rate themselves on each rubric criterion; identify one area to strengthen",
                "Partner discussion: share target areas and suggest how to address them",
                "Exit ticket: 'The most important pattern I learned in this unit is… because…'"),
        ],
    ),
]

_CHEM_SUMMARY_ROWS = [
    (1, "Observed Na in water and burning Mg; drew initial model; posted DQB questions",
        "Elements in the same group have the same number of outer electrons, explaining similar behaviour",
        "Introduces the phenomenon: why do elements show patterns in properties?",
        "DQB created; initial model drawn showing atomic structure ideas"),
    (2, "Drew atomic models for Group I; plotted atomic radius and ionisation energy graphs for Period 3",
        "Atomic radius increases down a group; decreases across a period due to nuclear charge",
        "Physical property trends reflect underlying atomic structure changes",
        "Model updated: atomic size trends added"),
    (3, "Observed Group I reactions with water; wrote and balanced equations; predicted Rb/Cs behaviour",
        "Group I reactivity increases down the group as ionisation energy decreases",
        "Reactivity trends explain why K reacts more violently with water than Li",
        "Model updated: Group I reactivity arrow added; DQB questions answered"),
    (4, "Demonstrated Mg reactions; compared Group I and Group II reactivity",
        "Group II metals are less reactive than Group I due to higher ionisation energy",
        "Electron configuration differences between groups explain reactivity differences",
        "Comparison table completed; model updated with Group II data"),
    (5, "Observed halogen colours/states; investigated displacement reactions; researched applications",
        "Halogen reactivity decreases down Group VII as atomic size increases",
        "Displacement reactions are evidence of halogen reactivity order",
        "Model updated: halogen reactivity trend added; applications noted on DQB"),
    (6, "Drew noble gas configurations; surveyed Period 3 properties; plotted melting point graph",
        "Full outer shells make noble gases unreactive; Period 3 transitions from metal to non-metal",
        "Period 3 trend shows the full range from reactive metal to non-metal to inert gas",
        "Model updated: Period 3 electron configs drawn; metal/non-metal transition marked"),
    (7, "Solved prediction problems; constructed class answer to driving question; completed Summary Table",
        "Periodic trends allow prediction of unfamiliar element properties",
        "Complete understanding: electron configuration drives all observed trends",
        "Final model complete; DQB fully answered; Summary Table reviewed"),
]

_CHEM_FE_SECTIONS = [
    ("1. Trends in Physical Properties",
     "Describe how atomic radius, ionisation energy, and melting point change across Period 3 and down Group I. Include diagrams.\n\n"
     "Guiding Questions:\n- Why does atomic radius increase down a group?\n"
     "- Why does atomic radius decrease across a period?\n"
     "- What happens to melting point across Period 3 and why?",
     "Atomic radius increases down a group because each successive element has an additional electron shell, "
     "placing the outer electrons further from the nucleus. Across a period, the nuclear charge increases "
     "while the number of shells stays the same, so the nucleus pulls electrons closer, decreasing atomic radius. "
     "Ionisation energy follows the reverse trend: decreasing down a group (easier to remove an outer electron "
     "that is further away) and generally increasing across a period. Melting point across Period 3 peaks at "
     "silicon (giant covalent structure) and drops sharply for non-metals (simple molecular structures)."),
    ("2. Chemical Properties of Group I and Group II",
     "Compare the chemical properties of Group I and Group II metals. Include equations.\n\n"
     "Guiding Questions:\n- How does Group I reactivity change down the group? Why?\n"
     "- How do Group I and Group II metals differ in reactivity?\n"
     "- Write balanced equations for Na reacting with water and Mg reacting with steam.",
     "Group I metals become more reactive down the group as ionisation energy decreases — the outer electron "
     "is easier to remove from K than Na because it is in a higher energy shell further from the nucleus. "
     "Group II metals are less reactive than Group I because they must lose two electrons (higher ionisation energy). "
     "2Na + 2H₂O → 2NaOH + H₂ (vigorous, room temperature). "
     "Mg + H₂O (steam) → MgO + H₂ (requires steam; does not react with cold water)."),
    ("3. Properties of Halogens (Group VII)",
     "Describe the trend in halogen properties down Group VII. Include displacement reactions.\n\n"
     "Guiding Questions:\n- How do the state, colour, and reactivity of halogens change down Group VII?\n"
     "- Explain using atomic structure why reactivity decreases down the group.\n"
     "- Write the equation for chlorine displacing bromide ions from solution.",
     "Down Group VII: state changes from gas (F₂, Cl₂) to liquid (Br₂) to solid (I₂); colour darkens. "
     "Reactivity decreases because larger atoms have their outer shells further from the nucleus, "
     "reducing the attraction for electrons being gained. Fluorine is the most reactive non-metal. "
     "Cl₂ + 2KBr → 2KCl + Br₂ — chlorine displaces bromide because Cl is more reactive than Br. "
     "Br₂ cannot displace Cl⁻ from solution. This provides experimental evidence for the halogen reactivity order."),
    ("4. Period 3 Trends and Noble Gases",
     "Describe the trend in properties across Period 3 and explain why noble gases are unreactive.\n\n"
     "Guiding Questions:\n- What is the trend in metallic character across Period 3?\n"
     "- Why are noble gases unreactive?\n"
     "- How do the oxides of Period 3 elements change from basic to acidic?",
     "Across Period 3 (Na → Ar), elements change from highly reactive metals to non-metals to an inert gas. "
     "Metallic character decreases as nuclear charge increases and elements are more likely to gain electrons "
     "rather than lose them. Noble gases have a full outer electron shell, giving them maximum stability and "
     "no tendency to react. Period 3 oxides change from basic (Na₂O, MgO) through amphoteric (Al₂O₃) "
     "to acidic (SO₃, Cl₂O₇), reflecting the change from metallic to non-metallic character across the period."),
    ("5. Phenomenon Explanation",
     "Answer the driving question using all evidence gathered in this unit.\n\n"
     "Driving Question: Why do elements show patterns in their properties across groups and periods, "
     "and how can we use these patterns to predict the behaviour of unfamiliar elements?\n\n"
     "Use concepts: electron configuration, atomic radius, ionisation energy, reactivity trends, group and period trends.",
     "Elements show patterns in their properties because their chemical and physical behaviour is determined "
     "by their electron configuration — specifically the number and arrangement of electrons in the outermost shell. "
     "Elements in the same group have the same number of outer electrons, giving them similar chemical properties. "
     "Moving down a group, each element has more electron shells: outer electrons are further from the nucleus, "
     "easier to remove (lower ionisation energy), and in the case of metals this means greater reactivity. "
     "For non-metals like halogens, larger atoms attract electrons less effectively, so reactivity decreases. "
     "Across a period, nuclear charge increases while the number of shells is constant: atoms shrink, "
     "ionisation energy rises, and character shifts from metallic to non-metallic. These regular, predictable "
     "trends mean that once we know the position of an element in the periodic table, we can confidently "
     "predict its properties — a powerful tool for scientists in medicine, environmental chemistry, and materials science."),
]

_CHEM_META = {
    "subject": "Chemistry",
    "grade": "Grade 10",
    "substrand": "Sub-Strand 1.5: The Periodicity",
    "lessons": 7,
    "driving_question": "Why do elements show patterns in their properties across groups and periods, and how can we use these patterns to predict the behaviour of unfamiliar elements?",
    "substrand_overview_rows": [
        ("Phenomenon", "Sodium reacting vigorously with water; magnesium burning with a bright white flame; potassium exploding in water — why do these elements behave so differently, yet follow a predictable pattern?"),
        ("Driving Question", "Why do elements show patterns in their properties across groups and periods, and how can we use these patterns to predict the behaviour of unfamiliar elements?"),
        ("Sub-Strand SLOs — Knowledge", _CHEM_SLO_K),
        ("Sub-Strand SLOs — Skills", _CHEM_SLO_S),
        ("Sub-Strand SLOs — Attitudes/Values", _CHEM_SLO_A),
        ("Number of Lessons", "7 lessons (14 periods of 40 minutes each)"),
        ("Core Competencies", "Critical Thinking and Problem Solving · Communication and Collaboration · Digital Literacy · Self-efficacy · Learning to Learn"),
        ("Core Values", "Curiosity · Responsibility · Integrity · Care and Compassion · Excellence"),
        ("Assessment", "Formative: observation, equation writing, prediction tasks, summary table\nSummative: Final Explanation document"),
        ("Resources", "Periodic table; videos of Group I reactions; PhET balancing chemical equations; Khan Academy — periodic trends; magnesium ribbon; dilute HCl (teacher demo)"),
    ],
    "summary_table_headers": [
        "Lesson", "What I Did", "Key Understanding Built",
        "Connection to Driving Question", "Model / DQB Evolution"
    ],
    "summary_table_rows": _CHEM_SUMMARY_ROWS,
    "summary_instructions": (
        "Use this table to track your learning across all 7 lessons. "
        "For each lesson record: what you did, what you understood, how it connects to the driving question, "
        "and how your model and DQB evolved. Your completed Summary Table is evidence of your chemical thinking."
    ),
    "final_explanation_sections": _CHEM_FE_SECTIONS,
    "final_explanation_rubric": {
        "headers": ["Criterion", "Beginning (1)", "Developing (2)", "Proficient (3)", "Advanced (4)"],
        "rows": [
            ("1. Physical Property Trends",
             "Cannot describe any trend",
             "States one trend but cannot explain it using atomic structure",
             "Correctly describes atomic radius and ionisation energy trends for groups and periods with atomic structure explanation",
             "Includes quantitative data; draws labelled diagrams; connects trends to melting point and bonding type"),
            ("2. Group I and II Reactivity",
             "Cannot explain why Group I metals react with water",
             "Describes reactions but cannot explain the reactivity trend",
             "Explains Group I and II reactivity using ionisation energy; writes and balances correct equations",
             "Compares Group I and II in depth; explains trend with electron configuration evidence; predicts behaviour of Rb/Cs"),
            ("3. Halogen Properties",
             "Cannot describe halogen states or reactivity",
             "Lists halogen states and colours but cannot explain the reactivity trend",
             "Explains decreasing reactivity down Group VII using atomic size; writes displacement equations correctly",
             "Analyses halogen displacement reactions as evidence; connects to real-world applications with quantitative reasoning"),
            ("4. Period 3 and Noble Gases",
             "Cannot describe Period 3 trends or noble gas stability",
             "Identifies metals and non-metals in Period 3 but cannot explain the trend",
             "Describes metallic/non-metallic transition across Period 3; explains noble gas stability using electron configuration",
             "Explains oxide character changes across Period 3; links noble gas stability to electron configuration with depth"),
            ("5. Phenomenon Explanation",
             "Cannot connect periodic trends to the driving question",
             "Makes partial connection but misses key factors",
             "Explains the phenomenon using electron configuration, group and period trends correctly",
             "Complete evidence-based explanation covering all 5 SLOs; applies trends to predict unfamiliar element properties"),
        ],
    },
    "lessons_data": _CHEM_LESSONS,
}

# ═══════════════════════════════════════════════════════════════════════════════
# META — Mathematics: Similarity and Enlargement (12 lessons)
# ═══════════════════════════════════════════════════════════════════════════════

_MATH_SLO_K = (
    "- Determine the centre of enlargement and the linear scale factor for similar figures\n"
    "- Construct the image of an object under an enlargement given the centre and linear scale factor\n"
    "- Determine the area and volume scale factors of similar figures and solids\n"
    "- Relate linear scale factor, area scale factor, and volume scale factor in enlargements\n"
    "- Apply similarity and enlargement to real-life situations"
)
_MATH_SLO_S = (
    "- Identify corresponding sides and angles of similar figures\n"
    "- Calculate linear, area, and volume scale factors from given measurements\n"
    "- Construct enlarged or reduced images using compass and ruler\n"
    "- Solve problems involving scale drawings, maps, and real-life applications"
)
_MATH_SLO_A = (
    "- Appreciate the use of similarity and enlargement in art, architecture, photography, and mapping\n"
    "- Show precision and care when constructing geometric figures\n"
    "- Value the power of mathematical ratios to describe the real world"
)

def _math_lesson(num, title, inquiry, know, skills, attit, purpose, materials, safety, p1):
    return {
        "number": num, "title": title,
        "inquiry_question": inquiry,
        "slo_knowledge": know, "slo_skills": skills, "slo_attitudes": attit,
        "overview_purpose": purpose,
        "materials": materials, "safety": safety,
        "period1_heading": "Period 1 (40 minutes)",
        "period1_table": p1,
        "period2_heading": "Period 2 (40 minutes)",
        "period2_table": [],
        "reflections": [],
    }

_MATH_LESSONS = [
    _math_lesson(
        1, "Introduction to the Phenomenon",
        "How can a full-size photograph of a person be resized to fit a small passport photo box without distortion?",
        "- Recall that shapes can be resized while keeping the same proportions\n- Identify the driving question for the unit",
        "- Observe and describe how a photograph is resized; draw an initial model of how dimensions change",
        "- Show curiosity about how mathematics is used in photography, maps, and design",
        "Launches the unit phenomenon. Students watch a video of fitting a passport photo into a document, generate questions, and record initial thinking about how resizing works.",
        "Video clip of resizing a passport photograph (teacher-prepared or downloaded); printed photographs in different sizes; Driving Question Board; blank paper; ruler",
        "No specific hazards.",
        [
            _pr("Watch video of resizing a passport photograph; record: what changed and what stayed the same",
                "Video of passport photo fitting; two printed photographs — one large, one small — of the same image",
                "Pose phenomenon without introduction; ask 'What do you notice? What is the same and what is different?'",
                "Think–Pair–Share: individual → partner → whole class; focus on dimensions and shape",
                "Exit ticket: 'What stayed the same when the photograph was made smaller?'"),
            _pr("Measure dimensions of large and small photographs; calculate the ratio of corresponding sides",
                "Printed photographs; ruler; calculator",
                "Students measure length and width of each photograph; calculate length ratio and width ratio",
                "Discussion: 'Are the two ratios the same? What does this tell us?'",
                "Oral check: if a 20 cm × 15 cm photo is reduced so the length becomes 4 cm, what is the new width?"),
            _pr("Draw initial model: how do you think dimensions change when a shape is enlarged or reduced?",
                "Blank squared paper; ruler; pencil",
                "Prompt: 'Draw two rectangles that you think are the same shape but different sizes. What did you have to do to keep the shape the same?'",
                "Pair-share models; identify key ideas: multiply by a number, keep angle the same",
                "Collect initial models for comparison at unit end"),
            _pr("Create the Driving Question Board: what do we need to know to answer the passport photo question?",
                "DQB poster; sticky notes; markers",
                "Facilitate student question generation; group by theme (scale factor, area change, construction)",
                "Read all questions; prioritise which will be answered in this unit",
                "Photograph DQB for record"),
            _pr("Discuss: where else in real life are shapes enlarged or reduced? (maps, blueprints, models, projectors)",
                "Photographs of maps, architectural plans, scale models, projected images",
                "Brainstorm applications; connect to students' own experience (school plans, football field maps)",
                "KWL chart: what do we know, want to know, and will learn about similarity and enlargement?",
                "Individual reflection: 'Which application of similarity do you find most interesting? Why?'"),
        ],
    ),
    _math_lesson(
        2, "Properties of Similar Figures",
        "What makes two figures 'similar', and how can we recognise similar figures from their measurements?",
        "- Define similar figures as shapes with equal corresponding angles and proportional corresponding sides\n- State the two conditions required for similarity",
        "- Identify corresponding angles and sides in pairs of triangles and rectangles\n- Test whether two figures are similar by checking angle equality and side ratios\n- Distinguish between congruent and similar figures",
        "- Appreciate that similarity is a precise mathematical relationship, not just a vague impression of 'looking the same'",
        "Students investigate the two conditions for similarity through hands-on comparison of shape pairs.",
        "Pairs of similar and non-similar triangles (printed); protractor; ruler; calculator; geometry sets",
        "Take care with compass points.",
        [
            _pr("Examine 4 pairs of triangles; measure all angles and sides; record in a table",
                "Printed triangle pairs (similar and non-similar); protractor; ruler; measurement table",
                "Students measure independently; teacher circulates to check accuracy",
                "Discussion: 'Which pairs have equal angles? Which have proportional sides? Are these always both true?'",
                "Exit ticket: state the two conditions for similarity"),
            _pr("Test: are equal angles sufficient for similarity? Test pairs of rectangles with equal angles (90°) but different ratios",
                "Pairs of rectangles with same angles but different side ratios; ruler; calculator",
                "Students measure and calculate side ratios; determine whether rectangles are similar",
                "Discussion: for triangles, equal angles guarantee similarity; for other shapes, both conditions must be checked",
                "Oral probe: 'Can two squares of different sizes be similar? Explain.'"),
            _pr("Sort 8 shape pairs into similar / not similar; justify each decision",
                "Sorting cards with 8 shape pairs; answer record sheet",
                "Students work in pairs; must give a reason for each decision using angle or ratio evidence",
                "Class debrief: share decisions; discuss disagreements",
                "Formative check: peer-mark sorting decisions using criteria from the lesson"),
            _pr("Connect to phenomenon: 'Are the two passport photographs similar figures? How do you know?'",
                "Two printed photographs of different sizes of the same image",
                "Students measure corresponding dimensions and calculate ratios; confirm similarity",
                "Discussion: 'If the photographs are similar, what mathematical relationship links their dimensions?'",
                "Written summary: 'Two figures are similar if and only if…' — students complete in own words"),
            _pr("Update DQB: mark which questions about similar figures have been answered",
                "DQB poster; sticky notes",
                "Facilitate: 'Which DQB questions can we now partially or fully answer?'",
                "Students update their models to include similarity conditions",
                "Review DQB with class"),
        ],
    ),
    _math_lesson(
        3, "Centre of Enlargement",
        "What is the centre of enlargement and how is it used to construct an enlarged image?",
        "- Define centre of enlargement as the fixed point from which an enlargement is measured\n- Construct the image of a shape under enlargement given the centre and scale factor\n- Locate the centre of enlargement given an object and its image",
        "- Use rays from the centre of enlargement to locate image vertices\n- Construct enlarged and reduced images using ruler and rays\n- Verify constructions by measuring corresponding sides",
        "- Show precision when drawing construction lines\n- Appreciate that the centre of enlargement is the key that links object and image",
        "Students learn to construct enlargements and locate the centre of enlargement through guided construction activities.",
        "Squared paper or dot paper; ruler; pencil; compass; geometry sets; printed construction examples",
        "Take care with compass points.",
        [
            _pr("Examine a given enlargement: draw rays through corresponding vertices; identify where they meet",
                "Printed example of a triangle and its enlargement on squared paper; ruler",
                "Demonstrate: draw a ray from each pair of corresponding vertices; all rays meet at the centre of enlargement",
                "Students repeat with a second example; identify the centre of enlargement independently",
                "Oral check: 'What is the centre of enlargement and where is it in relation to the object and image?'"),
            _pr("Construct an enlarged image: given triangle ABC and centre O, scale factor 2 — construct A'B'C'",
                "Squared paper; ruler; pencil; triangle ABC and centre O marked",
                "Guided steps: draw ray OA, measure OA, mark A' at distance 2×OA; repeat for B and C; join A'B'C'",
                "Students check by measuring: is each side of A'B'C' exactly double the corresponding side of ABC?",
                "Formative task: construct a scale factor 3 enlargement of a given quadrilateral from a given centre"),
            _pr("Construct a reduction (fractional scale factor): scale factor ½ applied to a large triangle",
                "Squared paper; ruler; large triangle with centre marked",
                "Discuss: 'What does scale factor ½ mean? Where will the image be relative to the object?'",
                "Students construct the reduction; measure to verify",
                "Discussion: image is between the centre and the object when 0 < k < 1"),
            _pr("Locate the centre of enlargement given an object and image (reverse problem)",
                "Printed object-image pairs; ruler",
                "Students draw rays through pairs of corresponding vertices; find intersection",
                "Peer-check: compare located centres with partners",
                "Exit ticket: sketch an object and image; mark the centre of enlargement and explain how you found it"),
            _pr("Apply to phenomenon: the passport photo centre of enlargement is the camera/scanner focal point — discuss",
                "Diagram of camera and film/sensor; photograph scaling diagram",
                "Connect mathematical construction to real photography: the lens acts as the centre of enlargement",
                "Discussion: 'When you zoom in on a photo on your phone, what is the centre of enlargement?'",
                "Written reflection: describe in your own words what the centre of enlargement represents"),
        ],
    ),
    _math_lesson(
        4, "Linear Scale Factor",
        "How do we calculate the linear scale factor and use it to find missing side lengths of similar figures?",
        "- Define linear scale factor (LSF) as the ratio of corresponding side lengths\n- Calculate LSF given two corresponding sides\n- Use LSF to find missing side lengths in similar figures",
        "- Calculate LSF from given measurements\n- Find missing side lengths using the ratio LSF = image length ÷ object length\n- Solve problems involving maps, models, and scale drawings using LSF",
        "- Show accuracy in ratio calculations\n- Appreciate that LSF is the fundamental multiplier linking all corresponding measurements",
        "Students practise calculating and applying the linear scale factor to find missing measurements in similar figures.",
        "Squared paper; ruler; calculator; LSF problem worksheets; map or scale drawing example",
        "No specific hazards.",
        [
            _pr("Given pairs of similar triangles, calculate the ratio of each pair of corresponding sides; confirm all ratios are equal",
                "Printed similar triangle pairs with labelled side lengths; calculator",
                "Students calculate 3 ratios per pair; confirm they are equal (the LSF)",
                "Discussion: 'What does LSF = 2.5 mean in words? What does LSF = 0.4 mean?'",
                "Quick oral check: 'If LSF = 3 and a side of the object is 4 cm, what is the corresponding image side?'"),
            _pr("Use LSF to find missing side lengths in similar figures",
                "LSF problem worksheet with 4 problems",
                "Model one example: identify corresponding sides, calculate LSF, use it to find unknowns",
                "Students complete remaining problems; peer-check with worked solutions",
                "Formative quiz: 3 problems finding missing lengths in similar figures"),
            _pr("Apply LSF to a map problem: a map has scale 1:50,000; find real distances from map measurements",
                "Printed map excerpt or diagram; ruler; calculator",
                "Students measure map distances with ruler; multiply by scale factor to get real distances",
                "Discussion: 'A map scale of 1:50,000 means what? Is 1:50,000 an enlargement or a reduction?'",
                "Written task: measure 3 distances on the map and convert to real-world distances"),
            _pr("Apply LSF to a model problem: a model car has LSF 1/20 compared to the real car",
                "Model car dimensions diagram; real car dimensions; calculator",
                "Students calculate real car dimensions from model measurements",
                "Extension: 'If the real car is 4.8 m long, what is the model's length in cm?'",
                "Peer-check calculations"),
            _pr("Update DQB and model: how does LSF link to the passport photo phenomenon?",
                "DQB poster; model diagrams",
                "Students calculate the LSF between their two passport photographs from Lesson 1",
                "Discussion: 'What LSF would reduce a 15 cm × 12 cm photo to 3.5 cm × 2.8 cm?'",
                "DQB update: mark which questions about scale factor are now answered"),
        ],
    ),
    _math_lesson(
        5, "Area Scale Factor",
        "If the linear scale factor is k, what is the relationship between the areas of similar figures?",
        "- State that area scale factor (ASF) = k² where k is the linear scale factor\n- Calculate the area of an image given the area of the object and the LSF\n- Explain why ASF = k²",
        "- Calculate ASF from given LSF or from measured areas\n- Find unknown areas using the formula ASF = k²\n- Verify the relationship by direct measurement on constructed figures",
        "- Appreciate that area grows as the square of linear dimensions — a powerful and non-obvious mathematical result\n- Show systematic verification skills",
        "Students discover that area scales as the square of the linear scale factor through direct measurement and calculation.",
        "Squared paper; ruler; calculator; area calculation worksheet; video clip on area scale factor",
        "No specific hazards.",
        [
            _pr("Draw a 2 × 3 rectangle; enlarge by LSF = 2; draw the 4 × 6 image; count squares in both",
                "Squared paper; pencil; ruler",
                "Students draw both rectangles; count squares (area of object = 6 sq units; image = 24 sq units)",
                "Discussion: 'The lengths doubled — but the area did not double. What happened to the area?'",
                "Oral check: 'If LSF = 3, how many times bigger is the area?'"),
            _pr("Tabulate: for LSF = 1, 2, 3, 4 — calculate expected area ratio; verify by construction",
                "Results table; squared paper; ruler; calculator",
                "Students complete table: LSF | LSF² | measured area ratio; confirm ASF = k²",
                "Discussion: explain WHY area scales as k² (both length and width are multiplied by k)",
                "Exit ticket: 'If LSF = 5, the area scale factor is ___. If the object area is 8 cm², the image area is ___."),
            _pr("Solve ASF problems: given object area and LSF, find image area; given areas, find LSF",
                "ASF problem worksheet",
                "Model one of each type; students complete remaining problems independently",
                "Peer-check with worked solutions; discuss common errors (confusing LSF and ASF)",
                "Formative quiz: 4 ASF problems of mixed types"),
            _pr("Watch short video on area scale factor applications; identify passport photo area change",
                "Video on area scale factor in photography/printing (Khan Academy or teacher-prepared)",
                "Students calculate: if passport photo LSF = 0.25 (relative to original), what fraction of the original area does it occupy?",
                "Discussion: 'Is the passport photo area ¼ or 1/16 of the original? Why?'",
                "Written explanation: 'The area scale factor equals k² because…'"),
            _pr("Update DQB and model: add ASF = k² to the model with a worked example",
                "DQB poster; student models; squared paper",
                "Students annotate their models with the ASF formula and a worked example",
                "Partner check: is the formula and example correct?",
                "Review DQB: are the area-related questions now answered?"),
        ],
    ),
    _math_lesson(
        6, "Volume Scale Factor",
        "If the linear scale factor is k, what is the relationship between the volumes of similar solids?",
        "- State that volume scale factor (VSF) = k³ where k is the linear scale factor\n- Calculate the volume of an image solid given the volume of the object and the LSF\n- Explain why VSF = k³",
        "- Calculate VSF from given LSF or measured volumes\n- Find unknown volumes using the formula VSF = k³\n- Verify the relationship using simple 3D shapes (cuboids)",
        "- Show curiosity about how a small change in linear scale produces a large change in volume\n- Appreciate real-world implications (scale models, packaging, medical imaging)",
        "Students discover that volume scales as the cube of the linear scale factor through systematic investigation with cuboids.",
        "Connecting cubes or cuboid models; ruler; calculator; VSF problem worksheet; volume scale factor video",
        "No specific hazards.",
        [
            _pr("Build a 1×1×1 cube; then a 2×2×2 cube; count unit cubes in each; calculate the ratio",
                "Connecting cubes or sugar cube models",
                "Students build both; count cubes (1 vs. 8); calculate ratio",
                "Discussion: 'The lengths doubled — the volume multiplied by 8. How does 2 relate to 8?'",
                "Oral check: 'If LSF = 3, how many times bigger is the volume?'"),
            _pr("Tabulate: for LSF = 1, 2, 3, 4 — calculate expected volume ratio; verify by construction",
                "Results table; connecting cubes",
                "Students complete table: LSF | LSF³ | measured volume ratio; confirm VSF = k³",
                "Discussion: explain WHY volume scales as k³ (length × width × height all multiplied by k)",
                "Exit ticket: 'If LSF = 4, VSF = ___. If object volume is 5 cm³, image volume = ___.'"),
            _pr("Solve VSF problems: given object volume and LSF, find image volume; given volumes, find LSF",
                "VSF problem worksheet",
                "Model one of each type; students complete remaining problems independently",
                "Peer-check with worked solutions",
                "Formative quiz: 4 VSF problems of mixed types"),
            _pr("Apply VSF to a model problem: a model ship has LSF 1/50 — what fraction of the real ship's volume does it have?",
                "Model ship dimensions; real ship volume calculation",
                "Students calculate VSF = (1/50)³ = 1/125,000; discuss the huge difference",
                "Extension: if the real ship holds 50,000 m³ of cargo, what volume does the model represent?",
                "Written explanation of why small models are so much less than real objects in volume"),
            _pr("Update model: add VSF = k³ alongside ASF = k² and LSF = k; create a summary of all three",
                "Student models; summary table template",
                "Students create a side-by-side comparison: LSF | ASF | VSF and their relationships",
                "Pair-check summaries for accuracy",
                "Review with class; prepare for next lesson on the full relationship"),
        ],
    ),
    _math_lesson(
        7, "Relationships: LSF, ASF, and VSF",
        "How are the linear, area, and volume scale factors mathematically related, and how do we move between them?",
        "- State the relationships: ASF = LSF², VSF = LSF³\n- Calculate any one scale factor given another\n- Solve problems that require moving between linear, area, and volume scale factors",
        "- Calculate LSF from ASF (LSF = √ASF) and from VSF (LSF = ∛VSF)\n- Solve multi-step problems involving all three scale factors\n- Apply relationships to real-world contexts (maps, models, photographs)",
        "- Show algebraic reasoning skills in relating the three scale factors\n- Appreciate the elegance of the mathematical connection between dimensions",
        "Students consolidate and apply the three-way relationship between LSF, ASF, and VSF through problem-solving.",
        "Calculator; LSF/ASF/VSF relationship worksheet; summary table from Lesson 6",
        "No specific hazards.",
        [
            _pr("Review the three formulae: ASF = k², VSF = k³ — confirm with 2–3 numerical examples",
                "Student summary tables from Lesson 6; calculator",
                "Quick oral review: teacher calls out k, students call back ASF and VSF",
                "Students verify 3 examples independently; check with partner",
                "Oral check: if k = 1.5, what are ASF and VSF?"),
            _pr("Reverse calculation: given ASF = 16, find LSF; given VSF = 27, find LSF",
                "Relationship worksheet",
                "Guided steps: LSF = √16 = 4; LSF = ∛27 = 3; then verify by computing ASF and VSF back",
                "Students complete 4 reverse calculations independently",
                "Formative quiz: 3 problems — some forward, some reverse"),
            _pr("Mixed problems: given area of object and image, find LSF; then find VSF",
                "Mixed problem worksheet",
                "Model: 'Object area = 9 cm², image area = 36 cm² → ASF = 4 → LSF = 2 → VSF = 8'",
                "Students complete remaining problems; peer-check",
                "Discussion: 'In what order do you need to calculate when the problem gives you areas or volumes?'"),
            _pr("Apply all three to a real scenario: a room plan has LSF 1:100 — find the real room area and volume from plan measurements",
                "Floor plan with measurements; calculator",
                "Students measure plan room dimensions; calculate real dimensions, area, and volume",
                "Extension: 'The room needs to be carpeted. What area of carpet is needed in m²?'",
                "Written task: explain all steps of the calculation for someone who hasn't studied scale factors"),
            _pr("Update DQB and model: final summary of all three relationships on one diagram",
                "DQB poster; student models",
                "Students create a one-page summary diagram: k → k² (ASF) → k³ (VSF)",
                "Review DQB: mark all questions about scale factor relationships as answered",
                "Self-assessment: which relationship do I find most challenging? Why?"),
        ],
    ),
    _math_lesson(
        8, "Constructing Similar Figures",
        "How do we accurately construct the image of a complex shape under an enlargement?",
        "- Construct the image of a polygon under enlargement using a given centre and scale factor\n- Verify constructions by measuring corresponding sides and angles",
        "- Accurately construct enlarged and reduced polygons using ruler and rays\n- Verify similarity by measuring and calculating ratios\n- Construct scale drawings of real objects",
        "- Show precision and patience in geometric construction\n- Appreciate that accurate construction is a foundational skill in engineering and design",
        "Students practise constructing enlargements and reductions of polygons, applying all the skills from previous lessons.",
        "Squared paper; ruler; compass; geometry set; protractor; construction exercise sheets",
        "Take care with compass points.",
        [
            _pr("Construct the image of triangle PQR under enlargement, centre O, scale factor 3",
                "Squared paper; ruler; triangle PQR and centre O marked",
                "Students draw rays OP, OQ, OR; mark P', Q', R' at 3× the distances; join the image",
                "Measure sides of P'Q'R' and divide by corresponding sides of PQR; confirm = 3",
                "Oral check: 'Is the image similar to the object? How do you know?'"),
            _pr("Construct the image of a quadrilateral ABCD under reduction, scale factor ⅓, centre O outside the shape",
                "Squared paper; ruler; quadrilateral and centre marked",
                "Guided steps; students note that when centre is outside the shape, rays must be extended",
                "Verify by measuring; discuss: 'Where does the image appear relative to the object?'",
                "Exit ticket: sketch the object, image, and centre; label rays and measurements"),
            _pr("Construct a scale drawing of the classroom on squared paper using LSF 1:50",
                "Measuring tape; squared paper (1 cm = 50 cm); ruler",
                "Students measure classroom dimensions in groups; calculate scaled dimensions; draw the plan",
                "Compare plans from different groups; discuss measurement accuracy",
                "Written task: calculate the real floor area and the plan area; verify ASF = k²"),
            _pr("Solve a construction problem: given the image and the centre, reconstruct the original object (scale factor ½ — so reverse by applying factor 2)",
                "Printed image with centre marked; ruler",
                "Students draw reverse rays from centre through image vertices; mark object vertices at half distance",
                "Peer-check constructions; discuss: 'How do you know which direction to draw the rays?'",
                "Formative assessment: independently construct the original object for a new example"),
            _pr("Review construction accuracy: measure all corresponding sides and angles; calculate actual vs. theoretical ratios",
                "Completed construction sheets; ruler; protractor; calculator",
                "Students calculate % error for their constructions: (actual ratio − theoretical ratio) / theoretical ratio × 100",
                "Discussion: sources of error; how to minimise them",
                "Self-assessment: identify one construction skill to improve"),
        ],
    ),
    _math_lesson(
        9, "Applications: Maps and Scale Drawings",
        "How is similarity and enlargement used in maps, architectural plans, and other real-world scale drawings?",
        "- Interpret map scales as linear scale factors\n- Calculate real distances from map distances and vice versa\n- Apply LSF to reading and creating scale drawings",
        "- Measure map distances and convert to real distances using LSF\n- Calculate required map measurements from real dimensions\n- Create a simple scale drawing of a real object",
        "- Appreciate the importance of maps and scale drawings in navigation, construction, and planning\n- Show accuracy in measurement and conversion",
        "Students apply their scale factor knowledge to interpret maps and create their own scale drawings.",
        "Printed map extract with stated scale; ruler; calculator; graph paper for scale drawing; measuring tape",
        "No specific hazards.",
        [
            _pr("Read a map scale: convert 1:25,000 to a sentence; calculate real distances for 5 map measurements",
                "Printed map with scale 1:25,000; ruler; calculator",
                "Students measure map segments with ruler; multiply by 25,000; convert to km",
                "Discussion: 'Is this an enlargement or a reduction? What is the LSF as a decimal?'",
                "Exit ticket: 'A 3 cm measurement on this map represents ___ km in reality'"),
            _pr("Calculate map distances from real distances: if two towns are 12 km apart, how far apart are they on a 1:50,000 map?",
                "Calculator; map problem worksheet",
                "Students solve 4 problems converting real distances to map distances",
                "Peer-check with worked solutions",
                "Discussion: when would you need to calculate map distance from real distance? (making a map)"),
            _pr("Examine an architectural floor plan; interpret scale; calculate real room dimensions and floor areas",
                "Printed floor plan with stated scale (e.g., 1:100); ruler; calculator",
                "Students measure each room on the plan; calculate real dimensions; calculate floor area",
                "Discuss: 'Which room is largest? How much carpet would the whole flat need?'",
                "Written task: present room dimensions and areas in a table"),
            _pr("Create your own scale drawing: measure a school classroom and draw a floor plan at 1:50",
                "Measuring tape; graph paper; ruler; pencil",
                "Groups measure the classroom; calculate scaled dimensions; draw the plan",
                "Compare plans from different groups; identify and discuss any discrepancies",
                "Reflection: 'What was the hardest part of making an accurate scale drawing?'"),
            _pr("Update DQB: how does map reading use similarity and enlargement?",
                "DQB poster; student models",
                "Students connect map scale to LSF and ASF; add map application to their models",
                "Review DQB: mark map-related questions as answered",
                "Peer discussion: 'Can you think of a situation where you would need to use a map scale in everyday life?'"),
        ],
    ),
    _math_lesson(
        10, "Applications: Area and Volume in Real Life",
        "How do area and volume scale factors apply to real problems such as painting surfaces, filling containers, and comparing models with originals?",
        "- Apply ASF to calculate the amount of material needed to cover similar surfaces\n- Apply VSF to calculate capacities of similar containers\n- Interpret the meaning of ASF and VSF in real-world contexts",
        "- Calculate paint or fabric required using ASF\n- Calculate capacity ratios using VSF\n- Solve multi-step real-life problems using LSF, ASF, and VSF together",
        "- Appreciate that understanding area and volume scale has direct practical implications for cost, materials, and efficiency\n- Show perseverance when solving multi-step problems",
        "Students apply area and volume scale factors to real-world problems involving materials, containers, and costs.",
        "Calculator; real-life application worksheet; container diagrams; paint coverage data",
        "No specific hazards.",
        [
            _pr("Problem: a tin of paint covers 10 m². If the wall dimensions are scaled by factor 3, how many tins are needed for the larger wall?",
                "Paint coverage problem card; calculator",
                "Students identify: larger wall area = 3² × original area = 9 × original area → 9 tins needed",
                "Discussion: 'Why do we square the scale factor? Draw a diagram to show why.'",
                "Exit ticket: solve a similar problem with LSF = 4 and different original area"),
            _pr("Problem: two similar water tanks have LSF = 2. If the small tank holds 500 litres, how much does the large tank hold?",
                "Tank diagram; calculator",
                "Students calculate VSF = 2³ = 8; large tank = 8 × 500 = 4,000 litres",
                "Discussion: 'Why does a small change in linear dimensions produce a large change in volume?'",
                "Peer problem: create and swap your own tank problem"),
            _pr("Problem: a statue and a model have LSF = 1:20. The model uses 0.5 kg of bronze. How much does the statue use?",
                "Statue/model problem card; calculator",
                "Students identify this as a VSF problem (mass ∝ volume): VSF = 20³ = 8,000; statue mass = 4,000 kg",
                "Discussion: 'Does this seem reasonable for a bronze statue? What real statues could be this heavy?'",
                "Written explanation: why do we use VSF for mass calculations?"),
            _pr("Costing problem: if painting a model costs KES 200 (using ASF), what does painting the full-size version cost at same rate?",
                "Costing problem worksheet; calculator",
                "Students identify ASF needed; calculate proportional cost",
                "Peer-check and discuss: 'Why does cost scale with area rather than volume for painting?'",
                "Formative quiz: 3 real-life ASF and VSF problems"),
            _pr("Connect to phenomenon: calculate the exact LSF, ASF for the passport photo scenario from Lesson 1 measurements",
                "Measurements from Lesson 1; calculator",
                "Students revisit their Lesson 1 photograph measurements; calculate LSF and ASF",
                "Discussion: 'Now you can fully explain the passport photo phenomenon using mathematics'",
                "DQB update: mark all area/volume application questions as answered"),
        ],
    ),
    _math_lesson(
        11, "Problem Solving with Similarity",
        "How do we identify and solve similarity problems when they are embedded in unfamiliar real-world contexts?",
        "- Identify similarity in unfamiliar geometric and real-world contexts\n- Select the appropriate scale factor (LSF, ASF, or VSF) for a given problem\n- Solve multi-step similarity problems",
        "- Analyse word problems to identify whether LSF, ASF, or VSF is needed\n- Apply geometric reasoning to identify similar triangles within larger figures\n- Present solutions with clear working and justification",
        "- Show persistence and systematic thinking when tackling multi-step problems\n- Appreciate that recognising mathematical structure is a transferable skill",
        "Students consolidate all similarity and enlargement knowledge by solving a range of multi-step and contextualised problems.",
        "Problem set worksheet; calculator; geometry set; squared paper",
        "No specific hazards.",
        [
            _pr("Identify similar triangles in a diagram: triangle within a triangle (AA condition); state the LSF",
                "Diagram with overlapping similar triangles; ruler",
                "Students identify corresponding angles (using parallel lines or given angle data); state AA similarity",
                "Calculate LSF from given side lengths",
                "Exit ticket: given a diagram, identify similar triangles and state why they are similar"),
            _pr("Solve a shadow problem: a 2 m pole casts a 3 m shadow; a nearby tree casts a 12 m shadow — find the tree height",
                "Shadow diagram; problem card",
                "Guide students to identify the two similar triangles formed by the pole/shadow and tree/shadow",
                "Students set up ratio: 2/3 = h/12 → h = 8 m",
                "Peer-check; discuss: 'What assumption must you make for this method to work?'"),
            _pr("Solve a photography problem: a camera lens produces an image that is 1/50 of the object height — find the image dimensions for a person 1.7 m tall",
                "Photography diagram; calculator",
                "Students calculate image height; discuss whether this is LSF or a different type of ratio",
                "Extension: find the area scale factor and the area of the image on the film/sensor",
                "Written explanation: which type of scale factor is relevant here and why?"),
            _pr("Mixed problem set: 4 problems requiring students to select the correct scale factor type",
                "Mixed problem worksheet covering maps, models, areas, and volumes",
                "Students work independently; teacher circulates and asks probing questions",
                "Peer-mark with worked solutions; discuss common errors",
                "Class debrief: which type of problem is most commonly confused? Why?"),
            _pr("Review Summary Table: ensure all 11 lessons are recorded accurately",
                "Individual Summary Tables",
                "Students complete any missing rows; peer-review for accuracy",
                "Teacher reviews sample tables; provide written feedback",
                "Reflection: 'What is the most powerful idea I have learned in this unit?'"),
        ],
    ),
    _math_lesson(
        12, "Unit Consolidation and Final Explanation Preparation",
        "How can I use all the mathematical ideas from this unit to fully explain the passport photo phenomenon?",
        "- Summarise all key concepts: similar figures, centre of enlargement, LSF, ASF, VSF, and applications\n- Construct a complete mathematical explanation of the driving question\n- Self-assess readiness for the Final Explanation task",
        "- Solve review problems covering all unit concepts\n- Complete the Summary Table\n- Draft the Final Explanation with evidence from all lessons",
        "- Show confidence in applying all unit concepts to unfamiliar situations\n- Appreciate the power of mathematical reasoning to explain real phenomena",
        "Students complete all summary work, solve a review problem set, and prepare their Final Explanation document.",
        "Summary Table (student copies); DQB poster; Final Explanation task description; rubric; calculator; geometry set",
        "No specific hazards.",
        [
            _pr("Review quiz: 5 problems covering similar figures, LSF, ASF, VSF, and construction",
                "Review problem set; calculator; geometry set",
                "Students complete quiz independently; teacher circulates",
                "Peer-mark with worked solutions; identify any remaining gaps",
                "Teacher notes common errors for class debrief"),
            _pr("Answer the driving question as a class: 'How can a full-size photo be resized to fit a passport box without distortion?'",
                "DQB poster; all evidence from 12 lessons",
                "Facilitate class construction of a complete mathematical explanation",
                "Students write the class explanation in their own words in notebooks",
                "Discuss: 'What other real-world applications use the same mathematical ideas?'"),
            _pr("Complete the Summary Table: verify all 12 lessons are recorded; peer-review",
                "Individual Summary Tables",
                "Students complete missing rows; swap with partner for review",
                "Peer feedback: is each entry accurate and does it connect to the driving question?",
                "Teacher spot-checks and provides written feedback on 5–6 tables"),
            _pr("Introduce the Final Explanation document and rubric; self-assess on each criterion",
                "Final Explanation task description; rubric",
                "Students read the rubric; rate themselves 1–4 on each criterion",
                "Partner discussion: 'Which criterion is your strongest? Which needs most work?'",
                "Students write a 5-minute plan for their Final Explanation"),
            _pr("Reflection: how has your mathematical model of the passport photo phenomenon evolved over 12 lessons?",
                "Initial model from Lesson 1; current model; DQB",
                "Students compare initial and final models side by side; annotate what changed and why",
                "Class share: 3–4 students describe their model evolution",
                "Exit ticket: 'In three sentences, explain how the passport photo fits into a small box without distortion'"),
        ],
    ),
]

_MATH_SUMMARY_ROWS = [
    (1, "Observed passport photo resizing; measured photo dimensions; drew initial model; created DQB",
        "A photograph can be resized while keeping the same proportions — dimensions change by a constant ratio",
        "Introduces the phenomenon: how does resizing work without distortion?",
        "DQB created; initial model drawn"),
    (2, "Compared shape pairs; tested two conditions for similarity (equal angles, proportional sides)",
        "Similar figures have equal corresponding angles AND proportional corresponding sides",
        "The passport photos are similar figures — both conditions confirmed",
        "Model updated: similarity conditions added"),
    (3, "Drew rays to locate centre of enlargement; constructed enlarged and reduced images",
        "Centre of enlargement is the fixed point from which all enlargement rays originate",
        "The camera lens acts as the centre of enlargement when taking a photograph",
        "Model updated: centre of enlargement and rays shown"),
    (4, "Calculated LSF from side ratios; solved missing length problems; applied to maps and models",
        "LSF = image length ÷ object length; links all corresponding lengths of similar figures",
        "Passport photo LSF calculated from actual measurements",
        "Model updated: LSF formula and value added"),
    (5, "Discovered ASF = k² by counting squares; applied to photography and area problems",
        "Area scale factor equals the square of the linear scale factor",
        "Passport photo area is k² × original area — explains why the photo is much smaller",
        "Model updated: ASF = k² shown with worked example"),
    (6, "Built cuboid models; discovered VSF = k³; applied to model and container problems",
        "Volume scale factor equals the cube of the linear scale factor",
        "A small linear reduction produces a very large volume reduction (k³ effect)",
        "Model updated: VSF = k³ added; all three scale factors now on one diagram"),
    (7, "Practised converting between LSF, ASF, and VSF; solved multi-step problems",
        "ASF = LSF²; VSF = LSF³; can calculate any one from another",
        "The three scale factors are mathematically linked — understanding one gives the others",
        "Summary table created: LSF | ASF | VSF relationships"),
    (8, "Constructed enlargements and reductions; created a classroom scale drawing",
        "Accurate constructions require careful measurement of rays from the centre",
        "Scale drawings are enlarged/reduced versions — the same relationship as the passport photo",
        "Classroom floor plan created at 1:50"),
    (9, "Read map scales; converted map distances to real distances; interpreted a floor plan",
        "Map scale is a linear scale factor; area on the map relates to real area by ASF = k²",
        "Maps are mathematical reductions of reality — the same principle as passport photos",
        "Model updated: map application added"),
    (10, "Solved paint, tank, and statue problems using ASF and VSF",
        "Material needed scales with area (ASF); capacity scales with volume (VSF)",
        "Practical applications confirm the mathematical relationships discovered in lessons 5–6",
        "DQB: area/volume application questions answered"),
    (11, "Identified similar triangles; solved shadow, photography, and mixed context problems",
        "Similar triangles arise naturally in real-world contexts; selecting the right scale factor is key",
        "Problem-solving with similarity applies to photography — directly relevant to the phenomenon",
        "Summary Table completed; problem-solving skills consolidated"),
    (12, "Completed review quiz; answered driving question; reviewed Summary Table; planned Final Explanation",
        "All unit concepts connect: similarity → scale factors → applications → phenomenon explanation",
        "Complete understanding: passport photo resizing is a mathematical enlargement using a specific LSF",
        "Final model complete; DQB fully answered; ready for Final Explanation"),
]

_MATH_FE_SECTIONS = [
    ("1. Similar Figures",
     "Define similar figures and state the two conditions for similarity. Give an example with measurements.\n\n"
     "Guiding Questions:\n- What are the two conditions for two figures to be similar?\n"
     "- How do you test whether two triangles are similar?\n"
     "- Are all squares similar? Are all rectangles similar? Explain.",
     "Two figures are similar if and only if their corresponding angles are equal AND their corresponding "
     "sides are in the same ratio. For triangles, equal angles alone (AA condition) is sufficient because "
     "the side ratios are then automatically equal. All squares are similar to each other (all angles 90° "
     "and all sides in ratio k). Rectangles are not all similar — they also need the ratio of length to "
     "width to be equal. Example: a 6 × 4 rectangle and a 9 × 6 rectangle are similar (ratio = 1.5); "
     "a 6 × 4 and a 10 × 4 rectangle are not similar (ratios 10/6 ≠ 4/4)."),
    ("2. Centre of Enlargement and Linear Scale Factor",
     "Explain the centre of enlargement and linear scale factor. Show how to construct an enlargement.\n\n"
     "Guiding Questions:\n- What is the centre of enlargement and how do you locate it?\n"
     "- How do you construct the image of a shape given the centre and LSF?\n"
     "- What does LSF = 0.5 mean? What does LSF = 3 mean?",
     "The centre of enlargement is the fixed point from which all enlargement rays are drawn. "
     "To find it, draw rays through corresponding pairs of vertices — they all meet at the centre. "
     "To construct an enlargement with centre O and LSF k: draw a ray from O through each vertex, "
     "measure the distance OA, mark A' at distance k × OA along the ray. LSF = 3 means the image "
     "is 3 times the linear dimensions of the object. LSF = 0.5 means the image is half the linear "
     "dimensions of the object (a reduction). The image appears between the centre and the object when 0 < k < 1."),
    ("3. Area and Volume Scale Factors",
     "State and explain the formulae ASF = k² and VSF = k³. Show why each is true with a diagram.\n\n"
     "Guiding Questions:\n- Why does area scale as k² and not k?\n"
     "- Why does volume scale as k³?\n"
     "- If LSF = 3, an object has area 4 cm², and volume 2 cm³, find the image area and volume.",
     "When linear dimensions are multiplied by k, area is multiplied by k² because area = length × width "
     "and both dimensions are multiplied by k: (kl) × (kw) = k²(lw). Volume is multiplied by k³ because "
     "volume = length × width × height and all three are multiplied by k: (kl)(kw)(kh) = k³(lwh). "
     "If LSF = 3: ASF = 9, so image area = 9 × 4 = 36 cm². VSF = 27, so image volume = 27 × 2 = 54 cm³."),
    ("4. Real-Life Applications",
     "Describe three real-life applications of similarity and enlargement. Include calculations.\n\n"
     "Guiding Questions:\n- How is similarity used in maps and scale drawings?\n"
     "- Give an example where VSF is needed to solve a real problem.\n"
     "- How does the passport photo phenomenon relate to similarity and enlargement?",
     "Maps: a map scale of 1:50,000 means the LSF is 1/50,000. A 4 cm map distance represents "
     "4 × 50,000 = 200,000 cm = 2 km. The ASF = (1/50,000)² = 1/2,500,000,000 — real areas are "
     "2.5 billion times larger than map areas. "
     "Similar containers: two similar tanks with LSF = 2; VSF = 8; if small tank holds 500 L, large holds 4,000 L. "
     "Passport photo: the photograph is a mathematical reduction of the original image. The camera applies "
     "a linear scale factor to project the 3D scene onto a 2D sensor; printing further reduces by another LSF. "
     "The image remains undistorted because all dimensions are multiplied by the same constant k."),
    ("5. Phenomenon Explanation",
     "Answer the driving question fully, using all unit evidence.\n\n"
     "Driving Question: How can a full-size photograph of a person be resized to fit a small passport photo box without distortion?\n\n"
     "Use all key concepts: similar figures, centre of enlargement, LSF, ASF, VSF.",
     "A passport photograph is a mathematically similar figure to the original image of the person. "
     "When a photograph is taken, the camera lens acts as the centre of enlargement, projecting the "
     "real scene onto the camera sensor at a linear scale factor determined by the focal length and "
     "distance. The resulting image has the same shape as the original (corresponding angles are equal "
     "and corresponding distances are in constant ratio) but is much smaller in linear, area, and volume terms. "
     "When the digital image is then resized to fit a 3.5 cm × 4.5 cm passport box, a second linear scale "
     "factor is applied. As long as the same LSF is applied to both dimensions, the shape is preserved — "
     "the person looks recognisable, just smaller. The area of the passport photo is LSF² times the area "
     "of the original digital image, making it a tiny fraction of the original. This demonstrates that "
     "similarity and enlargement are not just abstract geometry — they are the mathematical foundation of "
     "photography, mapping, architectural design, and any situation where shapes must be scaled without distortion."),
]

_MATH_META = {
    "subject": "Mathematics",
    "grade": "Grade 10",
    "substrand": "Sub-Strand 2.1: Similarity and Enlargement",
    "lessons": 12,
    "driving_question": "How can a full-size photograph of a person be resized to fit a small passport photo box without distortion?",
    "substrand_overview_rows": [
        ("Phenomenon", "A passport photograph — a full-size person reduced to a tiny 3.5 cm × 4.5 cm box without distortion. How is this mathematically possible?"),
        ("Driving Question", "How can a full-size photograph of a person be resized to fit a small passport photo box without distortion?"),
        ("Sub-Strand SLOs — Knowledge", _MATH_SLO_K),
        ("Sub-Strand SLOs — Skills", _MATH_SLO_S),
        ("Sub-Strand SLOs — Attitudes/Values", _MATH_SLO_A),
        ("Number of Lessons", "12 lessons (24 periods of 40 minutes each)"),
        ("Core Competencies", "Critical Thinking and Problem Solving · Communication and Collaboration · Digital Literacy · Self-efficacy · Creativity and Innovation"),
        ("Core Values", "Curiosity · Responsibility · Precision · Excellence · Integrity"),
        ("Assessment", "Formative: similarity tests, construction accuracy checks, scale factor calculations, problem sets, summary table\nSummative: Final Explanation document"),
        ("Resources", "Printed photographs; maps and floor plans; connecting cubes; geometry sets; squared paper; calculator; Khan Academy — Similarity and Scale Factor; GeoGebra constructions"),
    ],
    "summary_table_headers": [
        "Lesson", "What I Did", "Key Understanding Built",
        "Connection to Driving Question", "Model / DQB Evolution"
    ],
    "summary_table_rows": _MATH_SUMMARY_ROWS,
    "summary_instructions": (
        "Use this table to track your learning across all 12 lessons. "
        "For each lesson record: what you did, what you understood, how it connects to the driving question, "
        "and how your model and DQB evolved. Your completed Summary Table is evidence of your mathematical thinking."
    ),
    "final_explanation_sections": _MATH_FE_SECTIONS,
    "final_explanation_rubric": {
        "headers": ["Criterion", "Beginning (1)", "Developing (2)", "Proficient (3)", "Advanced (4)"],
        "rows": [
            ("1. Similar Figures",
             "Cannot state conditions for similarity",
             "States one condition but not both",
             "Correctly states and applies both conditions; gives example with measurements",
             "Distinguishes AA for triangles from two-condition requirement for other shapes; applies to novel examples"),
            ("2. Centre of Enlargement and LSF",
             "Cannot define centre of enlargement or LSF",
             "Defines one but cannot construct an enlargement",
             "Correctly constructs an enlargement; calculates LSF from measurements",
             "Constructs both enlargements and reductions; locates centre from object/image pair; explains negative scale factors"),
            ("3. ASF and VSF",
             "Cannot state the formulae",
             "States formulae but makes errors in calculation",
             "Correctly calculates ASF and VSF; applies to find image area and volume",
             "Explains WHY ASF = k² and VSF = k³ with diagrams; solves multi-step problems using all three"),
            ("4. Real-Life Applications",
             "Cannot connect theory to any real-life application",
             "Identifies one application but calculation is incomplete",
             "Correctly solves two application problems (e.g., map scale and container volume)",
             "Solves three or more applications; selects correct scale factor type; explains reasoning clearly"),
            ("5. Phenomenon Explanation",
             "Cannot explain how a photo is resized",
             "Mentions scale factor but explanation is incomplete or inaccurate",
             "Fully explains the photograph resizing using LSF, ASF, similarity, and centre of enlargement",
             "Complete, mathematically precise explanation with calculations; connects to additional real-world applications"),
        ],
    },
    "lessons_data": _MATH_LESSONS,
}

# ═══════════════════════════════════════════════════════════════════════════════
# SUBJECT REGISTRY
# ═══════════════════════════════════════════════════════════════════════════════

ALL_META = {
    "biology":   _BIO_META,
    "physics":   _PHY_META,
    "chemistry": _CHEM_META,
    "maths":     _MATH_META,
}

# ═══════════════════════════════════════════════════════════════════════════════
# BUILDERS
# ═══════════════════════════════════════════════════════════════════════════════

def build_lesson_sequence(subject_key: str) -> "Document":
    """Build the CBE Lesson Sequence document for a May18 subject."""
    meta = ALL_META[subject_key]
    lessons = meta["lessons_data"]

    doc = Document()
    _set_page_landscape(doc)

    _doc_title(doc, f"{meta['subject']} | {meta['grade']} | {meta['substrand']}")
    _doc_subtitle(doc,
                  f"CBE Lesson Sequence  ·  {meta['lessons']} Lessons  ·  "
                  f"{meta['lessons'] * 2} Periods")

    # Sub-strand overview using meta rows (source docs have no Section A to parse)
    _build_table0_overview(doc, meta)
    _tbl_no_spacing(doc)
    _tbl_no_spacing(doc)

    summary_rows = meta["summary_table_rows"]

    for lesson in lessons:
        lesson_num = lesson["number"]
        summary_row = (summary_rows[lesson_num - 1]
                       if lesson_num - 1 < len(summary_rows) else None)

        _add_page_break(doc)
        _build_table_A(doc, lesson)
        _tbl_no_spacing(doc)
        _build_table_B(doc, lesson)
        _tbl_no_spacing(doc)
        _build_table_C_period(doc, lesson, 1)
        _tbl_no_spacing(doc)
        if lesson.get("period2_table"):
            _build_table_C_period(doc, lesson, 2)
            _tbl_no_spacing(doc)
        _build_table_D(doc, lesson)
        _tbl_no_spacing(doc)
        _build_table_E(doc, summary_row)
        _tbl_no_spacing(doc)
        _tbl_no_spacing(doc)

    return doc


def build_final_explanation(subject_key: str) -> "Document":
    """Build the Final Explanation document for a May18 subject."""
    meta = ALL_META[subject_key]

    doc = Document()
    _set_page_landscape(doc)

    substrand_topic = meta["substrand"].split(":")[-1].strip().upper()
    fe_title = f"FINAL EXPLANATION: {substrand_topic}"

    _doc_title(doc, f"{meta['subject']} | {meta['grade']} | {meta['substrand']}")
    _doc_subtitle(doc, "Final Explanation Document")

    # Header block
    t0 = _new_table(doc, 5, 2)
    _col_widths(t0, [2.083, 7.417])

    c = _merge_row(t0, 0)
    _shade(c, C_NAVY)
    _cell_para(c, fe_title, bold=True, size_pt=11, color_hex=C_WHITE)

    c = _merge_row(t0, 1)
    _shade(c, C_TEAL)
    _cell_para(c, f"{meta['subject']} {meta['grade']} — Student Assessment Document",
               bold=True, size_pt=11, color_hex=C_WHITE)

    for ri, label in enumerate(["Student Name", "Class", "Date"], start=2):
        _shade(t0.rows[ri].cells[0], C_LT_BLUE)
        _shade(t0.rows[ri].cells[1], C_WHITE)
        _cell_para(t0.rows[ri].cells[0], label, bold=True, size_pt=9)
        _cell_para(t0.rows[ri].cells[1],
                   "_______________________________________________", size_pt=9)

    _tbl_no_spacing(doc)

    # Instructions block
    t1 = _new_table(doc, 2, 1)
    _col_widths(t1, [9.5])
    _shade(t1.rows[0].cells[0], C_TEAL)
    _cell_para(t1.rows[0].cells[0], "INSTRUCTIONS FOR STUDENTS",
               bold=True, size_pt=11, color_hex=C_WHITE)
    instructions = (
        f"You have completed all {meta['lessons']} lessons of {meta['substrand']}. "
        f"Write your COMPLETE EXPLANATION by answering all sections below.\n"
        f"Driving Question: {meta['driving_question']}\n\n"
        "Use evidence from investigations, discussions, and models from across this unit. "
        "Use scientific/mathematical vocabulary accurately. Include labeled diagrams where appropriate. "
        "Connect your explanations back to the anchoring phenomenon throughout."
    )
    _shade(t1.rows[1].cells[0], C_WHITE)
    _cell_para_auto(t1.rows[1].cells[0], instructions)
    _tbl_no_spacing(doc)

    # One table per FE section (3r × 1c)
    for section_data in meta["final_explanation_sections"]:
        section_title = section_data[0]
        prompts = section_data[1]
        model_answer = section_data[2] if len(section_data) > 2 else ""

        t = _new_table(doc, 3, 1)
        _col_widths(t, [9.5])

        _shade(t.rows[0].cells[0], C_MED_BLUE)
        _cell_para(t.rows[0].cells[0], section_title.upper(), bold=True, size_pt=11,
                   color_hex=C_WHITE)

        _shade(t.rows[1].cells[0], C_LT_BLUE)
        _cell_para_auto(t.rows[1].cells[0], prompts)

        _shade(t.rows[2].cells[0], C_WHITE)
        _cell_para_auto(t.rows[2].cells[0], model_answer)

        _tbl_no_spacing(doc)

    # Assessment rubric
    rubric = meta["final_explanation_rubric"]
    rub_rows = rubric["rows"]
    num_cols = len(rubric["headers"])
    remaining = 9.5 - 2.0
    other_w = round(remaining / (num_cols - 1), 3) if num_cols > 1 else remaining
    col_widths = [2.0] + [other_w] * (num_cols - 1)

    trub = _new_table(doc, 2 + len(rub_rows), num_cols)
    _col_widths(trub, col_widths)

    c = _merge_row(trub, 0)
    _shade(c, C_NAVY)
    _cell_para(c, "ASSESSMENT RUBRIC", bold=True, size_pt=11, color_hex=C_WHITE,
               align=WD_ALIGN_PARAGRAPH.CENTER)

    hdr_fills = [C_MED_BLUE, C_TEAL, C_MED_BLUE, C_TEAL, C_MED_BLUE][:num_cols]
    for cell, hdr, fill in zip(trub.rows[1].cells, rubric["headers"], hdr_fills):
        _shade(cell, fill)
        _cell_para(cell, hdr, bold=True, size_pt=9, color_hex=C_WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    alt = [C_WHITE, C_LT_GRAY, C_WHITE, C_LT_GRAY]
    for ri, row_data in enumerate(rub_rows):
        row = trub.rows[ri + 2]
        _shade(row.cells[0], C_LT_BLUE)
        _cell_para(row.cells[0], row_data[0], bold=True, size_pt=9)
        for cell, val, fill in zip(row.cells[1:], row_data[1:], alt):
            _shade(cell, fill)
            _cell_para_auto(cell, val)

    _tbl_no_spacing(doc)
    return doc


def build_summary_table(subject_key: str) -> "Document":
    """Build the Summary Table document for a May18 subject."""
    meta = ALL_META[subject_key]

    doc = Document()
    _set_page_landscape(doc)

    _doc_title(doc, f"{meta['subject']} | {meta['grade']} | {meta['substrand']}")
    _doc_subtitle(doc, "Summary Table")

    # Instructions table
    t_instr = _new_table(doc, 2, 1)
    _col_widths(t_instr, [9.5])
    _shade(t_instr.rows[0].cells[0], C_TEAL)
    _cell_para(t_instr.rows[0].cells[0], "SUMMARY TABLE — INSTRUCTIONS",
               bold=True, size_pt=11, color_hex=C_WHITE)
    _shade(t_instr.rows[1].cells[0], C_WHITE)
    _cell_para_auto(t_instr.rows[1].cells[0], meta.get("summary_instructions", ""))
    _tbl_no_spacing(doc)

    headers = meta["summary_table_headers"]
    rows = meta["summary_table_rows"]
    num_cols = len(headers)

    # Column widths: lesson number narrow, rest equal
    lesson_w = 0.7
    rest_w = round((9.5 - lesson_w) / (num_cols - 1), 3)
    col_ws = [lesson_w] + [rest_w] * (num_cols - 1)

    tsummary = _new_table(doc, 1 + len(rows), num_cols)
    _col_widths(tsummary, col_ws)

    # Header row
    hdr_fills = [C_NAVY, C_MED_BLUE, C_TEAL, C_MED_BLUE, C_TEAL][:num_cols]
    for cell, hdr, fill in zip(tsummary.rows[0].cells, headers, hdr_fills):
        _shade(cell, fill)
        _cell_para(cell, hdr, bold=True, size_pt=9, color_hex=C_WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    alt_fills = [C_WHITE, C_LT_GRAY]
    for ri, row_data in enumerate(rows):
        row = tsummary.rows[ri + 1]
        fill = alt_fills[ri % 2]
        for ci, (cell, val) in enumerate(zip(row.cells, row_data)):
            _shade(cell, fill if ci > 0 else C_LT_BLUE)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if ci == 0:
                _cell_para(cell, str(val), bold=True, size_pt=9,
                            align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                txt = str(val)
                if len(txt) > 80:
                    _cell_para_lines(cell, _prose_to_bullets(txt))
                else:
                    _cell_para_auto(cell, txt)

    _tbl_no_spacing(doc)
    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    for subj in ("biology", "physics", "chemistry", "maths"):
        paths = OUTPUTS[subj]
        print(f"\n=== {subj.upper()} ===")

        doc = build_lesson_sequence(subj)
        doc.save(paths["lesson_seq"])
        print(f"  Saved: {paths['lesson_seq'].name}")

        doc = build_final_explanation(subj)
        doc.save(paths["final_exp"])
        print(f"  Saved: {paths['final_exp'].name}")

        doc = build_summary_table(subj)
        doc.save(paths["summary_tbl"])
        print(f"  Saved: {paths['summary_tbl'].name}")

    print("\nDone. All 12 files written to", OUT_DIR)
