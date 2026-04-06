"""
Generate CBE Lesson Plans in Word (DOCX) Format
Much cleaner than HTML for tables and formatting
"""
import os
import sys
from pathlib import Path
from typing import Dict
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, str(Path(__file__).parent))
from template_loader import TemplateLoader

load_dotenv()


class DocxLessonGenerator:
    """Generate lessons in Word format"""
    
    def __init__(self):
        print("Initializing DOCX Lesson Generator...")
        self.template_loader = TemplateLoader()
        print("✓ Template loaded")
        self._init_claude()
    
    def _init_claude(self):
        """Initialize Claude API"""
        import anthropic
        self.anthropic_key = os.getenv("ANTHROPIC_API_KEY")
        if not self.anthropic_key:
            raise ValueError("Missing ANTHROPIC_API_KEY in .env")
        self.claude = anthropic.Anthropic(api_key=self.anthropic_key)
        print("✓ Claude API initialized")
    
    def generate_lesson_content(self, subject: str, grade: str, strand: str, 
                                substrand: str, lesson_topic: str, 
                                lesson_number: int, phenomenon: str = None) -> str:
        """Generate lesson content using Claude"""
        
        print(f"\n{'='*70}")
        print(f"LESSON {lesson_number}: {lesson_topic}")
        print(f"{'='*70}")
        
        # Get system prompt but modify for structured output
        system_prompt = f"""You are a CBE curriculum developer for Kenya.

Generate a lesson plan for:
- Subject: {subject}
- Grade: {grade}
- Topic: {lesson_topic}
- Phenomenon: {phenomenon or 'N/A'}

Return the lesson in this EXACT structure (plain text, no HTML):

===SECTION A: SPECIFIC LEARNING OUTCOMES===
Knowledge:
- Point 1
- Point 2

Skills:
- Point 1
- Point 2

Attitudes:
- Point 1

Science and Engineering Practices:
- Point 1

===SECTION B: OVERVIEW===
Key Inquiry Question: [Question here]
Purpose in Storyline: [Purpose here]
Safety Notes: [Notes or n/a]

===SECTION C: IMPLEMENTATION FRAMEWORK===
[For each row, use this format:]

ROW 1:
Learner Experience: [20+ words describing what students DO]
Resource Link: [List resources with ARES placeholders]
Teacher Moves: [Facilitation with probing questions, WAIT TIME noted]
Sensemaking Strategy: [Think-Pair-Share or other strategies]
Formative Assessment: [Checkboxes and quick checks]

ROW 2:
[Same format]

ROW 3:
[Same format - minimum 3 rows]

===SECTION D: TEACHER REFLECTION===
1. [Question 1]
2. [Question 2]
3. [Question 3]
4. [Question 4]
5. [Question 5]

===SECTION E: SUMMARY TABLE PROMPT===
Teacher says: "Add Lesson {lesson_number} to your Summary Table."

Prompt:
- What did I observe?
- What did I learn?
- How does this explain the phenomenon?

Use Kenyan context (KSh, KPLC, local examples). Include ARES resources. Use Think-Pair-Share."""
        
        print("  → Claude: Generating lesson...")
        
        response = self.claude.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=8000,
            messages=[{"role": "user", "content": system_prompt}]
        )
        
        print("  ✓ Lesson generated")
        return response.content[0].text
    
    def create_docx_lesson(self, content: str, lesson_number: int, 
                           lesson_topic: str, output_dir: str = "data/outputs/docx") -> str:
        """Convert structured text to formatted Word document"""
        
        print("  → Creating Word document...")
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Title
        title = doc.add_heading(f'Lesson {lesson_number}: {lesson_topic}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(0, 112, 192)
        
        # Subtitle
        subtitle = doc.add_paragraph('Kenya Competency-Based Curriculum')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(12)
        subtitle.runs[0].font.color.rgb = RGBColor(89, 89, 89)
        
        doc.add_paragraph()  # Spacing
        
        # Parse content sections
        sections = content.split('===SECTION')
        
        for section in sections[1:]:  # Skip first empty split
            if 'A: SPECIFIC LEARNING OUTCOMES' in section:
                self._add_outcomes_section(doc, section)
            elif 'B: OVERVIEW' in section:
                self._add_overview_section(doc, section)
            elif 'C: IMPLEMENTATION FRAMEWORK' in section:
                self._add_framework_section(doc, section)
            elif 'D: TEACHER REFLECTION' in section:
                self._add_reflection_section(doc, section)
            elif 'E: SUMMARY TABLE PROMPT' in section:
                self._add_summary_section(doc, section)
        
        # Save
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        filename = f"Lesson_{lesson_number:03d}_{lesson_topic.replace(' ', '_')}.docx"
        filepath = output_path / filename
        
        doc.save(str(filepath))
        print(f"  ✓ Saved: {filepath}")
        
        return str(filepath)
    
    def _add_outcomes_section(self, doc, section_text):
        """Add Learning Outcomes table"""
        heading = doc.add_heading('A. Specific Learning Outcomes', 2)
        heading.runs[0].font.color.rgb = RGBColor(0, 112, 192)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Outcomes'
        
        # Make header bold and colored
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            self._shade_cell(cell, RGBColor(112, 173, 71))
        
        # Parse categories
        categories = ['Knowledge:', 'Skills:', 'Attitudes:', 'Science and Engineering Practices:']
        for category in categories:
            if category in section_text:
                row_cells = table.add_row().cells
                row_cells[0].text = category.replace(':', '')
                
                # Extract bullet points
                start = section_text.find(category)
                end_markers = [c for c in categories if c != category]
                end = len(section_text)
                for marker in end_markers:
                    marker_pos = section_text.find(marker, start)
                    if marker_pos > start:
                        end = min(end, marker_pos)
                
                outcomes_text = section_text[start:end]
                outcomes = [line.strip('- ').strip() for line in outcomes_text.split('\n') 
                           if line.strip().startswith('-')]
                row_cells[1].text = '\n'.join(f'• {o}' for o in outcomes if o)
    
    def _add_overview_section(self, doc, section_text):
        """Add Overview table"""
        heading = doc.add_heading('B. Overview', 2)
        heading.runs[0].font.color.rgb = RGBColor(0, 112, 192)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Element'
        hdr_cells[1].text = 'Details'
        
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            self._shade_cell(cell, RGBColor(91, 155, 213))
        
        # Parse overview items
        items = [
            ('Key Inquiry Question', 'Key Inquiry Question:'),
            ('Purpose in Storyline', 'Purpose in Storyline:'),
            ('Safety Notes', 'Safety Notes:')
        ]
        
        for label, marker in items:
            if marker in section_text:
                row_cells = table.add_row().cells
                row_cells[0].text = label
                
                # Extract text after marker
                start = section_text.find(marker) + len(marker)
                end = section_text.find('===', start) if '===' in section_text[start:] else len(section_text)
                text = section_text[start:end].strip().split('\n')[0]
                row_cells[1].text = text
    
    def _add_framework_section(self, doc, section_text):
        """Add 5-column Implementation Framework table"""
        heading = doc.add_heading('C. Lesson Implementation Framework', 2)
        heading.runs[0].font.color.rgb = RGBColor(0, 112, 192)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Header
        headers = ['Learner Experience', 'Resource Link', 'Teacher Moves', 
                   'Sensemaking Strategy', 'Formative Assessment']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            self._shade_cell(hdr_cells[i], RGBColor(198, 89, 17))
        
        # Parse rows
        rows = section_text.split('ROW ')[1:]  # Split by ROW markers
        
        for row_text in rows:
            row_cells = table.add_row().cells
            
            # Extract each column
            columns = [
                ('Learner Experience:', 'Resource Link:'),
                ('Resource Link:', 'Teacher Moves:'),
                ('Teacher Moves:', 'Sensemaking Strategy:'),
                ('Sensemaking Strategy:', 'Formative Assessment:'),
                ('Formative Assessment:', 'ROW ')
            ]
            
            for i, (start_marker, end_marker) in enumerate(columns):
                if start_marker in row_text:
                    start = row_text.find(start_marker) + len(start_marker)
                    end = row_text.find(end_marker, start) if end_marker in row_text[start:] else len(row_text)
                    text = row_text[start:end].strip()
                    row_cells[i].text = text
    
    def _add_reflection_section(self, doc, section_text):
        """Add Teacher Reflection"""
        heading = doc.add_heading('D. Teacher Reflection', 2)
        heading.runs[0].font.color.rgb = RGBColor(0, 112, 192)
        
        # Extract numbered items
        lines = section_text.split('\n')
        for line in lines:
            if line.strip() and line.strip()[0].isdigit():
                doc.add_paragraph(line.strip(), style='List Number')
    
    def _add_summary_section(self, doc, section_text):
        """Add Summary Table Prompt"""
        heading = doc.add_heading('E. Summary Table Prompt', 2)
        heading.runs[0].font.color.rgb = RGBColor(0, 112, 192)
        
        # Add content
        lines = [l.strip() for l in section_text.split('\n') if l.strip()]
        for line in lines[1:]:  # Skip section header
            p = doc.add_paragraph(line)
            if line.startswith('- '):
                p.style = 'List Bullet'
    
    def _shade_cell(self, cell, color):
        """Add background color to table cell"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), f'{color.r:02x}{color.g:02x}{color.b:02x}')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def generate_lesson(self, subject: str, grade: str, strand: str,
                       substrand: str, lesson_topic: str, lesson_number: int = 1,
                       phenomenon: str = None) -> Dict:
        """Generate complete lesson in DOCX format"""
        
        # Generate content
        content = self.generate_lesson_content(
            subject, grade, strand, substrand, lesson_topic, lesson_number, phenomenon
        )
        
        # Create DOCX
        filepath = self.create_docx_lesson(content, lesson_number, lesson_topic)
        
        return {
            'lesson_number': lesson_number,
            'lesson_topic': lesson_topic,
            'filepath': filepath,
            'status': 'success'
        }


def main():
    """Main execution"""
    print("\n" + "="*70)
    print("CBE LESSON PLAN GENERATOR - WORD FORMAT")
    print("="*70 + "\n")
    
    generator = DocxLessonGenerator()
    
    result = generator.generate_lesson(
        subject="Physics",
        grade="Grade 10",
        strand="Electricity",
        substrand="Current Electricity",
        lesson_topic="Introduction to Current Electricity",
        lesson_number=1,
        phenomenon="Why do some electrical devices get hot when plugged in?"
    )
    
    print(f"\n{'='*70}")
    print("SUCCESS!")
    print(f"{'='*70}")
    print(f"Lesson saved to: {result['filepath']}")
    print(f"\nOpen in Microsoft Word to view and edit!")


if __name__ == "__main__":
    main()
