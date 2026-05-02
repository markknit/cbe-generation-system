"""
Reformat Seavuria Lesson Plan documents into the locked Biology Plant Nutrition format.

Produces 9 output files (3 doc types × 3 subjects):
  - CBE_LessonSequence.docx  — per-lesson 5-table blocks (A SLOs, B Materials,
                                C 6-col implementation, D Reflection, E Summary Prompt)
  - FinalExplanation.docx    — 8-table student document
  - SummaryTable.docx        — 3-table student tracking document

Reference format: data/outputs/docx/Grade 10 Bio/Bio 2.1 Plant Nutrition/docx/
"""

import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.table import Table as DocxTable

# ── paths ─────────────────────────────────────────────────────────────────────
SRC_DIR = Path("data/raw/Seavuria Lesson Plans")
OUT_DIR = Path("data/outputs/Seavuria_Reformatted")
OUT_DIR.mkdir(parents=True, exist_ok=True)

# ── color constants ───────────────────────────────────────────────────────────
C_NAVY      = "1F3864"   # titles, phase col header, Table A banner
C_TEAL      = "1F6B75"   # subtitles, Table B / resource / sensemaking headers
C_MED_BLUE  = "2E75B6"   # learner exp / teacher actions / assessment headers
C_PURPLE    = "7030A0"   # Table E header, summary DQB col header
C_ORANGE    = "C55A11"   # Table D header
C_LT_BLUE   = "D5E8F0"   # SLO labels, key inquiry bg, summary lesson# col
C_TEAL_LT   = "D9EEF1"   # Observe phase row
C_GREEN_LT  = "E2EFDA"   # Explain phase row
C_PURPLE_LT = "EAD1F5"   # Predict phase row, Table E row labels, DQB col
C_ORANGE_LT = "FCE4D6"   # DQB phase row, Table D content
C_LT_GRAY   = "F2F2F2"   # alternating content rows
C_WHITE     = "FFFFFF"

PHASE_FILLS = [C_PURPLE_LT, C_TEAL_LT, C_GREEN_LT, C_ORANGE_LT, C_LT_BLUE]
PHASE_NAMES = [
    "Predict Phase",
    "Observe Phase",
    "Explain Phase",
    "Driving Question Board (DQB) Creation",
    "Model Building Phase",
]

FONT = "Arial"

# ── META — subject metadata ───────────────────────────────────────────────────
META = {
    "chemistry": {
        "subject":    "Chemistry",
        "grade":      "Grade 10",
        "strand":     "Strand 1.0: Inorganic Chemistry",
        "substrand":  "Sub-Strand 1.4: Chemical Bonding",
        "lessons":    13,
        "driving_question": "How do the bonds between atoms determine the properties and uses of substances?",
        "phenomenon": (
            "A diamond and a graphite pencil tip are both made of pure carbon — "
            "yet diamond is the hardest natural substance and graphite is soft enough to write with. "
            "How can the same element behave so differently?"
        ),
        "substrand_overview_rows": [
            ("Subject",           "Chemistry"),
            ("Grade",             "Grade 10"),
            ("Strand",            "Strand 1.0: Inorganic Chemistry"),
            ("Sub-Strand",        "Sub-Strand 1.4: Chemical Bonding"),
            ("Number of Lessons", "13"),
            ("Number of Periods", "26"),
            ("Anchoring Phenomenon",
             "A diamond and a graphite pencil tip are both made of pure carbon — yet diamond is the "
             "hardest natural substance and graphite is soft enough to write with. How can the same element behave so differently?"),
            ("Driving Question",  "How do the bonds between atoms determine the properties and uses of substances?"),
            ("Sub-Strand SLOs: Knowledge",
             "Explain ionic, covalent, dative covalent, metallic, hydrogen, and Van der Waals bonding; "
             "describe the structures and properties of substances formed by each bond type."),
            ("Sub-Strand SLOs: Skills",
             "Draw Lewis dot diagrams; construct and interpret models of bonding; "
             "predict properties from structure; solve problems involving bonding and properties."),
            ("Sub-Strand SLOs: Attitudes & Values",
             "Appreciate how atomic-level bonding determines the macroscopic properties and uses of materials; "
             "connect chemistry to everyday Kenyan life."),
            ("Assessment",        "Formative: observation, discussion, models, summary table.\nSummative: Final Explanation document."),
            ("Core Values",       "Curiosity, critical thinking, connection to context"),
            ("Resources",         "Periodic table, molecular model kits, samples of ionic/covalent/metallic substances"),
            ("Notes",             ""),
        ],
        "final_explanation_sections": [
            (
                "SECTION 1: WHY DO ATOMS FORM BONDS?",
                "What makes an atom stable or unstable?\n"
                "What are valence electrons and why are they important?\n"
                "What is the octet rule? When does the duplet rule apply?\n"
                "How does bonding help atoms achieve stability?",
                "Atoms are stable when their outermost shell (valence shell) is full. "
                "Valence electrons are the electrons in the outermost shell — they determine how an atom bonds. "
                "Most atoms follow the octet rule: they become stable with 8 valence electrons (like noble gases). "
                "Hydrogen and helium follow the duplet rule: stable with 2 electrons.\n\n"
                "Carbon has 4 valence electrons (outer shell: 2,4) and needs 4 more to complete its octet. "
                "It achieves stability by forming 4 bonds with other atoms — either ionic or covalent depending on what it bonds with. "
                "Bonding allows atoms to reach the electron configuration of the nearest noble gas, which is the lowest-energy, most stable arrangement.",
            ),
            (
                "SECTION 2: HOW DO IONIC BONDS FORM AND WHAT PROPERTIES DO THEY GIVE SUBSTANCES?",
                "How do ionic bonds form through electron transfer?\n"
                "What types of atoms form ionic bonds?\n"
                "How can we represent ionic bonding using Lewis diagrams?\n"
                "What properties do ionic compounds have (melting point, conductivity, solubility, brittleness)?\n"
                "Why do ionic compounds have these properties?",
                "Ionic bonds form when a metal atom transfers one or more electrons to a non-metal atom. "
                "The metal loses electrons to form a positive ion (cation); the non-metal gains electrons to form a negative ion (anion). "
                "The strong electrostatic attraction between oppositely charged ions is the ionic bond.\n\n"
                "Example (Lewis diagram): Na· + ·Cl: → Na⁺ [:Cl:]⁻ — sodium transfers its 1 valence electron to chlorine, "
                "giving both the configuration of the nearest noble gas.\n\n"
                "Ionic compounds form giant ionic lattices — a regular 3D arrangement of millions of ions. "
                "This explains their properties:\n"
                "• High melting/boiling points: strong electrostatic forces require a lot of energy to break.\n"
                "• Conduct electricity when molten or dissolved (ions are free to move), but NOT as solids (ions locked in lattice).\n"
                "• Soluble in water: polar water molecules surround and separate the ions.\n"
                "• Brittle: shifting layers brings like charges together, creating repulsion that shatters the crystal.",
            ),
            (
                "SECTION 3: HOW DO COVALENT BONDS FORM AND WHAT STRUCTURES DO THEY CREATE?",
                "How do covalent bonds form through electron sharing?\n"
                "What is the difference between simple molecular and giant covalent structures?\n"
                "How do diamond and graphite differ in structure despite both being carbon?\n"
                "What properties do simple molecular substances have vs. giant covalent structures?\n"
                "Draw Lewis diagrams for simple molecules (H₂, H₂O, CO₂, CH₄).",
                "Covalent bonds form when two non-metal atoms share electrons to achieve full outer shells. "
                "Each shared pair of electrons counts toward the octet of both atoms.\n\n"
                "Lewis diagrams: H–H (H₂: 1 shared pair), H–O–H with 2 lone pairs on O (H₂O), "
                "O=C=O (CO₂: 2 double bonds), H–C–H with 4 single bonds (CH₄).\n\n"
                "Simple molecular structures (H₂O, CO₂, CH₄): small molecules held together by strong covalent bonds, "
                "but the forces BETWEEN molecules (Van der Waals / hydrogen bonds) are weak. "
                "Result: low melting/boiling points, do not conduct electricity.\n\n"
                "Giant covalent structures: millions of atoms bonded in a continuous 3D network.\n"
                "• Diamond: each carbon forms 4 strong covalent bonds to 4 other carbons in a tetrahedral structure. "
                "No free electrons → does not conduct. All bonds must break to melt → very high melting point → extremely hard.\n"
                "• Graphite: each carbon forms 3 covalent bonds in flat hexagonal layers; the 4th valence electron is delocalised "
                "and free to move between layers → conducts electricity. Layers are held by weak Van der Waals forces → "
                "layers slide over each other → soft and slippery.",
            ),
            (
                "SECTION 4: WHAT ARE DATIVE COVALENT, METALLIC, HYDROGEN, AND VAN DER WAALS BONDS?",
                "What is a dative covalent bond and how does it differ from a regular covalent bond?\n"
                "How do metallic bonds form and what properties do they give?\n"
                "What are hydrogen bonds and when do they form?\n"
                "What are Van der Waals forces and how do they affect properties?",
                "Dative covalent (coordinate) bond: both electrons in the shared pair come from the SAME atom "
                "(the donor). Example: the ammonium ion NH₄⁺ forms when NH₃ donates its lone pair to H⁺. "
                "Once formed, a dative bond is indistinguishable from a normal covalent bond.\n\n"
                "Metallic bonds: metal atoms release their valence electrons into a 'sea' of delocalised electrons. "
                "The positive metal ions are held together by electrostatic attraction to the electron sea. "
                "This explains: high melting points (strong metallic bonds), electrical conductivity (delocalised electrons carry charge), "
                "malleability/ductility (layers of ions slide without breaking bonds), and metallic lustre.\n\n"
                "Hydrogen bonds: a relatively strong intermolecular force forming between a δ+ hydrogen (bonded to N, O, or F) "
                "and a lone pair on N, O, or F of another molecule. They are NOT covalent bonds — they are between molecules. "
                "Hydrogen bonds give water its unusually high boiling point, surface tension, and solvent properties.\n\n"
                "Van der Waals forces: temporary dipoles caused by uneven electron distribution create weak attractions between all molecules. "
                "Stronger for larger molecules (more electrons). Explains why simple molecular substances have low but non-zero melting/boiling points.",
            ),
            (
                "SECTION 5: HOW DOES BOND TYPE EXPLAIN PHYSICAL PROPERTIES AND REAL-LIFE USES?",
                "Create a comparison table of ionic, simple molecular, giant covalent, and metallic substances "
                "(covering melting point, conductivity, hardness, solubility).\n"
                "Choose at least 4 substances and explain how their bonding and structure make them suitable for their uses "
                "(e.g. NaCl, diamond, graphite, copper, water).\n"
                "Answer the driving question: How do the bonds between atoms determine the properties and uses of substances?",
                "Comparison Table:\n"
                "| Type | Example | M.Pt | Electricity | Hardness | Water Solubility |\n"
                "| Ionic | NaCl | High | Yes (aq/l) | Hard, brittle | Soluble |\n"
                "| Simple Molecular | H₂O | Low | No | Soft | Variable |\n"
                "| Giant Covalent | Diamond | Very high | No (except graphite) | Very hard | Insoluble |\n"
                "| Metallic | Copper | High | Yes | Hard, malleable | Insoluble |\n\n"
                "Applications:\n"
                "• NaCl: used as food preservative and in medicine (saline drips) — dissolves in water, safe for consumption.\n"
                "• Diamond: used in cutting tools and drill bits — hardest natural substance due to strong 3D covalent network.\n"
                "• Graphite: used in pencils and as electrodes — soft (layers slide) and conducts electricity (delocalised electrons).\n"
                "• Copper: used in electrical wiring — excellent conductor, malleable, does not corrode easily.\n\n"
                "Driving Question Answer: The bonds between atoms determine atomic/molecular structure, which directly controls macroscopic properties. "
                "Diamond and graphite are both pure carbon — but in diamond, each carbon forms 4 bonds in a 3D tetrahedral network "
                "(making it the hardest substance), while in graphite, each carbon forms 3 bonds in flat layers with delocalised electrons "
                "(making it soft and conductive). The bond type IS the structure IS the property.",
            ),
            (
                "SECTION 6: HOW DOES BOND TYPE AND STRUCTURE EXPLAIN THE USES OF SUBSTANCES IN EVERYDAY LIFE?",
                "How do the properties of substances determine their uses?\n"
                "Can you explain specific real-world applications using bonding concepts?\n"
                "Choose at least 4 substances and explain how their bonding and structure make them suitable for their uses.\n"
                "For each substance explain: what type of bonding it has; what structure this creates; what properties result; "
                "how these properties make it useful.",
                "Substance Applications:\n\n"
                "Sodium chloride (NaCl) — Table salt and de-icing roads:\n"
                "Ionic bonding → giant ionic lattice → soluble in water (ions dissociate), safe to consume, lowers freezing point of water. "
                "Used as food preservative, in oral rehydration salts (saline drips), and to melt ice on roads in cold climates.\n\n"
                "Diamond — Cutting tools, drill bits, jewelry:\n"
                "Giant covalent structure — each carbon bonded to 4 others in a 3D tetrahedral network → "
                "all bonds must break to scratch or melt it → extremely hard (hardest natural substance, Mohs 10) and very high melting point. "
                "Used in industrial cutting tools, drill bits for oil exploration, and gemstones.\n\n"
                "Graphite — Pencil leads, lubricants, electrodes:\n"
                "Layered giant covalent structure — each carbon forms 3 bonds in flat hexagonal layers; "
                "one delocalised electron per carbon moves freely between layers → conducts electricity; "
                "layers held by weak Van der Waals forces → slide easily (soft and lubricating). "
                "Used in pencil leads (leaves a mark when layers slide onto paper), industrial lubricants, and as electrodes in electrolysis.\n\n"
                "Metals (copper, aluminium, iron) — Electrical wiring, construction, cookware:\n"
                "Metallic bonding — sea of delocalised electrons around positive metal ions → "
                "electrons carry charge freely (excellent conductor); ions can slide past each other (malleable and ductile); "
                "strong metallic bonds → high melting points and structural strength.\n\n"
                "Water (H₂O) — Universal solvent, biological processes:\n"
                "Polar covalent bonds + hydrogen bonding between molecules → "
                "unusually high boiling point for such a small molecule (100°C), remains liquid at room temperature, "
                "excellent solvent for ionic and polar substances. Essential for life and all biochemical reactions.\n\n"
                "Silicon dioxide (SiO₂) — Glass, electronics, construction:\n"
                "Giant covalent structure (like diamond but with O atoms between Si atoms) → "
                "very hard, very high melting point (~1700°C), transparent when pure, chemically inert. "
                "Used in glass making, optical fibres, microchip substrates, and as a construction material (sand, concrete).",
            ),
        ],
        "final_explanation_rubric": {
            "headers": ["Criterion", "Beginning (1)", "Developing (2)", "Proficient (3)", "Advanced (4)"],
            "rows": [
                ("1. Completeness",
                 "Addresses fewer than 4 sections; major gaps in content",
                 "Addresses 4–5 sections; some important concepts missing",
                 "Addresses all 6 sections with all key concepts included",
                 "Addresses all 6 sections comprehensively with depth and additional connections"),
                ("2. Scientific Accuracy",
                 "Multiple significant errors or misconceptions in bonding concepts",
                 "Some minor errors; generally accurate but may have 2–3 misconceptions",
                 "Scientifically accurate throughout; no significant errors",
                 "Scientifically accurate with sophisticated understanding; addresses exceptions and nuance"),
                ("3. Explanation Quality",
                 "Mostly describes what happens without explaining why; limited cause-effect",
                 "Provides some explanations but reasoning is incomplete or unclear",
                 "Clearly explains how bonding determines properties using cause-and-effect throughout",
                 "Deep, sophisticated explanations with multiple lines of reasoning and original synthesis"),
                ("4. Use of Evidence & Examples",
                 "Few or no specific examples; claims lack support",
                 "Some examples provided but not well-explained or connected to bonding theory",
                 "Appropriate examples throughout; evidence clearly supports claims",
                 "Rich, well-chosen examples; evidence is thoroughly analysed and connected to theory"),
                ("5. Structure-Property Connections",
                 "Does not clearly connect atomic/molecular structure to macroscopic properties",
                 "Makes some connections but they are superficial or incomplete",
                 "Consistently and clearly connects structure to properties across all bond types",
                 "Makes sophisticated connections; explains subtle differences between similar structures"),
                ("6. Scientific Vocabulary",
                 "Limited use of scientific terms; frequent use of everyday language instead",
                 "Uses some scientific vocabulary but inconsistently or occasionally incorrectly",
                 "Uses scientific vocabulary accurately and appropriately throughout",
                 "Uses sophisticated scientific vocabulary precisely; demonstrates mastery of terminology"),
                ("7. Visual Representations",
                 "Missing diagrams or diagrams are inaccurate/unlabeled",
                 "Some diagrams included but may lack labels, clarity, or accuracy",
                 "Appropriate diagrams included (Lewis structures, structural models); clearly labeled",
                 "Excellent diagrams that enhance understanding; creative and scientifically sophisticated"),
                ("8. Communication & Organisation",
                 "Difficult to follow; lacks clear organisation; many writing errors",
                 "Somewhat organised; writing is understandable but may be unclear in places",
                 "Well-organised and clear; easy to follow; few writing errors",
                 "Exceptionally clear and engaging; sophisticated organisation; polished, error-free writing"),
            ],
        },
        "summary_table_headers": [
            "Lesson # and Title",
            "What did I observe or investigate?",
            "What did I learn?",
            "How does this help explain the diamond/graphite phenomenon?",
            "DQB and Model Update",
        ],
        "summary_table_rows": [
            (
                "Lesson 1: Phenomenon Launch and Atomic Stability",
                "Observed diamond (hard, transparent) and graphite (soft, black) — both pure carbon but very different. "
                "Investigated atomic structure and electron arrangements.",
                "Atoms bond to achieve stable electron configurations. Carbon has 4 valence electrons and needs 4 more. "
                "The bonding pattern determines properties.",
                "Both are carbon, so the difference must come from how atoms are bonded/arranged, not from different atoms.",
                "DQB Started: Why are diamond and graphite so different? How do atoms stick together?\n"
                "Initial Model: Carbon atoms connecting — unclear how or why differently.",
            ),
            (
                "Lesson 2: Ionic Bonding — Electron Transfer",
                "Investigated how metals transfer electrons to non-metals. Used Lewis diagrams to show sodium donating "
                "its electron to chlorine. Identified conditions that favour ionic bond formation.",
                "Ionic bonds form when a metal transfers electrons to a non-metal, creating oppositely charged ions. "
                "The electrostatic attraction between ions is the ionic bond. Metals lose electrons; non-metals gain them.",
                "Carbon bonds with itself (non-metal to non-metal), so it cannot form ionic bonds — ionic bonding is not "
                "responsible for the diamond/graphite difference. The difference must involve covalent bonding.",
                "DQB Updated: Answered 'What is an ionic bond?' Added question: Can carbon form ionic bonds?\n"
                "Model Revised: Ionic bonding ruled out for carbon. Carbon must share electrons (covalent).",
            ),
            (
                "Lesson 3: Ionic Compounds — Properties and Structure",
                "Investigated properties of ionic compounds: melting point, solubility, electrical conductivity. "
                "Built a model of a crystal lattice using balls and sticks.",
                "Ionic compounds form giant ionic lattices — millions of ions in a regular 3D arrangement. "
                "This gives them high melting points, brittleness, and ability to conduct electricity when dissolved or molten.",
                "Ionic compounds like NaCl have fixed, ordered structures. Carbon-based structures must work differently — "
                "carbon forms networks, not ionic lattices. The structure of carbon must explain its unique properties.",
                "DQB Updated: Added 'How does structure explain properties?' \n"
                "Model Revised: Ionic lattice structure shown. Carbon model still unclear — needs covalent investigation.",
            ),
            (
                "Lesson 4: Covalent Bonding — Electron Sharing",
                "Investigated how non-metals share electrons to form covalent bonds. Drew Lewis diagrams for H₂, H₂O, CO₂, CH₄. "
                "Explored single, double, and triple bonds.",
                "Covalent bonds form when non-metal atoms share electron pairs. Each shared pair is one bond. "
                "Both atoms count the shared electrons toward their octet. Carbon can form 4 covalent bonds.",
                "Carbon forms 4 covalent bonds because it has 4 valence electrons and needs 4 more. "
                "This means carbon can bond to 4 other carbons — building a giant 3D network. This is the key to understanding diamond!",
                "DQB Updated: Added 'How does carbon bond to itself?' \n"
                "Model Revised: Carbon atom with 4 bond sites shown. Starting to see how diamond network forms.",
            ),
            (
                "Lesson 5: Simple Molecular vs Giant Covalent Structures",
                "Compared properties of simple molecular substances (e.g. water, iodine) with giant covalent structures (diamond, silica). "
                "Investigated why some covalent substances have low melting points and others extremely high.",
                "Simple molecular: small molecules with strong intramolecular bonds but weak intermolecular forces → low melting points. "
                "Giant covalent: atoms bonded in a continuous network → all bonds must break to melt → very high melting points.",
                "Diamond must be a giant covalent structure — it melts at ~3500°C, far too high for a simple molecular substance. "
                "Graphite is also giant covalent but has a different structure that makes it soft. The arrangement of bonds matters!",
                "DQB Updated: Answered 'Why does diamond have such a high melting point?' \n"
                "Model Revised: Diamond shown as giant covalent 3D network. Graphite structure still to investigate.",
            ),
            (
                "Lesson 6: Diamond and Graphite — Structure Explains Properties",
                "Built 3D models of diamond and graphite structures. Compared bonding arrangements. "
                "Investigated why graphite conducts electricity but diamond does not.",
                "Diamond: each carbon forms 4 bonds to 4 others in a tetrahedral 3D network — no free electrons, very hard, "
                "does not conduct. Graphite: each carbon forms 3 bonds in flat hexagonal layers — the 4th electron is delocalised "
                "between layers, conducts electricity, layers slide easily (soft and lubricating).",
                "The mystery is SOLVED: same element (carbon) but different bonding arrangements → completely different properties. "
                "Diamond's 3D bonding makes it the hardest substance. Graphite's layered structure with free electrons makes it soft and conductive.",
                "DQB Updated: Answered central driving question for carbon! Diamond = 3D covalent, Graphite = layered covalent.\n"
                "Model Revised: Both structures drawn accurately showing bond arrangements and explaining properties.",
            ),
            (
                "Lesson 7: Dative Covalent Bonding",
                "Investigated dative covalent bonds where both electrons come from one donor atom. "
                "Formed ammonium ion (NH₄⁺) by NH₃ donating its lone pair to H⁺. Drew coordinate bond diagrams.",
                "A dative (coordinate) covalent bond is the same as a normal covalent bond once formed — the only difference "
                "is the origin of the electrons (both from one atom). It creates no special properties distinct from normal covalent bonds.",
                "Dative bonds further expand our understanding of how carbon compounds can form complex structures "
                "(e.g. in organic chemistry, coordinate bonds appear in many compounds). Same principle: bond type → structure → property.",
                "DQB Updated: Added 'Are there other types of covalent bonds?' Answered with dative bonding.\n"
                "Model Revised: Added dative bond notation to model toolkit.",
            ),
            (
                "Lesson 8: Metallic Bonding",
                "Investigated metallic bonding using the electron sea model. Tested properties of metals: conductivity, "
                "malleability, ductility, high melting points.",
                "Metallic bonds: metal atoms release valence electrons into a 'sea' of delocalised electrons. "
                "Positive ions are held by electrostatic attraction to the electron sea. "
                "This explains conductivity (free electrons), malleability (ions slide without breaking bonds), and high melting points.",
                "Metals have a completely different bonding model from carbon — electron sea vs covalent network. "
                "Bond type determines structure determines property: this principle applies to ALL substances, not just carbon.",
                "DQB Updated: Added 'How do metals conduct electricity?' Answered via electron sea model.\n"
                "Model Revised: Metallic bonding diagram added alongside ionic and covalent models.",
            ),
            (
                "Lesson 9: Hydrogen Bonds and Van der Waals Forces",
                "Investigated intermolecular forces — forces BETWEEN molecules (not within). "
                "Tested water's surface tension, high boiling point, and ice floating as evidence of hydrogen bonding.",
                "Hydrogen bonds are stronger intermolecular forces (not covalent bonds!) between δ+H and lone pairs on N, O, F. "
                "Van der Waals forces are weak temporary dipole attractions between all molecules — increase with molecular size. "
                "These explain boiling points of simple molecular substances.",
                "Even for simple molecular carbon compounds (like CH₄), Van der Waals forces between molecules determine "
                "physical properties. The stronger the intermolecular forces, the higher the boiling point — even in simple molecules.",
                "DQB Updated: Added 'What forces exist between molecules?' Answered with hydrogen bonds + Van der Waals.\n"
                "Model Revised: Added intermolecular forces layer to model — both intramolecular and intermolecular bonds matter.",
            ),
            (
                "Lesson 10: Comparing Bond Types and Properties",
                "Constructed a comprehensive comparison table of all bond types. Matched substances to their bond types "
                "by reasoning from their properties (melting point, conductivity, hardness, solubility).",
                "Each bond type produces a predictable set of properties: ionic → high M.Pt, conducts when dissolved; "
                "metallic → conducts, malleable; giant covalent → very high M.Pt, very hard, non-conducting (except graphite); "
                "simple molecular → low M.Pt, non-conducting.",
                "This lesson consolidates the answer to the driving question: bond type → structure → properties. "
                "Diamond (giant covalent) and graphite (giant covalent but different arrangement) perfectly fit predicted properties.",
                "DQB Updated: Created comprehensive bond-type → property map.\n"
                "Model Revised: Final comparison model with all four bond types shown side by side.",
            ),
            (
                "Lesson 11: Real-Life Applications of Chemical Bonding",
                "Researched applications of different bond types in everyday Kenyan life and technology. "
                "Presented case studies: NaCl in food, diamond in industry, copper in wiring, water in biology.",
                "Bond type determines suitability for specific uses: NaCl (ionic) dissolves in water — good for food and medicine; "
                "diamond (giant covalent) — hardest substance, used for cutting; copper (metallic) — malleable and conducts; "
                "graphite (giant covalent, layered) — used as lubricant and electrodes.",
                "The driving question is answered in real-world context: understanding bonding allows us to choose the right "
                "material for the right job. Diamond and graphite are perfect examples: same atoms, different bonds, different uses.",
                "DQB Updated: Answered 'How does this help in real life?' with specific Kenya examples.\n"
                "Model Revised: Real-world applications connected to bond type in final model.",
            ),
            (
                "Lesson 12: Review and Consolidation",
                "Reviewed all bond types and their properties through a class quiz and collaborative model building. "
                "Identified and corrected misconceptions. Prepared for Final Explanation.",
                "All five bond types (ionic, covalent, dative covalent, metallic, hydrogen/Van der Waals) explain specific "
                "properties through specific structures. The linking principle: atoms arrange themselves in the lowest-energy "
                "configuration, which is determined by how they bond.",
                "Full answer to driving question: Carbon's 4 valence electrons allow it to bond 4 ways — in diamond, "
                "these form a 3D tetrahedral network (hardest substance); in graphite, they form flat hexagonal layers "
                "(soft, lubricating, conductive). Bond type → structure → property.",
                "DQB Updated: All questions answered. DQB Complete.\n"
                "Model Revised: Complete and accurate model showing all bond types with property explanations.",
            ),
            (
                "Lesson 13: Final Explanation and Model",
                "Completed individual Final Explanation document synthesising all learning across 13 lessons. "
                "Presented and explained final model of chemical bonding.",
                "A complete understanding of chemical bonding — from atomic stability through all five bond types to "
                "real-world applications — explains why substances have the properties they do.",
                "Diamond and graphite are both pure carbon, but their completely different properties arise entirely from "
                "different bonding arrangements: 4-bond 3D network (diamond) vs 3-bond layered structure with delocalised electrons (graphite).",
                "DQB Completed: Answered driving question with evidence from all lessons.\n"
                "Final Model: Comprehensive model showing how atomic bonding determines substance properties.",
            ),
        ],
        "summary_instructions": (
            "Complete one row after each lesson. Be specific and use evidence from investigations. "
            "Track your Driving Question Board (DQB) — note new questions and when questions get answered. "
            "Update your models progressively as new evidence changes your thinking."
        ),
    },

    "mathematics": {
        "subject":    "Mathematics",
        "grade":      "Grade 10",
        "strand":     "Strand 1.0: Numbers and Algebra",
        "substrand":  "Sub-Strand 1.3: Quadratic Expressions and Equations",
        "lessons":    7,
        "driving_question": "How can we use quadratic equations to model, solve, and make sense of real-life problems in our communities?",
        "phenomenon": (
            "A harambee group plans to raise Ksh 12,000. Members share the cost equally, but 3 members fail to show up — "
            "so the remaining members each pay Ksh 200 more. How many people were originally in the group?"
        ),
        "substrand_overview_rows": [
            ("Subject",           "Mathematics"),
            ("Grade",             "Grade 10"),
            ("Strand",            "Strand 1.0: Numbers and Algebra"),
            ("Sub-Strand",        "Sub-Strand 1.3: Quadratic Expressions and Equations"),
            ("Number of Lessons", "7"),
            ("Number of Periods", "14"),
            ("Anchoring Phenomenon",
             "A harambee group plans to raise Ksh 12,000. Members share the cost equally, but 3 members fail to show up — "
             "so the remaining members each pay Ksh 200 more. How many people were originally in the group?"),
            ("Driving Question",
             "How can we use quadratic equations to model, solve, and make sense of real-life problems in our communities?"),
            ("Sub-Strand SLOs: Knowledge",
             "Define and identify quadratic expressions; state and derive the three quadratic identities; "
             "explain factorisation methods; describe the zero product property; "
             "form quadratic equations from real-world contexts."),
            ("Sub-Strand SLOs: Skills",
             "Expand and factorise quadratic expressions; solve quadratic equations by factorisation; "
             "interpret solutions in context; model real-life problems with quadratic equations."),
            ("Sub-Strand SLOs: Attitudes & Values",
             "Appreciate mathematics as a tool for solving real community problems; "
             "develop persistence and systematic problem-solving habits."),
            ("Assessment",        "Formative: classwork, summary table, problem sets.\nSummative: Final Explanation document."),
            ("Core Values",       "Curiosity, perseverance, community connection"),
            ("Resources",         "Graph paper, area model tiles, calculators"),
            ("Notes",             ""),
        ],
        "final_explanation_sections": [
            (
                "SECTION 1: WHAT IS A QUADRATIC EXPRESSION AND HOW DO WE FORM ONE?",
                "What is the standard form ax² + bx + c and what do a, b, c represent?\n"
                "Give three examples of quadratic expressions from real-life Kenyan contexts "
                "(e.g. shamba area, harambee contributions, matatu revenue).\n"
                "Explain how the harambee problem leads to a quadratic relationship.",
                "A quadratic expression has the standard form ax² + bx + c where:\n"
                "• a = coefficient of the squared term (a ≠ 0)\n"
                "• b = coefficient of the linear term\n"
                "• c = constant term\n\n"
                "Three Kenyan real-life examples:\n"
                "1. Shamba area: if a rectangular plot has length (x + 5) m and width (x + 3) m, "
                "area = (x+5)(x+3) = x² + 8x + 15 m². (a=1, b=8, c=15)\n"
                "2. Harambee contributions: if x people share Ksh 12,000 but 3 fail to show, "
                "cost per person increases by Ksh 200 → equation involves 12,000/x and 12,000/(x-3) → leads to x² - 3x - 180 = 0.\n"
                "3. Matatu revenue: if a matatu charges Ksh (100-2x) per passenger and carries (50+x) passengers, "
                "revenue = (100-2x)(50+x) = 5000 + 100x - 100x - 2x² = 5000 - 2x²; to maximise revenue involves a quadratic.\n\n"
                "Harambee problem as quadratic: Let x = original number of members. "
                "Each originally pays 12,000/x. After 3 leave, each pays 12,000/(x-3). "
                "Difference = 200: 12,000/(x-3) - 12,000/x = 200. "
                "Multiply through by x(x-3): 12,000x - 12,000(x-3) = 200x(x-3). "
                "Simplify: 36,000 = 200x² - 600x → x² - 3x - 180 = 0. This IS a quadratic equation!",
            ),
            (
                "SECTION 2: WHAT ARE THE THREE QUADRATIC IDENTITIES AND HOW ARE THEY DERIVED FROM AREA?",
                "State and prove all three identities using area diagrams:\n"
                "  (a + b)² = a² + 2ab + b²\n"
                "  (a - b)² = a² - 2ab + b²\n"
                "  (a + b)(a - b) = a² - b²\n"
                "Use a numerical example for each identity (e.g. calculate 103², 97², and 103 × 97).",
                "Identity 1: (a + b)² = a² + 2ab + b²\n"
                "Proof (area method): Draw a square with side (a+b). "
                "Divide into 4 regions: a²  (top-left), ab (top-right), ab (bottom-left), b² (bottom-right). "
                "Total area = a² + ab + ab + b² = a² + 2ab + b².\n"
                "Numerical example: 103² = (100+3)² = 10,000 + 600 + 9 = 10,609 ✓\n\n"
                "Identity 2: (a - b)² = a² - 2ab + b²\n"
                "Proof: (a-b)² = (a-b)(a-b) = a²-ab-ab+b² = a²-2ab+b².\n"
                "Numerical example: 97² = (100-3)² = 10,000 - 600 + 9 = 9,409 ✓\n\n"
                "Identity 3: (a + b)(a - b) = a² - b²\n"
                "Proof: expand (a+b)(a-b) = a²-ab+ab-b² = a²-b².\n"
                "Numerical example: 103 × 97 = (100+3)(100-3) = 100²-3² = 10,000-9 = 9,991 ✓\n\n"
                "These identities allow us to expand brackets quickly without full multiplication, "
                "and they reverse into factorisation patterns (e.g. x²-9 = (x+3)(x-3)).",
            ),
            (
                "SECTION 3: HOW DO WE FACTORISE QUADRATIC EXPRESSIONS?",
                "Explain the method for factorising when a = 1 (find two numbers that multiply to c and add to b).\n"
                "Explain the splitting method for factorising when a ≠ 1.\n"
                "Apply both methods with worked examples.\n"
                "Identify and factorise perfect square trinomials and differences of two squares.",
                "Method 1 (a = 1): Find two numbers p and q such that p × q = c AND p + q = b. "
                "Then x² + bx + c = (x + p)(x + q).\n"
                "Example: x² + 7x + 12. Need p × q = 12, p + q = 7 → p = 3, q = 4. "
                "So x² + 7x + 12 = (x + 3)(x + 4). Check: (x+3)(x+4) = x² + 4x + 3x + 12 = x² + 7x + 12 ✓\n\n"
                "Method 2 — Splitting (a ≠ 1): For ax² + bx + c, find p and q with p × q = a×c and p + q = b. "
                "Rewrite bx as px + qx, then factor by grouping.\n"
                "Example: 2x² + 7x + 3. a×c = 6. Need p+q=7, p×q=6 → p=1, q=6. "
                "Rewrite: 2x² + x + 6x + 3 = x(2x+1) + 3(2x+1) = (2x+1)(x+3). "
                "Check: (2x+1)(x+3) = 2x² + 6x + x + 3 = 2x² + 7x + 3 ✓\n\n"
                "Special cases:\n"
                "• Perfect square trinomial: x² + 6x + 9 = (x+3)² (because 3²=9 and 2×3=6).\n"
                "• Difference of two squares: x² - 16 = (x+4)(x-4) (because 4²=16).\n"
                "For the harambee equation x² - 3x - 180 = 0: need p×q=-180, p+q=-3 → p=12, q=-15. "
                "So x² - 3x - 180 = (x+12)(x-15).",
            ),
            (
                "SECTION 4: HOW DO WE FORM AND SOLVE A QUADRATIC EQUATION?",
                "State the four steps to solve a quadratic equation by factorisation.\n"
                "Solve the harambee textbook fund problem completely, showing all steps:\n"
                "  — define variable, form equation, simplify to standard form, factorise, "
                "apply zero product property, interpret, verify.\n"
                "Explain why one solution is rejected and what the answer means in context.",
                "Four steps to solve by factorisation:\n"
                "1. Rearrange to standard form: ax² + bx + c = 0\n"
                "2. Factorise the left side completely\n"
                "3. Apply zero product property: if AB = 0, then A = 0 OR B = 0\n"
                "4. Solve each factor equation, then interpret and verify in context\n\n"
                "Harambee Problem (full solution):\n"
                "Step 1 — Define variable: Let x = original number of members.\n"
                "Step 2 — Form equation: Each originally pays 12,000/x. After 3 leave, remaining x-3 members each pay "
                "12,000/(x-3). The increase is Ksh 200: 12,000/(x-3) - 12,000/x = 200\n"
                "Step 3 — Simplify: Multiply both sides by x(x-3):\n"
                "  12,000x - 12,000(x-3) = 200x(x-3)\n"
                "  12,000x - 12,000x + 36,000 = 200x² - 600x\n"
                "  36,000 = 200x² - 600x\n"
                "  Divide by 200: 180 = x² - 3x → x² - 3x - 180 = 0\n"
                "Step 4 — Factorise: (x + 12)(x - 15) = 0\n"
                "Step 5 — Solve: x = -12 or x = 15\n"
                "Step 6 — Reject x = -12 (impossible: number of people cannot be negative)\n"
                "Step 7 — Interpret: Original group = 15 members. Each paid 12,000/15 = Ksh 800.\n"
                "After 3 left, 12 members each paid 12,000/12 = Ksh 1,000. Increase = 200. ✓",
            ),
            (
                "SECTION 5: HOW DO QUADRATIC EQUATIONS HELP US SOLVE REAL-LIFE PROBLEMS?",
                "Solve at least two of the following problems, showing all working:\n"
                "  — Shamba area problem (farmer with 100m fencing)\n"
                "  — Classroom foundation dimensions\n"
                "  — Matatu fare optimization\n"
                "  — Pythagorean theorem application\n"
                "For each: define variable, form equation, solve, interpret solution, verify answer.",
                "Problem 1 — Shamba area:\n"
                "A farmer has 100 m of fencing for a rectangular plot against a wall (only 3 sides needed). "
                "The area must be 1,200 m². Find the dimensions.\n"
                "Let width = x, then length = 100 - 2x.\n"
                "Area: x(100-2x) = 1,200 → 100x - 2x² = 1,200 → 2x² - 100x + 1,200 = 0 → x² - 50x + 600 = 0\n"
                "Factorise: (x-20)(x-30) = 0 → x = 20 or x = 30.\n"
                "If x = 20: width = 20 m, length = 60 m. Area = 1,200 ✓\n"
                "If x = 30: width = 30 m, length = 40 m. Area = 1,200 ✓\n"
                "Both solutions are valid — two possible rectangles fit the constraints.\n\n"
                "Problem 2 — Pythagorean theorem:\n"
                "A ladder 10 m long leans against a wall. The foot is 2 m further from the wall than the height it reaches. "
                "Find the height.\n"
                "Let height = h. Foot distance = h + 2. By Pythagoras: h² + (h+2)² = 10²\n"
                "h² + h² + 4h + 4 = 100 → 2h² + 4h - 96 = 0 → h² + 2h - 48 = 0\n"
                "Factorise: (h+8)(h-6) = 0 → h = -8 (reject, negative) or h = 6 m.\n"
                "The ladder reaches 6 m up the wall. Foot is 8 m from wall. Check: 6² + 8² = 36+64 = 100 = 10² ✓\n\n"
                "Quadratic equations are powerful tools for modelling real-world situations involving area, distance, "
                "rates, and revenue. The process — define variable, form equation, solve, interpret, verify — "
                "applies across all contexts.",
            ),
        ],
        "final_explanation_rubric": {
            "headers": ["Criterion", "Below Expectation", "Approaches Expectation", "Meets Expectation", "Exceeds Expectation"],
            "rows": [
                ("1. Understanding Quadratic Concepts",
                 "Cannot identify quadratic expressions; explanations unclear or with major misconceptions",
                 "Identifies quadratics; states identities but struggles to explain derivations",
                 "Solid understanding; explains standard form, derives identities from area, describes zero product property",
                 "Deep understanding; uses multiple representations; makes connections across concepts"),
                ("2. Factorisation Skills",
                 "Cannot reliably factorise; attempts lack clear method; frequent errors",
                 "Factorises simple (a=1) cases with some success; struggles with a≠1",
                 "Factorises most expressions correctly; handles a=1 and a≠1; recognizes some special cases",
                 "Factorises all types accurately and efficiently; explains strategy; recognizes patterns; checks by expanding"),
                ("3. Problem-Solving and Equation Solving",
                 "Cannot set up equations from word problems; misapplies zero product property",
                 "Solves simple equations with guidance; struggles to set up from word problems",
                 "Solves equations correctly; defines variables; interprets solutions; shows steps",
                 "Excellent equation solving with clear variable definition, all steps shown, context interpretation, and verification"),
                ("4. Real-World Application and Communication",
                 "Does not connect maths to real-world contexts; communication unclear; no units",
                 "Some connections but superficial; communication unclear in places; inconsistent notation",
                 "Connects equations to real-world contexts; communicates clearly; uses correct notation and units",
                 "Strong connections with reflection on meaning; polished communication; organized presentation"),
            ],
        },
        "summary_table_headers": [
            "Lesson # and Title",
            "What did I investigate or work on?",
            "What did I learn?",
            "How does this help solve the harambee problem?",
            "DQB and Model Update",
        ],
        "summary_table_rows": [
            (
                "Lesson 1: Introduction to Quadratic Expressions",
                "Exploring patterns in area models; recognizing expressions with squared terms; sorting algebraic expressions.",
                "A quadratic expression contains a term with a variable squared (x²) written as ax² + bx + c. "
                "Quadratics model situations involving area, growth, and non-linear relationships.",
                "The harambee problem involves a relationship between contributions and total amount. "
                "Recognizing this as a quadratic situation is the first step.",
                "DQB Started: Posted harambee problem and initial questions.\n"
                "Model: Identified variables — let x = original number of members.",
            ),
            (
                "Lesson 2: Expanding Quadratic Expressions",
                "Practicing the distributive property; expanding products like (x + 3)(x + 5) using area models.",
                "Expanding means multiplying out brackets to write a quadratic in standard form. "
                "Used FOIL and area diagrams.",
                "If I set up an equation for the harambee problem, I may need to expand brackets to simplify it.",
                "DQB Updated: Added 'How do we expand quadratic expressions?'\n"
                "Model: Practiced expanding expressions to prepare for equation setup.",
            ),
            (
                "Lesson 3: Factorising Quadratic Expressions",
                "Breaking down quadratic expressions into products of binomials; finding factor pairs; using the ac method.",
                "Factorising is the reverse of expanding. Find two numbers that multiply to give ac and add to give b, "
                "then rewrite and group. Factorising helps solve equations.",
                "Once I form a quadratic equation from the harambee problem, factorising will allow me to find the possible values.",
                "DQB Updated: Added 'How do we factorise quadratics?'\n"
                "Model: Explored factorising sample equations.",
            ),
            (
                "Lesson 4: Solving Quadratic Equations by Factorising",
                "Setting equations equal to zero; factorising; applying the zero-product property.",
                "To solve: (1) rearrange to standard form, (2) factorise, (3) set each factor to zero, "
                "(4) solve for the variable. Must check solutions in context.",
                "This is the key method I'll use to solve the harambee equation.",
                "DQB Updated: Added 'How do we solve quadratic equations?'\n"
                "Model: Practiced solving and checking for extraneous solutions.",
            ),
            (
                "Lesson 5: Formation of Quadratic Equations from Real-World Contexts",
                "Translating word problems into algebraic equations; defining variables; writing equations from context.",
                "Forming equations requires: (1) define a variable, (2) express other quantities in terms of it, "
                "(3) use relationships to write an equation, (4) simplify to standard form.",
                "This lesson directly addresses the harambee problem. "
                "Equation: 12000/(x−3) − 12000/x = 200.",
                "DQB Updated: Added 'How do we translate real-world problems into equations?'\n"
                "Model Revised: Let x = original members. Set up the harambee equation.",
            ),
            (
                "Lesson 6: Solving and Interpreting Quadratic Equations in Context",
                "Solving equations formed in Lesson 5; interpreting solutions; rejecting non-viable answers.",
                "Not all mathematical solutions make sense in real life. Must check against context and reject "
                "impossible values (negative people, fractional members).",
                "After solving the harambee equation, found x = 15 or x = −12. Rejected x = −12.",
                "DQB Updated: Added 'How do we interpret and validate solutions?'\n"
                "Model Revised: Solved harambee equation — original number = 15 members.",
            ),
            (
                "Lesson 7: Reflection and Application",
                "Reviewing entire process from problem to solution; presenting the harambee solution; "
                "exploring extensions and new applications.",
                "Quadratic equations are powerful tools for modelling real-world situations. "
                "The process — define, model, solve, interpret — applies across many contexts.",
                "Solved the harambee problem: 15 original members, each paying Ksh 800. "
                "After 3 left, 12 members each paid Ksh 1,000. Difference = 200. ✓",
                "DQB Completed: Answered driving question with full solution.\n"
                "Final Model: Verified 12,000/15 = 800, 12,000/12 = 1,000, difference = 200 ✓",
            ),
        ],
        "summary_instructions": (
            "Complete one row after each lesson. Be specific in your reflections. "
            "Column 2: describe the main activity or concept explored. "
            "Column 3: summarize key mathematical ideas in your own words. "
            "Column 4: explain how this lesson helps understand or solve the harambee problem. "
            "Column 5: note progress on your Driving Question Board or mathematical model."
        ),
    },

    "physics": {
        "subject":    "Physics",
        "grade":      "Grade 10",
        "strand":     "Strand 1.0: Mechanics and Thermal Physics",
        "substrand":  "Sub-Strand 1.4: Temperature and Thermal Expansion",
        "lessons":    6,
        "driving_question": "Why do materials change size when temperature changes — and how do engineers use this to keep us safe?",
        "phenomenon": (
            "A tightly sealed metal sufuria lid won't budge — but after running hot water over it for 10 seconds, it opens easily. "
            "Also: a glass bottle full of water placed in a freezer cracks overnight."
        ),
        "substrand_overview_rows": [
            ("Subject",           "Physics"),
            ("Grade",             "Grade 10"),
            ("Strand",            "Strand 1.0: Mechanics and Thermal Physics"),
            ("Sub-Strand",        "Sub-Strand 1.4: Temperature and Thermal Expansion"),
            ("Number of Lessons", "6"),
            ("Number of Periods", "12"),
            ("Anchoring Phenomena",
             "1. A tightly sealed metal sufuria lid won't budge — but after running hot water over it for 10 seconds, it opens easily.\n"
             "2. A glass bottle full of water placed in a freezer cracks overnight."),
            ("Driving Question",
             "Why do materials change size when temperature changes — and how do engineers use this to keep us safe?"),
            ("Sub-Strand SLOs: Knowledge",
             "Define temperature and thermal energy; describe temperature scales and conversion; "
             "explain thermal expansion in solids, liquids, and gases; describe anomalous expansion of water; "
             "explain engineering applications of thermal expansion."),
            ("Sub-Strand SLOs: Skills",
             "Measure temperature using different thermometers; calculate linear expansion using ΔL = α × L₀ × ΔT; "
             "apply thermal expansion principles to explain real phenomena; "
             "analyze engineering solutions involving thermal expansion."),
            ("Sub-Strand SLOs: Attitudes & Values",
             "Appreciate the connection between particle-level behavior and everyday phenomena; "
             "recognize how physics knowledge enables engineers to design safer structures."),
            ("Assessment",        "Formative: observation, particle diagrams, calculations, summary table.\nSummative: Final Explanation document."),
            ("Core Values",       "Curiosity, safety awareness, connection to Kenyan context"),
            ("Resources",         "Thermometers, metal ball-and-ring, bimetallic strip, materials samples"),
            ("Notes",             ""),
        ],
        "final_explanation_sections": [
            (
                "SECTION 1: WHAT IS TEMPERATURE AND HOW IS IT MEASURED?",
                "What is temperature at the particle level (how does it relate to kinetic energy)?\n"
                "What are the Celsius and Kelvin scales and how do we convert between them (K = °C + 273)?\n"
                "What types of thermometers exist and how does each one work?\n"
                "Which thermometer would be best for measuring the hot water used on the sufuria lid? Why?",
                "Temperature is a measure of the average kinetic energy of particles in a substance. "
                "When we say something is 'hot', we mean its particles are moving faster on average. "
                "Temperature does NOT measure total heat energy — a large cold object can contain more heat energy than a small hot one.\n\n"
                "Temperature scales:\n"
                "• Celsius (°C): water freezes at 0°C, boils at 100°C. Used in everyday life.\n"
                "• Kelvin (K): absolute scale; 0 K = absolute zero (particles have minimum kinetic energy). K = °C + 273.\n"
                "Example: 25°C = 25 + 273 = 298 K. Human body temperature 37°C = 310 K.\n\n"
                "Types of thermometers:\n"
                "• Liquid-in-glass (mercury or alcohol): liquid expands with temperature along a calibrated scale — simple, accurate for 0-100°C.\n"
                "• Digital/thermocouple: measures voltage change due to temperature — fast, wide range, used in laboratories.\n"
                "• Clinical thermometer: mercury, range 35-42°C, has a constriction to hold reading.\n"
                "• Infrared thermometer: measures heat radiation emitted by objects — non-contact, fast.\n\n"
                "Best for sufuria lid experiment: liquid-in-glass or digital thermometer — both can measure 70-100°C water temperature accurately. "
                "Infrared would also work for measuring lid surface temperature without contact.",
            ),
            (
                "SECTION 2: WHAT HAPPENS TO PARTICLES WHEN TEMPERATURE CHANGES?",
                "How do particles behave differently in solids, liquids, and gases?\n"
                "What happens to particle motion when temperature increases?\n"
                "How does increased particle motion lead to expansion?\n"
                "Do all states of matter expand equally? Why or why not?\n"
                "Connect to phenomenon: What is happening to the particles in the metal lid and glass bottle?",
                "Particle behaviour by state:\n"
                "• Solids: particles closely packed, vibrate in fixed positions, strong intermolecular forces.\n"
                "• Liquids: particles close but free to move past each other, moderate forces.\n"
                "• Gases: particles far apart, move rapidly and randomly, negligible forces.\n\n"
                "When temperature increases, particles absorb energy → kinetic energy increases → particles move faster "
                "and vibrate more vigorously. In solids, faster vibrations mean particles need MORE SPACE between them, "
                "so the material expands. Each bond between atoms is like a spring — at higher temperatures the average "
                "separation increases even though the spring oscillates.\n\n"
                "Expansion rates:\n"
                "• Gases expand most (particles far apart, free to move, large space between them).\n"
                "• Liquids expand moderately (particles closer, moderate forces).\n"
                "• Solids expand least (particles tightly held — but they DO expand).\n"
                "Gases expand ~1000× more than solids for the same temperature rise.\n\n"
                "Connection to phenomena:\n"
                "Sufuria lid: Metal is a solid. When hot water runs over the metal lid, particles gain energy, "
                "vibrate more, lid expands slightly — breaking the airtight seal and making it easier to twist off.\n"
                "Glass bottle: Water is a liquid that becomes solid (ice) when frozen. BUT water is anomalous — "
                "it expands rather than contracts when freezing (investigated in Section 4).",
            ),
            (
                "SECTION 3: HOW DO WE CALCULATE THERMAL EXPANSION?",
                "What is linear expansivity (α) and what are its units?\n"
                "Apply the formula ΔL = α × L₀ × ΔT with a fully worked example (show all 5 steps: "
                "identify given values, identify unknown, write formula, substitute with units, calculate).\n"
                "Choose ONE: (A) A 50 m steel railway rail at 20°C rises to 45°C (α = 12 × 10⁻⁶ °C⁻¹). "
                "OR (B) An aluminum sufuria lid, diameter 20 cm, heats from 25°C to 95°C (α = 24 × 10⁻⁶ °C⁻¹).\n"
                "Connect your calculation to the sufuria lid phenomenon.",
                "Linear expansivity (α) is a material constant that measures how much a 1-metre length of a material "
                "expands per degree Celsius of temperature rise. Units: °C⁻¹ (or K⁻¹).\n"
                "Formula: ΔL = α × L₀ × ΔT\n"
                "where ΔL = change in length, L₀ = original length, ΔT = temperature change.\n\n"
                "Worked example — Option B (Sufuria Lid):\n"
                "Step 1 — Given values: L₀ = 20 cm = 0.20 m; T₁ = 25°C; T₂ = 95°C; α = 24 × 10⁻⁶ °C⁻¹\n"
                "Step 2 — Unknown: ΔL = change in diameter\n"
                "Step 3 — Formula: ΔL = α × L₀ × ΔT\n"
                "Step 4 — Substitute: ΔL = (24 × 10⁻⁶ °C⁻¹) × (0.20 m) × (95 - 25)°C\n"
                "Step 5 — Calculate: ΔL = 24 × 10⁻⁶ × 0.20 × 70 = 24 × 10⁻⁶ × 14 = 336 × 10⁻⁶ m = 0.336 mm\n\n"
                "The lid diameter increases by 0.336 mm when heated with hot water. "
                "This small but significant expansion breaks the partial vacuum seal created when the pot cooled, "
                "allowing the lid to be twisted off easily. Without heating, the lid fits too tightly to turn.",
            ),
            (
                "SECTION 4: WHY DOES WATER BEHAVE UNUSUALLY WHEN IT FREEZES?",
                "How does water behave differently from most substances when it freezes (anomalous expansion)?\n"
                "At what temperature is water most dense? What happens between 4°C and 0°C?\n"
                "Explain using hydrogen bonding and the hexagonal ice crystal structure why water expands by ~9% on freezing.\n"
                "Explain step-by-step why the glass bottle cracked: what happened from room temperature to fully frozen?",
                "Most substances contract when cooling and freeze into a denser solid. "
                "WATER IS ANOMALOUS: it is densest at 4°C and EXPANDS when cooling from 4°C to 0°C and further when freezing.\n\n"
                "Density-temperature relationship for water:\n"
                "• Above 4°C: density decreases as temperature rises (normal for liquids).\n"
                "• From 4°C to 0°C: density DECREASES even as temperature drops (anomalous).\n"
                "• At 0°C: water freezes and expands by approximately 9%, giving ice a density of ~917 kg/m³ vs water's 1000 kg/m³. "
                "This is why ice floats on water!\n\n"
                "Molecular explanation: In liquid water, hydrogen bonds between H₂O molecules are constantly forming and breaking. "
                "As water cools below 4°C, molecules slow down and hydrogen bonds become permanent, arranging molecules into "
                "a hexagonal lattice (like a honeycomb). This open hexagonal structure has MORE space between molecules than liquid water — "
                "hence ice is LESS dense and takes up MORE volume (about 9% more) than liquid water.\n\n"
                "Why the glass bottle cracked (step-by-step):\n"
                "1. Room temperature (~25°C): bottle completely full of liquid water.\n"
                "2. Temperature drops: water cools, contracts slightly down to 4°C.\n"
                "3. Temperature drops further (4°C → 0°C): water EXPANDS anomalously, increasing pressure inside bottle.\n"
                "4. At 0°C: water begins freezing, expanding ~9% — volume increase pushes outward with enormous force.\n"
                "5. Glass cannot withstand the pressure: bottle cracks along weakest line.\n"
                "Prevention: never fill glass bottles completely before freezing — leave expansion space.",
            ),
            (
                "SECTION 5: HOW DO ENGINEERS USE THERMAL EXPANSION TO KEEP US SAFE?",
                "What problems does thermal expansion cause in structures (railway tracks, bridges, roads, pipelines, power lines)?\n"
                "Explain with diagrams how expansion joints work and where they are used.\n"
                "Explain how bimetallic strips work and name three applications (thermostat, fire alarm, circuit breaker).\n"
                "Give at least three specific Kenya examples (e.g. Thika Superhighway, SGR railway, mabati roofing, JKIA runway).\n"
                "Connect to phenomenon: How does this knowledge prevent problems like the cracked bottle?",
                "Problems caused by thermal expansion:\n"
                "• Railway tracks: rails expand on hot days → buckle and derail trains if no gaps allowed.\n"
                "• Bridges: bridge decks expand in heat and contract in cold → would crack and break supports without joints.\n"
                "• Roads: tarmac cracks and buckles in extreme heat without expansion slots.\n"
                "• Pipelines: hot oil/water pipelines must accommodate expansion or they burst.\n"
                "• Power lines: sag in heat, tighten and snap in cold if no slack provided.\n\n"
                "Expansion joints: deliberately designed gaps between sections of track, bridge, road, or building that allow "
                "the material to expand into the gap without buckling or cracking. They look like interlocking 'teeth' on bridges "
                "and roads — the teeth mesh at normal temperatures and separate slightly in heat.\n\n"
                "Bimetallic strips: two metals (e.g., brass and steel) bonded together. When heated, "
                "brass expands more than steel → the strip bends towards the steel side. "
                "Three applications:\n"
                "1. Thermostat: bimetallic strip bends at set temperature → breaks electrical circuit → turns off heater.\n"
                "2. Fire alarm: strip bends in heat → closes electrical circuit → triggers alarm.\n"
                "3. Circuit breaker: strip bends due to overheating current → cuts the circuit to prevent fire.\n\n"
                "Kenya examples:\n"
                "• SGR railway (Mombasa-Nairobi): expansion gaps designed into rails to prevent buckling in coastal/inland heat variation.\n"
                "• Thika Superhighway bridges: large expansion joints visible on all bridge decks.\n"
                "• Mabati (iron sheet) roofing: sheets contract at night (popping sound) and expand in daytime — gaps between sheets prevent buckling.\n"
                "• JKIA runway: concrete slabs have expansion joints to prevent cracking under aircraft loading and temperature change.\n\n"
                "Connection to cracked bottle: knowing that water expands on freezing allows engineers and consumers to design packaging "
                "that accommodates this expansion — e.g., PET plastic water bottles (flexible), or headspace left in glass bottles. "
                "The same principle used to prevent bridge cracking prevents bottle cracking: allow for the dimensional change!",
            ),
        ],
        "final_explanation_rubric": {
            "headers": ["Criterion", "Beginning (1)", "Developing (2)", "Proficient (3)", "Advanced (4)"],
            "rows": [
                ("1. Phenomenon Connection",
                 "Rarely refers to anchoring phenomena; explanation disconnected from observation",
                 "Mentions phenomena but connections superficial",
                 "Regularly connects concepts to phenomena; explanation grounded in observation",
                 "Seamlessly integrates phenomena throughout; makes additional real-world connections"),
                ("2. Scientific Accuracy & Vocabulary",
                 "Multiple scientific errors; misuses key terms",
                 "Some errors; mostly accurate; some terms used incorrectly",
                 "Scientifically accurate; correct use of terminology throughout",
                 "Highly accurate and precise; sophisticated scientific language; nuanced understanding"),
                ("3. Particle-Level Reasoning",
                 "Little or no particle-level explanation; stays at macroscopic level",
                 "Attempts particle reasoning but gaps or errors; unclear link to macroscopic",
                 "Clearly explains particle behavior and links to macroscopic observations",
                 "Sophisticated particle reasoning; explicit microscopic-to-macroscopic connections"),
                ("4. Mathematical Application",
                 "Calculation missing, incomplete, or major errors; no method shown",
                 "Calculation present but errors in setup, substitution, or arithmetic",
                 "Calculation correct; all steps shown; appropriate units and significant figures",
                 "Calculation exemplary; method crystal clear; result interpreted in context"),
                ("5. Engineering Applications & Completeness",
                 "Fewer than 3 sections; few or no engineering examples",
                 "3–4 sections; some engineering mentions but explanations vague",
                 "All 5 sections; multiple engineering applications with Kenya examples",
                 "All 5 sections comprehensively; rich Kenya examples; design reasoning discussed"),
            ],
        },
        "summary_table_headers": [
            "Lesson # and Title",
            "What did I observe or investigate?",
            "What did I learn?",
            "How does this explain the sufuria lid and frozen bottle phenomenon?",
            "DQB and Model Update",
        ],
        "summary_table_rows": [
            (
                "Lesson 1: Introduction to Temperature and Heat",
                "Measured temperature of different objects using thermometers. "
                "Discussed what 'hot' and 'cold' mean and shared everyday experiences.",
                "Temperature measures the average kinetic energy of particles. "
                "Heat is energy transferred from hotter to cooler objects. They are related but different.",
                "This sets the foundation: when we run hot water over the sufuria lid, heat transfers to the lid "
                "and increases its temperature. But we don't yet know why this makes it easier to open.",
                "DQB Started: Why do things expand when heated? What happens to particles when temperature increases?\n"
                "Initial Model: Particles in a solid vibrate but stay in place.",
            ),
            (
                "Lesson 2: Particle Model of Matter and Temperature",
                "Built particle diagrams for solids, liquids, and gases. Investigated how particle motion "
                "changes with temperature using demonstrations. Explored kinetic theory of matter.",
                "Temperature measures the average kinetic energy of particles. When temperature rises, particles "
                "move faster and need more space — this causes expansion. Solids vibrate in fixed positions; "
                "liquids flow; gases move freely. Higher temperature = faster particle motion.",
                "When hot water runs over the sufuria lid, energy transfers to the metal particles. "
                "They vibrate more vigorously and the metal expands — this is why heating helps open the lid. "
                "The particle model explains thermal expansion at the microscopic level.",
                "DQB Updated: Answered 'What happens to particles when something heats up?'\n"
                "Model Revised: Particle diagrams added showing vibrating solids and moving liquids/gases with temperature arrows.",
            ),
            (
                "Lesson 3: Thermal Expansion in Solids",
                "Heated a metal ball until it no longer fit through a ring. Observed a bimetallic strip bend when heated. "
                "Measured change in length of a metal rod with a dial gauge.",
                "Thermal expansion: most solids expand when heated because particles vibrate more vigorously and need more space. "
                "Different metals have different coefficients of linear expansion (α). "
                "Engineers use expansion joints to accommodate this.",
                "Sufuria lid: the metal lid expands when heated with hot water, breaking the tight seal and making it easier to twist off. "
                "Frozen bottle: we know solids usually expand when heated — water may behave differently (anomaly to investigate).",
                "DQB Updated: Answered 'Why does the lid open when heated?' (differential expansion breaks the seal.)\n"
                "New question: Does water expand or contract when it freezes?\n"
                "Model Revised: Added arrows showing particles moving further apart when heated.",
            ),
            (
                "Lesson 4: Thermal Expansion in Liquids",
                "Observed coloured water and alcohol rising in thin glass tubes when heated. Compared expansion "
                "rates of different liquids. Investigated how liquid thermometers work using expansion principles.",
                "Liquids expand more than solids for the same temperature rise because their particles are "
                "free to move further apart. This is why liquid-in-glass thermometers work — the liquid expands "
                "up the calibrated tube. Different liquids have different expansion rates.",
                "Liquids also expand when heated. The water in the glass bottle was a liquid when placed in the "
                "freezer. As it cooled, it first contracted slightly (normal liquid behaviour) before eventually "
                "freezing. The expansion happened when it froze — this is the anomaly we need to investigate.",
                "DQB Updated: Added 'Do liquids expand the same as solids?' Answered — liquids expand more.\n"
                "Model Revised: Added liquid expansion diagram showing particles spreading out more freely than solids.",
            ),
            (
                "Lesson 5: Anomalous Expansion of Water",
                "Investigated the density of water at different temperatures. Observed ice floating on water. "
                "Studied hydrogen bonding and hexagonal ice crystal structure models.",
                "Water is anomalous: it is densest at 4°C and EXPANDS when cooling from 4°C to 0°C. "
                "When water freezes, hydrogen bonds form a hexagonal lattice with MORE space than liquid water — "
                "so ice is less dense than water and expands by ~9%.",
                "MYSTERY SOLVED for the frozen bottle: water expands ~9% when it freezes due to hydrogen bond "
                "hexagonal lattice formation. A bottle completely full of water had no room for this expansion — "
                "the outward force cracked the glass. This is anomalous — most substances contract when freezing!",
                "DQB Updated: Answered 'Why did the glass bottle crack?' Water expands (anomalously) on freezing.\n"
                "Model Revised: Hydrogen bond hexagonal ice structure drawn. Water anomaly explained at molecular level.",
            ),
            (
                "Lesson 6: Engineering Applications and Design Challenge",
                "Investigated expansion joints on bridges and railway tracks. Studied bimetallic strip bending. "
                "Designed a solution to a thermal expansion engineering problem (e.g. railway track joint).",
                "Engineers must account for thermal expansion in all structures. Expansion joints allow materials "
                "to expand and contract safely. Bimetallic strips (two metals with different expansion rates) "
                "bend when heated — used in thermostats, fire alarms, and circuit breakers.",
                "The cracked bottle represents a failure to account for thermal expansion (water expanding on freezing). "
                "Engineers prevent similar failures: expansion joints in SGR railway tracks, flexible PET water bottles, "
                "gap-fitted mabati roofing, bridge deck joints on the Thika Superhighway.",
                "DQB Completed: Answered driving question — materials change size because particles move more at higher temperatures; "
                "engineers use expansion joints, bimetallic strips, and flexible materials to keep structures safe.\n"
                "Final Model: Complete model explaining thermal expansion in solids and liquids, "
                "including water's anomalous behavior, with engineering applications.",
            ),
        ],
        "summary_instructions": (
            "1. Complete one row after each lesson (or as directed by your teacher).\n"
            "2. Column 1: Record the lesson number and title.\n"
            "3. Column 2: Describe what you observed, measured, or investigated (activities, demonstrations, simulations, or data).\n"
            "4. Column 3: Summarize the key scientific concepts you learned. Use precise vocabulary "
            "(e.g. thermal expansion, kinetic energy, linear expansivity, coefficient of expansion).\n"
            "5. Column 4: Connect your learning to the anchoring phenomena — explain how this lesson helps you understand "
            "the sufuria lid or frozen bottle (or both).\n"
            "6. Column 5: Update your Driving Question Board (note new questions raised or questions answered) "
            "and update your scientific model sketch."
        ),
    },
}

SOURCE_COMPLETE = {
    "chemistry":  SRC_DIR / "Chemistry_Chemical_Bonding_CBE_LessonSequence_COMPLETE.docx",
    "mathematics": SRC_DIR / "Mathematics_Quadratic_Equations_CBE_LessonSequence_COMPLETE.docx",
    "physics":    SRC_DIR / "Physics_Temperature_Thermal_Expansion_CBE_LessonSequence_COMPLETE.docx",
}

SOURCE_FINAL_EXP = {
    "chemistry":  SRC_DIR / "FINAL_EXPLANATION_Chemical_Bonding.docx",
    "mathematics": SRC_DIR / "FINAL_EXPLANATION_Quadratic_Equations.docx",
    "physics":    SRC_DIR / "FINAL_EXPLANATION_Temperature_Thermal_Expansion.docx",
}

SOURCE_SUMMARY_TABLE = {
    "chemistry":  SRC_DIR / "SUMMARY_TABLE_Chemical_Bonding.docx",
    "mathematics": SRC_DIR / "SUMMARY_TABLE_Quadratic_Equations.docx",
    "physics":    SRC_DIR / "SUMMARY_TABLE_Temperature_Thermal_Expansion.docx",
}

OUTPUT_FILES = {
    "chemistry": {
        "lesson_seq":    OUT_DIR / "Chemistry_10_SubStrand1.4_ChemicalBonding_CBE_LessonSequence.docx",
        "final_exp":     OUT_DIR / "Chemistry_10_SubStrand1.4_ChemicalBonding_FinalExplanation.docx",
        "summary_table": OUT_DIR / "Chemistry_10_SubStrand1.4_ChemicalBonding_SummaryTable.docx",
    },
    "mathematics": {
        "lesson_seq":    OUT_DIR / "Mathematics_10_SubStrand1.3_QuadraticEquations_CBE_LessonSequence.docx",
        "final_exp":     OUT_DIR / "Mathematics_10_SubStrand1.3_QuadraticEquations_FinalExplanation.docx",
        "summary_table": OUT_DIR / "Mathematics_10_SubStrand1.3_QuadraticEquations_SummaryTable.docx",
    },
    "physics": {
        "lesson_seq":    OUT_DIR / "Physics_10_SubStrand1.4_ThermalExpansion_CBE_LessonSequence.docx",
        "final_exp":     OUT_DIR / "Physics_10_SubStrand1.4_ThermalExpansion_FinalExplanation.docx",
        "summary_table": OUT_DIR / "Physics_10_SubStrand1.4_ThermalExpansion_SummaryTable.docx",
    },
}


# ═══════════════════════════════════════════════════════════════════════════════
# LOW-LEVEL DOCX HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def _shade(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd if any
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_para(cell, text: str, bold=False, size_pt=10, color_hex=None,
               align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    _apply_run(p, text, bold=bold, size_pt=size_pt, color_hex=color_hex, italic=italic)
    return p


def _add_para(cell, text: str, bold=False, size_pt=10, color_hex=None,
              align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    _apply_run(p, text, bold=bold, size_pt=size_pt, color_hex=color_hex)
    return p


def _apply_run(para, text: str, bold=False, size_pt=10, color_hex=None, italic=False):
    if not text:
        return
    run = para.add_run(text)
    run.font.name  = FONT
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    if color_hex:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16),
        )
    else:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _white_run(para, text, bold=False, size_pt=10):
    _apply_run(para, text, bold=bold, size_pt=size_pt, color_hex=C_WHITE)


def _merge_row(table, row_idx):
    """Merge all cells in a row."""
    row = table.rows[row_idx]
    if len(row.cells) < 2:
        return row.cells[0]
    a = row.cells[0]
    for c in row.cells[1:]:
        a.merge(c)
    return a


def _col_widths(table, widths_inches):
    """Set column widths. widths_inches is a list of floats."""
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(0, tblGrid)
    # Clear existing gridCols
    for gc in tblGrid.findall(qn("w:gridCol")):
        tblGrid.remove(gc)
    for w in widths_inches:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 1440)))
        tblGrid.append(gc)
    # Set each cell width
    for row in table.rows:
        cells = row.cells
        for i, cell in enumerate(cells):
            if i < len(widths_inches):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.append(tcW)
                tcW.set(qn("w:w"), str(int(widths_inches[i] * 1440)))
                tcW.set(qn("w:type"), "dxa")


def _set_page_landscape(doc):
    """Set US Letter Landscape, 0.75\" margins."""
    sec = doc.sections[0]
    sec.page_width    = Inches(11.0)
    sec.page_height   = Inches(8.5)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)


def _doc_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    _apply_run(p, text, bold=True, size_pt=14, color_hex=C_NAVY)


def _doc_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    _apply_run(p, text, bold=False, size_pt=11, color_hex=C_TEAL)


def _tbl_no_spacing(doc):
    """Add a 0-height spacer paragraph to keep tables adjacent (no extra space)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"),  "0")
    spacing.set(qn("w:line"),   "240")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _add_page_break(doc):
    """Insert a page break before the next lesson."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def parse_doc_sections(doc_path: Path, section_prefixes: list) -> dict:
    """Parse content under H2 headings that match any of section_prefixes (prefix match).

    Returns {prefix: [(kind, text), ...]} where kind is 'h3', 'bullet', 'quote', or 'text'.
    'quote' is for block-quote style paragraphs (example student responses).
    """
    doc = Document(doc_path)
    results = {}
    current_key = None
    buffer = []

    def flush():
        if current_key is not None and buffer:
            results[current_key] = buffer[:]

    for p in doc.paragraphs:
        text = p.text.strip()
        style = p.style.name if p.style else ""
        style_l = style.lower()
        if not text:
            continue

        if "heading 2" in style_l:
            flush()
            buffer = []
            current_key = None
            for prefix in section_prefixes:
                if text.upper().startswith(prefix.upper()):
                    current_key = prefix
                    break
        elif current_key is not None:
            if "heading 3" in style_l:
                buffer.append(("h3", text))
            elif any(s in style_l for s in ("compact", "list bullet", "list paragraph")):
                buffer.append(("bullet", text))
            elif "block" in style_l:
                buffer.append(("quote", text))
            else:
                buffer.append(("text", text))

    flush()
    return results


def _cell_para_lines(cell, content: list, size_pt: int = 9):
    """Write rich content into a cell.

    content is a list of (kind, text) tuples where kind is one of:
    'h3'    -> bold subheading
    'bullet' -> bulleted line with left indent
    'quote'  -> italic, indented (used for example student responses)
    'text'   -> normal paragraph
    """
    first_para = True
    for kind, text in content:
        if first_para:
            p = cell.paragraphs[0]
            first_para = False
        else:
            p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)

        if kind == "h3":
            _apply_run(p, text, bold=True, size_pt=size_pt)
        elif kind == "bullet":
            p.paragraph_format.left_indent = Pt(12)
            _apply_run(p, f"• {text}", size_pt=size_pt)
        elif kind == "quote":
            p.paragraph_format.left_indent = Pt(18)
            _apply_run(p, text, size_pt=size_pt, italic=True)
        else:
            _apply_run(p, text, size_pt=size_pt)

    if not content:
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)


def _text_to_rich(text: str) -> list:
    """Convert a plain-text string to [(kind, str), ...] for _cell_para_lines().

    Detects and marks:
    • Lettered lists   a) b) c) or A) B) C)                              → 'bullet'
    • Numbered lists   1. 2. 3. or 1) 2) 3)                             → 'bullet'
    • Dash prefix      - text (possibly concatenated list "- a - b - c") → 'bullet' each
    • Label: - items   "Physical Setup: - item1 - item2"                 → h3 + bullets
    • Short label ending with colon  (≤ 70 chars)                        → 'h3'
    • Numbered section heading  8.1 Title                                → 'h3'
    • "Short Label: description" (label ≤ 40 chars, title-case start)    → h3 + text
    • Question lines ending with ?  (≤ 180 chars)                        → 'bullet'
    • Everything else                                                     → 'text'
    """
    import re

    _LABEL_EXCLUSIONS = ("e.g.", "note:", "for example", "i.e.", "e.g.:", "for instance",
                         "note that", "please note", "nb:", "ps:")

    def _split_dash_list(content: str) -> list:
        """Split a string that may be a concatenated dash list into individual items."""
        if content.count(" - ") >= 2:
            return [x.strip() for x in re.split(r"\s+-\s+", content) if x.strip()]
        return [content]

    lines = [ln.strip() for ln in text.replace("\r\n", "\n").split("\n")]
    result = []
    for ln in lines:
        if not ln:
            continue
        ln_lower = ln.lower()
        # Lettered list: a) b) or A) B)
        if re.match(r"^[a-zA-Z]\)\s", ln):
            result.append(("bullet", ln))
        # Numbered list: 1. 2. or 1) 2)
        elif re.match(r"^\d+[.)]\s", ln):
            result.append(("bullet", ln))
        # Dash/bullet prefix — may be a concatenated list "- item1 - item2 - item3"
        elif ln.startswith(("- ", "• ", "* ", "– ", "— ")):
            content = re.sub(r"^[-•*–—]\s+", "", ln)
            for item in _split_dash_list(content):
                result.append(("bullet", item))
        # "Label: - item1 - item2" → h3 label + bullets
        elif ": - " in ln:
            label, rest = ln.split(": - ", 1)
            if len(label) <= 60:
                result.append(("h3", label + ":"))
                for item in _split_dash_list(rest):
                    result.append(("bullet", item))
            else:
                result.append(("text", ln))
        # Short label ending with colon (standalone heading, ≤ 70 chars)
        elif (ln.endswith(":") and len(ln) <= 70
              and not any(ln_lower.startswith(x) for x in
                          ("e.g.", "note:", "(", "for example", "i.e."))):
            result.append(("h3", ln))
        # Numbered section headings like "8.1 Knowledge"
        elif re.match(r"^\d+\.\d*\s+\w", ln) and len(ln) <= 80:
            result.append(("h3", ln))
        # "Short Label: description text" — bold the label, normal text for rest
        # Matches lines starting with capital letter, label ≤ 40 chars before ": "
        elif (re.match(r"^[A-Z][^:\n]{2,38}: \S", ln)
              and not any(ln_lower.startswith(x) for x in _LABEL_EXCLUSIONS)):
            colon_pos = ln.index(": ")
            label = ln[:colon_pos]
            rest  = ln[colon_pos + 2:]
            result.append(("h3", label + ":"))
            result.append(("text", rest))
        # Stand-alone question lines → bullet
        elif ln.endswith("?") and 10 <= len(ln) <= 180:
            result.append(("bullet", ln))
        else:
            result.append(("text", ln))
    return result if result else [("text", text.strip())]


def _cell_para_auto(cell, text: str, size_pt: int = 9):
    """Write text into a cell with automatic bullet/heading detection via _text_to_rich()."""
    if not text or not text.strip():
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        return
    _cell_para_lines(cell, _text_to_rich(text), size_pt=size_pt)


def _enrich_content(content: list) -> list:
    """Post-process parsed content to improve bullet/heading detection.

    - 'text' items are re-evaluated through _text_to_rich() to catch embedded
      dash-lists, label:description patterns, question marks, etc.
    - 'bullet' items ending with ':' (≤ 60 chars) are promoted to 'h3' headings.
    """
    result = []
    for kind, text in content:
        if kind == "text":
            for new_kind, new_text in _text_to_rich(text):
                result.append((new_kind, new_text))
        elif kind == "bullet" and text.strip().endswith(":") and len(text.strip()) <= 60:
            result.append(("h3", text.strip()))
        else:
            result.append((kind, text))
    return result


def _build_section_table(doc, header_text: str, content: list,
                          header_fill: str = None, content_fill: str = None,
                          size_pt: int = 9):
    """Build a 2-row x 1-col section table with a coloured header and rich content."""
    if header_fill is None:
        header_fill = C_TEAL
    if content_fill is None:
        content_fill = C_WHITE
    t = _new_table(doc, 2, 1)
    _col_widths(t, [9.5])
    _shade(t.rows[0].cells[0], header_fill)
    _cell_para(t.rows[0].cells[0], header_text, bold=True, size_pt=11, color_hex=C_WHITE)
    _shade(t.rows[1].cells[0], content_fill)
    _cell_para_lines(t.rows[1].cells[0], _enrich_content(content), size_pt=size_pt)


def _set_tbl_border(table):
    """Apply thin black borders to all cells."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    # Remove existing tblBorders
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(tblBorders)


def _new_table(doc, rows, cols):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_tbl_border(tbl)
    return tbl


# ═══════════════════════════════════════════════════════════════════════════════
# SOURCE DOCUMENT PARSING  (unchanged from working version)
# ═══════════════════════════════════════════════════════════════════════════════

def iter_block_items(doc):
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            yield "paragraph", DocxParagraph(child, doc)
        elif tag == "tbl":
            yield "table", DocxTable(child, doc)


def is_impl_table(table) -> bool:
    if len(table.columns) != 5:
        return False
    header = table.rows[0].cells[0].text.strip()
    return "Learner" in header or "Experience" in header


def _cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs if p.text.strip()).strip()


def parse_source(doc_path: Path) -> list[dict]:
    doc = Document(doc_path)
    lessons = []
    current: Optional[dict] = None
    section = None
    impl_tables_for_current = []
    buffer: list[str] = []

    LESSON_HEADING_RE = re.compile(r"^LESSON\s+(\d+)\s*[:–—]\s*(.*)", re.IGNORECASE)
    SECTION_RE        = re.compile(r"^([A-E])\.\s+(.*)")

    def flush_buffer():
        return "\n".join(buffer).strip()

    def save_section():
        nonlocal buffer, section
        if current is None or section is None:
            buffer = []
            return
        text = flush_buffer()
        if section == "A":
            parts = re.split(r"\n(?=Knowledge:|Skills:|Attitudes?:)", text, flags=re.IGNORECASE)
            for part in parts:
                part = part.strip()
                if re.match(r"Knowledge:", part, re.I):
                    current["slo_knowledge"] += "\n" + part[10:].strip()
                elif re.match(r"Skills:", part, re.I):
                    current["slo_skills"] += "\n" + part[7:].strip()
                elif re.match(r"Attitudes?:", part, re.I):
                    current["slo_attitudes"] += "\n" + re.split(r"Attitudes?:", part, flags=re.I, maxsplit=1)[-1].strip()
                elif re.match(r"By the end", part, re.I):
                    if not current["slo_purpose"]:
                        current["slo_purpose"] = part
                elif part:
                    current["slo_knowledge"] += "\n" + part
        elif section == "B":
            lines = text.split("\n")
            for i, ln in enumerate(lines):
                if re.match(r"Key Inquiry Question", ln, re.I):
                    inline = ln.split(":", 1)[1].strip() if ":" in ln else ""
                    if inline:
                        current["inquiry_question"] = inline
                    else:
                        next_ln = lines[i+1].strip() if i+1 < len(lines) else ""
                        if next_ln and not re.match(r"(Purpose|Safety|Materials|Duration|Supporting)", next_ln, re.I):
                            current["inquiry_question"] = next_ln
                elif re.match(r"Purpose in Storyline|Purpose:", ln, re.I):
                    purpose_lines = []
                    for j in range(i+1, len(lines)):
                        if re.match(r"(Safety|Materials|Duration|Supporting)", lines[j], re.I):
                            break
                        if lines[j].strip():
                            purpose_lines.append(lines[j].strip())
                    current["overview_purpose"] = " ".join(purpose_lines)
                elif re.match(r"Safety", ln, re.I):
                    # Capture inline content (e.g. "Safety: - item1 - item2") + any continuation lines
                    inline = ln.split(":", 1)[1].strip() if ":" in ln else ""
                    safety_lines = [inline] if inline else []
                    for j in range(i+1, len(lines)):
                        if re.match(r"(Materials|Purpose|Key Inquiry)", lines[j], re.I):
                            break
                        if lines[j].strip():
                            safety_lines.append(lines[j].strip())
                    current["safety"] = " ".join(safety_lines)
                elif re.match(r"Materials?", ln, re.I) and ":" in ln:
                    # Capture inline content (e.g. "Materials: - item1 - item2") + continuation lines
                    inline = ln.split(":", 1)[1].strip() if ":" in ln else ""
                    mat_lines = [inline] if inline else []
                    for j in range(i+1, len(lines)):
                        if re.match(r"(Safety|Purpose|Duration|Key Inquiry)", lines[j], re.I):
                            break
                        if lines[j].strip():
                            mat_lines.append(lines[j].strip())
                    current["materials"] = " ".join(mat_lines)
        elif section == "D":
            lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
            current["reflections"] = lines
        elif section == "E":
            if text:
                current["summary_prompt"] = text
        buffer = []

    def save_impl_tables():
        if current is None:
            return
        for tbl in impl_tables_for_current:
            rows = []
            for row in tbl.rows[1:]:
                rows.append(tuple(_cell_text(c) for c in row.cells))
            if not current["period1_table"]:
                current["period1_table"] = rows
            else:
                current["period2_table"] = rows

    def new_lesson(number, title):
        nonlocal current, section, impl_tables_for_current
        if current is not None:
            save_section()
            save_impl_tables()
            lessons.append(current)
        impl_tables_for_current = []
        current = {
            "number": int(number),
            "title":  title.strip(),
            "inquiry_question": "",
            "slo_purpose":      "",
            "slo_knowledge":    "",
            "slo_skills":       "",
            "slo_attitudes":    "",
            "overview_purpose": "",
            "materials":        "",
            "safety":           "",
            "period1_table":    [],
            "period2_table":    [],
            "period1_heading":  "Period 1 (40 minutes)",
            "period2_heading":  "Period 2 (40 minutes)",
            "reflections":      [],
            "summary_prompt":   "",
        }
        section = None

    for kind, obj in iter_block_items(doc):
        if kind == "paragraph":
            text = obj.text.strip()
            if not text:
                continue
            style_l = (obj.style.name if obj.style else "").lower()
            is_heading = "heading" in style_l

            m = LESSON_HEADING_RE.match(text)
            if m and is_heading:
                save_section()
                new_lesson(m.group(1), m.group(2))
                section = None
                buffer = []
                continue

            sm = SECTION_RE.match(text)
            if sm and current is not None and is_heading:
                save_section()
                section = sm.group(1).upper()
                buffer = []
                continue

            if "heading 4" in style_l and current is not None:
                if "PERIOD 1" in text.upper():
                    current["period1_heading"] = text
                elif "PERIOD 2" in text.upper():
                    current["period2_heading"] = text
                continue

            if current is not None and section is None and "Supporting Driving Question" in text:
                q = text.split(":", 1)[-1].strip() if ":" in text else ""
                if q:
                    current["inquiry_question"] = q
                continue

            if current is not None and section is not None:
                buffer.append(text)

        elif kind == "table":
            if is_impl_table(obj) and current is not None:
                save_section()
                impl_tables_for_current.append(obj)
                section = None
                buffer = []

    if current is not None:
        save_section()
        save_impl_tables()
        lessons.append(current)

    return lessons


def parse_section_a(doc_path: Path) -> list[tuple[str, str]]:
    """Parse Section A (Sub-Strand Overview) from source document.

    Iterates all paragraphs, collects content grouped by Heading 3 subsections,
    between 'SECTION A' heading and the first 'SECTION B' or 'LESSON' heading.
    Returns list of (subsection_title, content_text) tuples.
    """
    doc = Document(doc_path)
    in_section_a = False
    current_h3 = None
    section_data: list[tuple[str, str]] = []
    buffer: list[str] = []

    SEPARATOR = set("─━ \t")

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else ""
        style_l = style.lower()

        if ("SECTION A" in text.upper() or "SECTION  A" in text.upper()) and "heading" in style_l:
            in_section_a = True
            continue

        if not in_section_a:
            continue

        if "heading 2" in style_l and ("SECTION B" in text.upper() or "LESSON" in text.upper()):
            break
        if "heading 1" in style_l and "LESSON" in text.upper():
            break

        if "heading 3" in style_l:
            if current_h3 is not None:
                section_data.append((current_h3, "\n".join(buffer).strip()))
            current_h3 = text
            buffer = []
        else:
            if current_h3 is not None:
                if not set(text).issubset(SEPARATOR):
                    buffer.append(text)

    if current_h3 is not None:
        section_data.append((current_h3, "\n".join(buffer).strip()))

    return section_data


# ═══════════════════════════════════════════════════════════════════════════════
# LESSON SEQUENCE BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

def _build_table0_overview(doc, meta, section_a_rows=None):
    """Sub-strand overview table: navy SUB-STRAND OVERVIEW banner + data rows (2c).

    If section_a_rows is provided (parsed from source), uses those instead of
    the hardcoded META substrand_overview_rows so all 13-14 H3 subsections appear.
    Label cells cycle through theme colors matching the other tables.
    """
    rows_data = section_a_rows if section_a_rows else meta["substrand_overview_rows"]
    t = _new_table(doc, 1 + len(rows_data), 2)
    _col_widths(t, [2.083, 7.417])

    # R0: navy merged banner
    c = _merge_row(t, 0)
    _shade(c, C_NAVY)
    _cell_para(c, "SUB-STRAND OVERVIEW", bold=True, size_pt=11, color_hex=C_WHITE)

    # Cycle through theme colors for label cells (matching Table A/B/C palette)
    _OVERVIEW_LABEL_FILLS = [C_LT_BLUE, C_PURPLE_LT, C_TEAL_LT, C_GREEN_LT, C_ORANGE_LT]

    for ri, (label, value) in enumerate(rows_data):
        row_idx = ri + 1
        label_fill = _OVERVIEW_LABEL_FILLS[ri % len(_OVERVIEW_LABEL_FILLS)]
        _shade(t.rows[row_idx].cells[0], label_fill)
        _shade(t.rows[row_idx].cells[1], C_WHITE)
        _cell_para(t.rows[row_idx].cells[0], label, bold=True, size_pt=9)
        if isinstance(value, list):
            _cell_para_lines(t.rows[row_idx].cells[1], value, size_pt=9)
        else:
            _cell_para_auto(t.rows[row_idx].cells[1], value, size_pt=9)


def _build_table_A(doc, lesson):
    """Table A — Specific Learning Outcomes (5r × 2c), total 9.5\".

    R0: navy merged LESSON banner
    R1: teal merged A. SPECIFIC LEARNING OUTCOMES
    R2: lt-blue Knowledge | content
    R3: lt-blue Skills    | content
    R4: lt-blue Attitudes | content
    """
    t = _new_table(doc, 5, 2)
    _col_widths(t, [2.083, 7.417])

    # R0: merged navy — LESSON N: Title
    c = _merge_row(t, 0)
    _shade(c, C_NAVY)
    _cell_para(c, f"LESSON {lesson['number']}: {lesson['title']}",
               bold=True, size_pt=11, color_hex=C_WHITE)

    # R1: merged teal — A. SPECIFIC LEARNING OUTCOMES
    c = _merge_row(t, 1)
    _shade(c, C_TEAL)
    _cell_para(c, "A. SPECIFIC LEARNING OUTCOMES", bold=True, size_pt=11, color_hex=C_WHITE)

    # R2: Knowledge
    _shade(t.rows[2].cells[0], C_LT_BLUE)
    _shade(t.rows[2].cells[1], C_WHITE)
    _cell_para(t.rows[2].cells[0], "Knowledge", bold=True, size_pt=9)
    _cell_para_auto(t.rows[2].cells[1], lesson["slo_knowledge"].strip() or "—")

    # R3: Skills
    _shade(t.rows[3].cells[0], C_LT_BLUE)
    _shade(t.rows[3].cells[1], C_WHITE)
    _cell_para(t.rows[3].cells[0], "Skills", bold=True, size_pt=9)
    _cell_para_auto(t.rows[3].cells[1], lesson["slo_skills"].strip() or "—")

    # R4: Attitudes
    _shade(t.rows[4].cells[0], C_LT_BLUE)
    _shade(t.rows[4].cells[1], C_WHITE)
    _cell_para(t.rows[4].cells[0], "Attitudes", bold=True, size_pt=9)
    _cell_para_auto(t.rows[4].cells[1], lesson["slo_attitudes"].strip() or "—")


def _build_table_B(doc, lesson):
    """Table B — Lesson Overview (5r × 2c), each field in its own row.

    R0: teal merged  B. LESSON OVERVIEW
    R1: purple-lt    Key Inquiry Question | content
    R2: teal-lt      Purpose in Storyline | content
    R3: orange-lt    Safety Considerations | content
    R4: lt-blue      Materials Needed      | content
    """
    t = _new_table(doc, 5, 2)
    _col_widths(t, [2.083, 7.417])

    # R0: merged teal header
    c = _merge_row(t, 0)
    _shade(c, C_TEAL)
    _cell_para(c, "B. LESSON OVERVIEW", bold=True, size_pt=11, color_hex=C_WHITE)

    rows_spec = [
        (1, C_PURPLE_LT, "Key Inquiry Question",  lesson.get("inquiry_question", "") or "—"),
        (2, C_TEAL_LT,   "Purpose in Storyline",  lesson.get("overview_purpose", "").strip() or "—"),
        (3, C_ORANGE_LT, "Safety Considerations",  lesson.get("safety", "") or "None noted."),
        (4, C_LT_BLUE,   "Materials Needed",       lesson.get("materials", "") or "See lesson plan."),
    ]
    for ri, fill, label, value in rows_spec:
        _shade(t.rows[ri].cells[0], fill)
        _shade(t.rows[ri].cells[1], C_WHITE)
        _cell_para(t.rows[ri].cells[0], label, bold=True, size_pt=9)
        _cell_para_auto(t.rows[ri].cells[1], value)


def _build_table_C_period(doc, lesson, period_num: int):
    """Table C — 6-col Implementation Framework for one period (7r × 6c), total 9.5\".

    Called once per period so Period 1 and Period 2 are separate tables.
    Col widths match Bio reference: [0.625, 1.775, 1.775, 1.775, 1.775, 1.775].
    """
    COL_W = [0.625, 1.775, 1.775, 1.775, 1.775, 1.775]
    COL_HEADER_FILLS = [C_NAVY, C_MED_BLUE, C_TEAL, C_MED_BLUE, C_TEAL, C_MED_BLUE]
    COL_HEADERS = ["Phase", "Learner Experience", "Resource",
                   "Teacher Actions", "Sensemaking Strategy", "Assessment Strategy"]

    if period_num == 1:
        period_heading = lesson.get("period1_heading", "Period 1 (40 minutes)")
        period_data = lesson["period1_table"]
    else:
        period_heading = lesson.get("period2_heading", "Period 2 (40 minutes)")
        period_data = lesson["period2_table"]

    t = _new_table(doc, 7, 6)
    _col_widths(t, COL_W)

    # R0: merged title with period heading
    c = _merge_row(t, 0)
    _shade(c, C_TEAL)
    _cell_para(c, f"C. LESSON IMPLEMENTATION FRAMEWORK — {period_heading}",
               bold=True, size_pt=11, color_hex=C_WHITE)

    # R1: col headers
    for cell, hdr, fill in zip(t.rows[1].cells, COL_HEADERS, COL_HEADER_FILLS):
        _shade(cell, fill)
        _cell_para(cell, hdr, bold=True, size_pt=9, color_hex=C_WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # R2–R6: 5 phase rows
    content_fills = [C_WHITE, C_LT_GRAY, C_WHITE, C_LT_GRAY, C_WHITE]
    for ri in range(5):
        row = t.rows[ri + 2]
        phase_fill = PHASE_FILLS[ri]
        phase_name = PHASE_NAMES[ri]

        _shade(row.cells[0], phase_fill)
        _cell_para(row.cells[0], phase_name, bold=True, size_pt=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        src_row = period_data[ri] if ri < len(period_data) else None

        for ci in range(5):
            cell = row.cells[ci + 1]
            _shade(cell, content_fills[ci])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            val = src_row[ci] if src_row and ci < len(src_row) else ""
            _cell_para_auto(cell, val)


def _build_table_D(doc, lesson):
    """Table D — Teacher Reflection (2r × 1c)."""
    t = _new_table(doc, 2, 1)
    _col_widths(t, [9.5])

    _shade(t.rows[0].cells[0], C_ORANGE)
    _cell_para(t.rows[0].cells[0], "D. TEACHER REFLECTION",
               bold=True, size_pt=11, color_hex=C_WHITE)

    _shade(t.rows[1].cells[0], C_ORANGE_LT)
    reflections = lesson["reflections"]
    if reflections:
        refl_text = "\n".join(f"{i + 1}. {item}" for i, item in enumerate(reflections))
    else:
        refl_text = (
            "After teaching this lesson, reflect on:\n"
            "1. Which phase generated the most discussion or engagement?\n"
            "2. Where did students struggle or show misconceptions?\n"
            "3. Was the inquiry question answered by the end of the lesson?\n"
            "4. What would you change if teaching this lesson again?"
        )
    _cell_para_auto(t.rows[1].cells[0], refl_text)


def _build_table_E(doc, summary_row=None):
    """Table E — Summary Table Prompt (4r × 2c) with pre-filled example content."""
    t = _new_table(doc, 4, 2)
    _col_widths(t, [2.083, 7.417])

    # R0: merged purple header
    c = _merge_row(t, 0)
    _shade(c, C_PURPLE)
    _cell_para(c, "E. SUMMARY TABLE PROMPT  (pre-filled example for this lesson)",
               bold=True, size_pt=11, color_hex=C_WHITE)

    prompts = [
        "What did I observe?",
        "What did I learn?",
        "How does this explain the phenomenon?",
    ]

    pre_filled = ["", "", ""]
    if summary_row and len(summary_row) >= 4:
        pre_filled[0] = summary_row[1] if len(summary_row) > 1 else ""
        pre_filled[1] = summary_row[2] if len(summary_row) > 2 else ""
        pre_filled[2] = summary_row[3] if len(summary_row) > 3 else ""

    for ri, (prompt, content) in enumerate(zip(prompts, pre_filled), start=1):
        _shade(t.rows[ri].cells[0], C_PURPLE_LT)
        _shade(t.rows[ri].cells[1], C_WHITE)
        _cell_para(t.rows[ri].cells[0], prompt, bold=True, size_pt=9)
        _cell_para_auto(t.rows[ri].cells[1], content or "")


def build_lesson_sequence_docx(subject_key: str, lessons: list[dict]) -> Document:
    meta = META[subject_key]
    doc  = Document()
    _set_page_landscape(doc)

    _doc_title(doc, f"{meta['subject']} | {meta['grade']} | {meta['substrand']}")
    _doc_subtitle(doc,
                  f"CBE Lesson Sequence  ·  {meta['lessons']} Lessons  ·  "
                  f"{meta['lessons'] * 2} Periods")

    # Parse full Section A from source for rich sub-strand overview
    src_path = SOURCE_COMPLETE[subject_key]
    section_a_rows = parse_section_a(src_path)

    _build_table0_overview(doc, meta, section_a_rows)
    _tbl_no_spacing(doc)
    _tbl_no_spacing(doc)  # two empty paras after overview (matching Bio reference)

    summary_rows = meta["summary_table_rows"]

    for i, lesson in enumerate(lessons):
        lesson_num = lesson["number"]
        summary_row = summary_rows[lesson_num - 1] if lesson_num - 1 < len(summary_rows) else None

        _add_page_break(doc)

        _build_table_A(doc, lesson)
        _tbl_no_spacing(doc)
        _build_table_B(doc, lesson)
        _tbl_no_spacing(doc)
        # Period 1 — always present
        _build_table_C_period(doc, lesson, 1)
        _tbl_no_spacing(doc)
        # Period 2 — only if the parser found content for it
        if lesson.get("period2_table"):
            _build_table_C_period(doc, lesson, 2)
            _tbl_no_spacing(doc)
        _build_table_D(doc, lesson)
        _tbl_no_spacing(doc)
        _build_table_E(doc, summary_row)
        _tbl_no_spacing(doc)
        _tbl_no_spacing(doc)  # two empty paras between lessons

    # ── Sections C, D, E — parsed from source document ──────────────────────
    cde = parse_doc_sections(
        SOURCE_COMPLETE[subject_key],
        ["SECTION C", "SECTION D", "SECTION E"],
    )
    cde_display = [
        ("SECTION C", "SECTION C: SUMMARY TABLE TEMPLATE & INSTRUCTIONS", C_MED_BLUE),
        ("SECTION D", "SECTION D: DIFFERENTIATION STRATEGIES",             C_TEAL),
        ("SECTION E", "SECTION E: FINAL SYNTHESIS & ASSESSMENT",           C_NAVY),
    ]
    _add_page_break(doc)
    for key, display_header, fill in cde_display:
        content = cde.get(key, [("text", "See source document.")])
        _build_section_table(doc, display_header, content, header_fill=fill)
        _tbl_no_spacing(doc)

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# FINAL EXPLANATION BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

def build_final_explanation_docx(subject_key: str) -> Document:
    meta = META[subject_key]
    doc  = Document()
    _set_page_landscape(doc)

    subj_upper = meta["subject"].upper()
    substrand_topic = meta["substrand"].split(":")[-1].strip().upper()
    fe_title = f"FINAL EXPLANATION: {substrand_topic}"

    _doc_title(doc, f"{meta['subject']} | {meta['grade']} | {meta['substrand']}")
    _doc_subtitle(doc, "Final Explanation Document")

    # ── Table 0: Header block (5r × 2c) ──────────────────────────────────────
    t0 = _new_table(doc, 5, 2)
    _col_widths(t0, [2.083, 7.417])

    c = _merge_row(t0, 0)
    _shade(c, C_NAVY)
    _cell_para(c, fe_title, bold=True, size_pt=11, color_hex=C_WHITE)

    c = _merge_row(t0, 1)
    _shade(c, C_TEAL)
    _cell_para(c, f"{meta['subject']} {meta['grade']} — Student Assessment Document",
               bold=True, size_pt=11, color_hex=C_WHITE)

    for ri, label in enumerate(["Student Name", "Class", "Date"], start=2):
        _shade(t0.rows[ri].cells[0], C_LT_BLUE)
        _shade(t0.rows[ri].cells[1], C_WHITE)
        _cell_para(t0.rows[ri].cells[0], label, bold=True, size_pt=9)
        _cell_para(t0.rows[ri].cells[1],
                   "_______________________________________________", size_pt=9)

    _tbl_no_spacing(doc)

    # ── Table 1: Instructions (2r × 1c) ─────────────────────────────────────
    t1 = _new_table(doc, 2, 1)
    _col_widths(t1, [9.5])

    _shade(t1.rows[0].cells[0], C_TEAL)
    _cell_para(t1.rows[0].cells[0], "INSTRUCTIONS FOR STUDENTS",
               bold=True, size_pt=11, color_hex=C_WHITE)

    instructions = (
        f"You have completed all {meta['lessons']} lessons of {meta['substrand']}. "
        f"Write your COMPLETE EXPLANATION by answering all sections below. "
        f"Driving Question: {meta['driving_question']}\n\n"
        "Use evidence from investigations, discussions, and models from across this unit. "
        "Use scientific vocabulary accurately. Include labeled diagrams where appropriate. "
        "Show all calculations with steps. Connect your explanations back to the "
        "anchoring phenomenon throughout."
    )
    _shade(t1.rows[1].cells[0], C_WHITE)
    _cell_para_auto(t1.rows[1].cells[0], instructions)

    _tbl_no_spacing(doc)

    # ── Introductory block: Driving Question, Phenomenon, What is Final Explanation ──
    fe_intro = parse_doc_sections(
        SOURCE_FINAL_EXP[subject_key],
        ["Driving Question", "Anchoring Phenom", "What is a Final Explanation", "Instructions"],
    )
    intro_sections = [
        ("Driving Question",          "DRIVING QUESTION",           C_NAVY,    C_LT_BLUE),
        ("Anchoring Phenom",          "ANCHORING PHENOMENON",       C_TEAL,    C_WHITE),
        ("What is a Final Explanation", "WHAT IS A FINAL EXPLANATION?", C_MED_BLUE, C_WHITE),
        ("Instructions",              "INSTRUCTIONS",               C_MED_BLUE, C_WHITE),
    ]
    for key, display, hdr_fill, body_fill in intro_sections:
        content = fe_intro.get(key)
        if content:
            _build_section_table(doc, display, content, header_fill=hdr_fill, content_fill=body_fill)
            _tbl_no_spacing(doc)

    # ── Tables 2–N: One per section (3r × 1c each) ───────────────────────────
    sections = meta["final_explanation_sections"]
    for i, section_data in enumerate(sections):
        section_title = section_data[0]
        prompts = section_data[1]
        model_answer = section_data[2] if len(section_data) > 2 else ""

        t = _new_table(doc, 3, 1)
        _col_widths(t, [9.5])

        _shade(t.rows[0].cells[0], C_MED_BLUE)
        _cell_para(t.rows[0].cells[0], section_title, bold=True, size_pt=11, color_hex=C_WHITE)

        _shade(t.rows[1].cells[0], C_LT_BLUE)
        _cell_para_auto(t.rows[1].cells[0], prompts)

        _shade(t.rows[2].cells[0], C_WHITE)
        _cell_para_auto(t.rows[2].cells[0], model_answer)

        _tbl_no_spacing(doc)

    # ── Table N+1: Rubric ────────────────────────────────────────────────────
    rubric = meta["final_explanation_rubric"]
    rub_rows = rubric["rows"]
    num_cols = len(rubric["headers"])

    # Col widths: criterion 2.0" + remaining split evenly
    remaining = 9.5 - 2.0
    other_w = round(remaining / (num_cols - 1), 3) if num_cols > 1 else remaining
    col_widths = [2.0] + [other_w] * (num_cols - 1)

    t7 = _new_table(doc, 2 + len(rub_rows), num_cols)
    _col_widths(t7, col_widths)

    c = _merge_row(t7, 0)
    _shade(c, C_NAVY)
    _cell_para(c, "ASSESSMENT RUBRIC", bold=True, size_pt=11, color_hex=C_WHITE,
               align=WD_ALIGN_PARAGRAPH.CENTER)

    hdr_fills = [C_MED_BLUE, C_TEAL, C_MED_BLUE, C_TEAL, C_MED_BLUE][:num_cols]
    for cell, hdr, fill in zip(t7.rows[1].cells, rubric["headers"], hdr_fills):
        _shade(cell, fill)
        _cell_para(cell, hdr, bold=True, size_pt=9, color_hex=C_WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    alt = [C_WHITE, C_LT_GRAY, C_WHITE, C_LT_GRAY]
    for ri, row_data in enumerate(rub_rows):
        row = t7.rows[ri + 2]
        _shade(row.cells[0], C_LT_BLUE)
        _cell_para(row.cells[0], row_data[0], bold=True, size_pt=9)
        for ci, (cell, val, fill) in enumerate(zip(row.cells[1:], row_data[1:], alt)):
            _shade(cell, fill)
            _cell_para_auto(cell, val)

    # ── Post-rubric sections: Example, Tips, Reminders ───────────────────────
    fe_tail = parse_doc_sections(
        SOURCE_FINAL_EXP[subject_key],
        ["Example", "Partial Example", "Tips for Students", "Tips for Success",
         "Tips for Teachers", "Final Reminders", "Final Checklist"],
    )
    tail_sections = [
        ("Example",          "EXAMPLE OF A STRONG STUDENT RESPONSE", C_GREEN_LT, C_WHITE),
        ("Partial Example",  "EXAMPLE OF A STRONG STUDENT RESPONSE", C_GREEN_LT, C_WHITE),
        ("Tips for Students","TIPS FOR STUDENTS",                     C_LT_BLUE,  C_WHITE),
        ("Tips for Success", "TIPS FOR SUCCESS",                      C_LT_BLUE,  C_WHITE),
        ("Tips for Teachers","TIPS FOR TEACHERS",                     C_PURPLE_LT, C_WHITE),
        ("Final Reminders",  "FINAL REMINDERS",                       C_ORANGE_LT, C_WHITE),
        ("Final Checklist",  "FINAL CHECKLIST FOR STUDENTS",          C_ORANGE_LT, C_WHITE),
    ]
    seen_example = False
    seen_tips_students = False
    seen_tips_teachers = False
    seen_final = False
    for key, display, hdr_fill, body_fill in tail_sections:
        content = fe_tail.get(key)
        if not content:
            continue
        # Only show first matching variant for each logical group
        if key in ("Example", "Partial Example"):
            if seen_example:
                continue
            seen_example = True
        elif key in ("Tips for Students", "Tips for Success"):
            if seen_tips_students:
                continue
            seen_tips_students = True
        elif key == "Tips for Teachers":
            if seen_tips_teachers:
                continue
            seen_tips_teachers = True
        elif key in ("Final Reminders", "Final Checklist"):
            if seen_final:
                continue
            seen_final = True
        _build_section_table(doc, display, content, header_fill=hdr_fill, content_fill=body_fill)
        _tbl_no_spacing(doc)

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# SUMMARY TABLE BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

def build_summary_table_docx(subject_key: str) -> Document:
    meta = META[subject_key]
    doc  = Document()
    _set_page_landscape(doc)

    substrand_topic = meta["substrand"].split(":")[-1].strip().upper()
    st_title = f"SUMMARY TABLE: {meta['subject'].upper()} {meta['grade'].upper()} -- {substrand_topic}"

    _doc_title(doc, f"{meta['subject']} | {meta['grade']} | {meta['substrand']}")
    _doc_subtitle(doc, "Summary Table")

    # ── Purpose Statement (parsed from source) ──────────────────────────────
    st_intro = parse_doc_sections(
        SOURCE_SUMMARY_TABLE[subject_key],
        ["Purpose Statement"],
    )
    purpose_content = st_intro.get("Purpose Statement", [])
    if purpose_content:
        _build_section_table(doc, "PURPOSE STATEMENT", purpose_content,
                             header_fill=C_NAVY, content_fill=C_WHITE)
        _tbl_no_spacing(doc)

    # ── Table 0: Header block (5r × 2c) ──────────────────────────────────────
    t0 = _new_table(doc, 5, 2)
    _col_widths(t0, [2.083, 7.417])

    c = _merge_row(t0, 0)
    _shade(c, C_NAVY)
    _cell_para(c, st_title, bold=True, size_pt=11, color_hex=C_WHITE)

    c = _merge_row(t0, 1)
    _shade(c, C_TEAL)
    _cell_para(c, "SUMMARY TABLE", bold=True, size_pt=11, color_hex=C_WHITE)

    _shade(t0.rows[2].cells[0], C_LT_BLUE)
    _shade(t0.rows[2].cells[1], C_WHITE)
    _cell_para(t0.rows[2].cells[0], "Sub-Strand", bold=True, size_pt=9)
    _cell_para(t0.rows[2].cells[1], meta["substrand"], size_pt=9)

    _shade(t0.rows[3].cells[0], C_LT_BLUE)
    _shade(t0.rows[3].cells[1], C_WHITE)
    _cell_para(t0.rows[3].cells[0], "Driving Question", bold=True, size_pt=9)
    _cell_para(t0.rows[3].cells[1], meta["driving_question"], size_pt=9)

    _shade(t0.rows[4].cells[0], C_ORANGE_LT)
    _shade(t0.rows[4].cells[1], C_WHITE)
    _cell_para(t0.rows[4].cells[0], "Anchoring Phenomenon", bold=True, size_pt=9)
    _cell_para(t0.rows[4].cells[1], meta["phenomenon"], size_pt=9)

    _tbl_no_spacing(doc)

    # ── Table 1: Instructions (2r × 1c) ─────────────────────────────────────
    t1 = _new_table(doc, 2, 1)
    _col_widths(t1, [9.5])

    _shade(t1.rows[0].cells[0], C_TEAL)
    _cell_para(t1.rows[0].cells[0], "HOW TO USE THIS TABLE", bold=True, size_pt=11, color_hex=C_WHITE)

    _shade(t1.rows[1].cells[0], C_WHITE)
    _cell_para_auto(t1.rows[1].cells[0], meta["summary_instructions"])

    _tbl_no_spacing(doc)

    # ── Table 2: Main summary table (N+1 rows × 5c) ──────────────────────────
    COL_W = [0.833, 2.167, 2.167, 2.167, 2.167]
    HDR_FILLS = [C_NAVY, C_MED_BLUE, C_TEAL, C_MED_BLUE, C_PURPLE]

    data_rows = meta["summary_table_rows"]
    t2 = _new_table(doc, 1 + len(data_rows), 5)
    _col_widths(t2, COL_W)

    headers = meta["summary_table_headers"]
    for cell, hdr, fill in zip(t2.rows[0].cells, headers, HDR_FILLS):
        _shade(cell, fill)
        _cell_para(cell, hdr, bold=True, size_pt=9, color_hex=C_WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row_data in enumerate(data_rows):
        row = t2.rows[ri + 1]
        alt_fill = C_WHITE if ri % 2 == 0 else C_LT_GRAY

        _shade(row.cells[0], C_LT_BLUE)
        _cell_para(row.cells[0], row_data[0] if row_data else "", bold=True, size_pt=9)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        for ci in range(1, 4):
            _shade(row.cells[ci], alt_fill)
            _cell_para_auto(row.cells[ci], row_data[ci] if len(row_data) > ci else "")
            row.cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        _shade(row.cells[4], C_PURPLE_LT)
        _cell_para_auto(row.cells[4], row_data[4] if len(row_data) > 4 else "")
        row.cells[4].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ── End-of-Unit Reflection Questions + Teacher Notes (parsed from source) ─
    st_tail = parse_doc_sections(
        SOURCE_SUMMARY_TABLE[subject_key],
        ["End-of-Unit Reflection", "Reflection Questions",
         "Teacher Notes", "DQB and Model Tracking"],
    )
    tail_configs = [
        ("DQB and Model Tracking",    "DQB AND MODEL TRACKING GUIDE",             C_PURPLE,    C_WHITE),
        ("End-of-Unit Reflection",    "END-OF-UNIT REFLECTION QUESTIONS",          C_MED_BLUE,  C_WHITE),
        ("Reflection Questions",      "END-OF-UNIT REFLECTION QUESTIONS",          C_MED_BLUE,  C_WHITE),
        ("Teacher Notes",             "TEACHER NOTES: FORMATIVE USE OF THIS TABLE",C_NAVY,      C_WHITE),
    ]
    seen_reflection = False
    for key, display, hdr_fill, body_fill in tail_configs:
        content = st_tail.get(key)
        if not content:
            continue
        if key in ("End-of-Unit Reflection", "Reflection Questions"):
            if seen_reflection:
                continue
            seen_reflection = True
        _tbl_no_spacing(doc)
        _build_section_table(doc, display, content, header_fill=hdr_fill, content_fill=body_fill)

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    for key in ("chemistry", "mathematics", "physics"):
        meta = META[key]
        print(f"\n{'='*60}")
        print(f"Processing: {meta['subject']} — {meta['substrand']}")

        # ── Parse source ──────────────────────────────────────────────────────
        src_path = SOURCE_COMPLETE[key]
        lessons  = parse_source(src_path)
        print(f"  Parsed {len(lessons)} lessons (expected {meta['lessons']})")
        if len(lessons) != meta["lessons"]:
            print(f"  WARNING: lesson count mismatch!")

        for lesson in lessons:
            p1 = len(lesson["period1_table"])
            p2 = len(lesson["period2_table"])
            print(f"    Lesson {lesson['number']:2d}: {lesson['title'][:50]!r}"
                  f"  P1={p1} P2={p2}")

        # Fill missing inquiry questions
        for lesson in lessons:
            if not lesson["inquiry_question"]:
                lesson["inquiry_question"] = lesson["title"]

        # ── Build & save 3 documents ─────────────────────────────────────────
        out = OUTPUT_FILES[key]

        doc_ls = build_lesson_sequence_docx(key, lessons)
        doc_ls.save(str(out["lesson_seq"]))
        print(f"  ✓ Saved: {out['lesson_seq'].name}")

        doc_fe = build_final_explanation_docx(key)
        doc_fe.save(str(out["final_exp"]))
        print(f"  ✓ Saved: {out['final_exp'].name}")

        doc_st = build_summary_table_docx(key)
        doc_st.save(str(out["summary_table"]))
        print(f"  ✓ Saved: {out['summary_table'].name}")

    print(f"\n{'='*60}")
    print(f"Done. 9 files written to {OUT_DIR}/")


if __name__ == "__main__":
    main()
