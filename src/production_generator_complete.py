"""
Production CBE Lesson Generator - COLUMN WIDTHS ADJUSTED
2-column tables now use 1/3 : 2/3 split for better content display
"""
import os
import sys
import re
from pathlib import Path
from typing import Dict, List, Optional
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).parent))

try:
    from resource_manager import ResourceManager
except ImportError:
    print("⚠ Resource manager not found!")
    sys.exit(1)

load_dotenv()


class ProductionLessonGenerator:
    """Production generator with improved column widths"""
    
    def __init__(self):
        print("\n" + "="*70)
        print("PRODUCTION CBE LESSON GENERATOR - COLUMN WIDTHS IMPROVED")
        print("="*70 + "\n")
        
        self.resources = ResourceManager()
        self.config = self.resources.get_config()
        
        print("→ Initializing Claude API...")
        import anthropic
        self.anthropic_key = os.getenv("ANTHROPIC_API_KEY")
        if not self.anthropic_key:
            raise ValueError("Missing ANTHROPIC_API_KEY")
        self.claude = anthropic.Anthropic(api_key=self.anthropic_key)
        print("✓ Claude API initialized\n")
    
    def extract_sections(self, content: str) -> Dict[str, str]:
        """Extract sections using regex"""
        sections = {}
        pattern = r'===SECTION ([A-E]):.*?===(.*?)(?====SECTION|$)'
        matches = re.findall(pattern, content, re.DOTALL)
        
        for section_letter, section_content in matches:
            sections[section_letter] = section_content.strip()
        
        return sections
    
    def generate_substrand_overview(self, subject: str, grade: str, 
                                    strand: str, substrand: str) -> str:
        """Generate sub-strand overview"""
        
        curriculum_pdf = self.resources.get_curriculum_pdf(subject, grade)
        
        prompt = f"""Generate Sub-Strand Overview for Kenya CBC.

Subject: {subject}, Grade: {grade}, Sub-Strand: {substrand}

Return in this format:

===SUB-STRAND OVERVIEW===
Component: Specific Learning Outcomes
Content: [List learning outcomes a-f, each on new line]

Component: Key Inquiry Questions  
Content: [2-3 questions]

Component: Core Competencies
Content: [List competencies]

Component: Values
Content: [List values]

Component: PCIs
Content: [List of PCIs]"""

        response = self.claude.messages.create(
            model=self.config['generation']['model'],
            max_tokens=4000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        return response.content[0].text
    
    def generate_lesson_content(self, subject: str, grade: str, strand: str,
                                substrand: str, topic: str, number: int,
                                phenomenon: str = None, total_lessons: int = 22) -> str:
        """Generate lesson content"""
        
        prompt = f"""Generate COMPLETE CBE lesson for Kenya. MUST include ALL sections especially Section C with 7 rows!

Subject: {subject}, Topic: {topic}, Lesson: {number}/{total_lessons}
Phenomenon: {phenomenon or 'N/A'}

**REQUIRED FORMAT:**

===SECTION A: LEARNING OUTCOMES===
Knowledge:
- Point 1
- Point 2

Skills:
- Point 1

Attitudes:
- Point 1

Practices:
- Point 1

===SECTION B: OVERVIEW===
KeyQuestion: [question]
Purpose: [2-3 sentences]
Safety: [notes or n/a]

===SECTION C: FRAMEWORK===

ROW 1 - PREDICT:
Experience: [40+ words describing PREDICT phase - what students DO]
Resources: [Materials with [ARES VIDEO: ...]]
Teacher: [Facilitation with "PROBING:", "WAIT TIME: 10 seconds"]
Sensemaking: [Think-Pair-Share with timing]
Assessment: [☐ Checklist items]

ROW 2 - OBSERVE:
Experience: [40+ words describing OBSERVE phase]
Resources: [Materials with [ARES SIMULATION: ...]]
Teacher: [Facilitation with "PROBING:", "WAIT TIME: 10 seconds"]
Sensemaking: [Data collection, Think-Pair-Share]
Assessment: [☐ Observation criteria]

ROW 3 - EXPLAIN:
Experience: [40+ words describing EXPLAIN phase]
Resources: [Materials]
Teacher: [Questions ONLY - NO LECTURING! "PROBING:", "WAIT TIME: 15 seconds"]
Sensemaking: [Think-Pair-Share, consensus building]
Assessment: [☐ Quality criteria, Exit Ticket question]

ROW 4 - DQB:
Experience: [Class generates questions about {topic}. Teacher creates visible Driving Questions Board.]
Resources: [Chart paper, markers, DQB template, [ARES VIDEO: Intro to DQB]]
Teacher: [Establishing DQB, recording ALL questions, "PROBING:", "WAIT TIME:"]
Sensemaking: [Question Generation - list 4-5 example student questions about {topic}]
Assessment: [☐ DQB Started, ☐ Questions posted, note for Summary Table]

ROW 5 - INITIAL MODEL:
Experience: [Students draw their initial model of how {topic} works]
Resources: [Drawing materials, model template, [ARES SIMULATION: {topic} Model]]
Teacher: [Facilitating modeling, "no wrong models at this stage", "PROBING:", "WAIT TIME:"]
Sensemaking: [Initial model creation, Think-Pair-Share to compare models]
Assessment: [☐ Model quality, ☐ Creative thinking, collect models for later comparison]

ROW 6 - KENYA CONTEXT:
Experience: [How does understanding {topic} help Kenya? Relevant organizations (like KPLC for electricity), costs in KSh, local applications, careers in Kenya]
Resources: [Kenya-specific materials, photos, cost examples in KSh, [ARES VIDEO: {topic} in Kenya]]
Teacher: [Facilitating Kenya connection, "PROBING:", "WAIT TIME:", discussing real Kenyan issues]
Sensemaking: [Think-Pair-Share on Kenya applications, local organizations, careers]
Assessment: [☐ Identifies Kenyan context, ☐ 2+ applications, ☐ Career connections]

ROW 7 - MODEL REVISION:
Experience: [Students revise their initial models from ROW 5 based on evidence from today's lesson]
Resources: [Initial models from earlier, revision materials, colored pencils, before/after comparison template]
Teacher: [Facilitating revision, "What evidence now?", "PROBING:", "WAIT TIME:", "Scientists revise models!"]
Sensemaking: [Model revision showing evidence-based improvements, before/after comparison, written reflection]
Assessment: [☐ Shows improvements, ☐ Labels accurately, ☐ Explains changes, ☐ Uses evidence, note for Summary Table]

===SECTION D: REFLECTION===
1. [Reflection question about POE effectiveness]
2. [Reflection about probing questions and wait time]
3. [Reflection about DQB and future use]
4. [Reflection about what models revealed]
5. [Reflection about Kenya connections]
6. [Reflection about key learning/misconceptions]
7. [Student Feedback: "What surprised you?" - record 5-10 responses]

===SECTION E: SUMMARY PROMPT===
Teacher says: "Add Lesson {number} to your Summary Table."
Prompt for students:
- What did I observe? (Note: DQB Started, Initial Model Created, Model Revised)
- What did I learn? (Key concepts about {topic})
- How does this explain the phenomenon? (Connection to {phenomenon if phenomenon else topic})

Generate the complete lesson now with ALL sections!"""

        response = self.claude.messages.create(
            model=self.config['generation']['model'],
            max_tokens=self.config['generation']['max_tokens'],
            messages=[{"role": "user", "content": prompt}]
        )
        
        return response.content[0].text
    
    def create_word_document(self, content: str, overview_content: str,
                            subject: str, grade: str, substrand: str,
                            topic: str, number: int, total_lessons: int) -> str:
        """Create Word document with improved column widths"""
        
        doc = Document()
        
        # LANDSCAPE with NARROW margins
        section = doc.sections[0]
        section.orientation = 1
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Add overview if present
        if overview_content:
            self._add_substrand_overview(doc, overview_content, substrand)
            doc.add_page_break()
        
        # Add lesson
        self._add_lesson(doc, content, subject, grade, substrand, topic, number, total_lessons)
        
        # Save
        output_dir = Path(self.config['paths']['output_docx'])
        output_dir.mkdir(parents=True, exist_ok=True)
        
        filename = f"Lesson_{number:03d}_{topic.replace(' ', '_')}.docx"
        filepath = output_dir / filename
        
        doc.save(str(filepath))
        return str(filepath)
    
    def _add_substrand_overview(self, doc, content: str, substrand: str):
        """Add sub-strand overview with ADJUSTED COLUMN WIDTHS"""
        title = doc.add_heading(f'SUB-STRAND: {substrand.upper()}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(0, 70, 135)
        
        subtitle = doc.add_paragraph('Sub-Strand Overview')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.italic = True
        doc.add_paragraph()
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # ADJUSTED WIDTHS: 2.5" label, 7.5" content (10" total usable width)
        for row in table.rows:
            row.cells[0].width = Inches(2.5)  # 25% - narrower label column
            row.cells[1].width = Inches(7.5)  # 75% - wider content column
        
        hdr = table.rows[0].cells
        hdr[0].text = 'Component'
        hdr[1].text = 'Description'
        for cell in hdr:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Parse components
        for component in ['Specific Learning Outcomes', 'Key Inquiry Questions', 
                         'Core Competencies', 'Values', 'PCIs']:
            if f'Component: {component}' in content:
                row = table.add_row()
                row.cells[0].width = Inches(2.5)
                row.cells[1].width = Inches(7.5)
                row.cells[0].text = component
                
                # Extract content
                start = content.find(f'Component: {component}')
                start = content.find('Content:', start) + len('Content:')
                end = content.find('Component:', start)
                if end == -1:
                    end = len(content)
                
                row.cells[1].text = content[start:end].strip()
    
    def _add_lesson(self, doc, content: str, subject: str, grade: str,
                   substrand: str, topic: str, number: int, total_lessons: int):
        """Add lesson with improved formatting"""
        
        # Title
        title = doc.add_heading(f'LESSON {number}: {topic.upper()}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(0, 70, 135)
        
        subtitle = doc.add_paragraph(f'{grade} {subject} - {substrand} - Lesson {number} of {total_lessons}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Extract sections
        sections = self.extract_sections(content)
        print(f"  → Sections extracted: {list(sections.keys())}")
        
        # Add sections
        if 'A' in sections:
            self._add_section_a(doc, sections['A'])
        if 'B' in sections:
            self._add_section_b(doc, sections['B'])
        if 'C' in sections:
            print(f"  ✓ Section C found ({len(sections['C'])} chars)")
            self._add_section_c(doc, sections['C'])
        else:
            print("  ✗ Section C NOT FOUND!")
        if 'D' in sections:
            self._add_section_d(doc, sections['D'])
        if 'E' in sections:
            self._add_section_e(doc, sections['E'])
    
    def _add_section_a(self, doc, section_text):
        """Add Learning Outcomes with ADJUSTED COLUMN WIDTHS"""
        doc.add_heading('A. Specific Learning Outcomes', 2)
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # ADJUSTED WIDTHS: 2.5" label, 7.5" content
        for row in table.rows:
            row.cells[0].width = Inches(2.5)
            row.cells[1].width = Inches(7.5)
        
        hdr = table.rows[0].cells
        hdr[0].text = 'Category'
        hdr[1].text = 'Outcomes'
        for c in hdr:
            c.paragraphs[0].runs[0].font.bold = True
        
        categories = ['Knowledge:', 'Skills:', 'Attitudes:', 'Practices:']
        for i, cat in enumerate(categories, 1):
            if cat in section_text:
                table.rows[i].cells[0].text = cat.replace(':', '')
                
                start = section_text.find(cat) + len(cat)
                next_cats = [c for c in categories if c != cat]
                end = len(section_text)
                for nc in next_cats:
                    pos = section_text.find(nc, start)
                    if pos > start:
                        end = min(end, pos)
                
                text = section_text[start:end]
                lines = [l.strip('- ').strip() for l in text.split('\n') if l.strip().startswith('-')]
                table.rows[i].cells[1].text = '\n'.join(f'• {l}' for l in lines if l)
        
        doc.add_paragraph()
    
    def _add_section_b(self, doc, section_text):
        """Add Overview with ADJUSTED COLUMN WIDTHS"""
        doc.add_heading('B. Overview', 2)
        
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # ADJUSTED WIDTHS: 2.5" label, 7.5" content
        for row in table.rows:
            row.cells[0].width = Inches(2.5)
            row.cells[1].width = Inches(7.5)
        
        hdr = table.rows[0].cells
        hdr[0].text = 'Element'
        hdr[1].text = 'Details'
        for c in hdr:
            c.paragraphs[0].runs[0].font.bold = True
        
        items = [
            ('Key Inquiry Question', 'KeyQuestion:'),
            ('Purpose in Storyline', 'Purpose:'),
            ('Safety Notes', 'Safety:')
        ]
        
        for i, (label, marker) in enumerate(items, 1):
            if marker in section_text:
                table.rows[i].cells[0].text = label
                
                start = section_text.find(marker) + len(marker)
                end = len(section_text)
                for next_marker in ['KeyQuestion:', 'Purpose:', 'Safety:']:
                    if next_marker != marker:
                        pos = section_text.find(next_marker, start)
                        if pos > start:
                            end = min(end, pos)
                
                text = section_text[start:end].strip()
                table.rows[i].cells[1].text = text
        
        doc.add_paragraph()
    
    def _add_section_c(self, doc, section_text):
        """Add Implementation Framework - 5 EQUAL COLUMNS"""
        doc.add_heading('C. Lesson Implementation Framework', 2)
        
        rows = section_text.split('ROW ')[1:]
        row_count = len(rows)
        
        print(f"    → Creating table with {row_count} content rows")
        
        if row_count == 0:
            doc.add_paragraph("⚠ ERROR: No rows detected")
            return
        
        table = doc.add_table(rows=row_count + 1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # EQUAL column widths - 2.0 inches each (Section C stays equal)
        for row in table.rows:
            for i in range(5):
                row.cells[i].width = Inches(2.0)
        
        # Headers
        headers = ['Learner Experience', 'Resource Link', 'Teacher Moves',
                   'Sensemaking Strategy', 'Formative Assessment']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Fill rows
        for row_idx, row_text in enumerate(rows, 1):
            if row_idx >= len(table.rows):
                break
            
            columns_data = {'Experience': '', 'Resources': '', 'Teacher': '', 
                          'Sensemaking': '', 'Assessment': ''}
            
            if 'Experience:' in row_text:
                start = row_text.find('Experience:') + len('Experience:')
                end = row_text.find('Resources:', start) if 'Resources:' in row_text[start:] else len(row_text)
                columns_data['Experience'] = row_text[start:end].strip()
            
            if 'Resources:' in row_text:
                start = row_text.find('Resources:') + len('Resources:')
                end = row_text.find('Teacher:', start) if 'Teacher:' in row_text[start:] else len(row_text)
                columns_data['Resources'] = row_text[start:end].strip()
            
            if 'Teacher:' in row_text:
                start = row_text.find('Teacher:') + len('Teacher:')
                end = row_text.find('Sensemaking:', start) if 'Sensemaking:' in row_text[start:] else len(row_text)
                columns_data['Teacher'] = row_text[start:end].strip()
            
            if 'Sensemaking:' in row_text:
                start = row_text.find('Sensemaking:') + len('Sensemaking:')
                end = row_text.find('Assessment:', start) if 'Assessment:' in row_text[start:] else len(row_text)
                columns_data['Sensemaking'] = row_text[start:end].strip()
            
            if 'Assessment:' in row_text:
                start = row_text.find('Assessment:') + len('Assessment:')
                columns_data['Assessment'] = row_text[start:].strip()
            
            table.rows[row_idx].cells[0].text = columns_data['Experience']
            table.rows[row_idx].cells[1].text = columns_data['Resources']
            table.rows[row_idx].cells[2].text = columns_data['Teacher']
            table.rows[row_idx].cells[3].text = columns_data['Sensemaking']
            table.rows[row_idx].cells[4].text = columns_data['Assessment']
        
        print(f"    ✓ Section C table created with {row_count} rows")
        doc.add_paragraph()
    
    def _add_section_d(self, doc, section_text):
        """Add Teacher Reflection"""
        doc.add_heading('D. Teacher Reflection', 2)
        
        lines = section_text.split('\n')
        for line in lines:
            line = line.strip()
            if line and line[0].isdigit() and '.' in line[:3]:
                doc.add_paragraph(line, style='List Number')
        
        doc.add_paragraph()
    
    def _add_section_e(self, doc, section_text):
        """Add Summary Prompt"""
        doc.add_heading('E. Summary Table Prompt', 2)
        
        lines = [l.strip() for l in section_text.split('\n') if l.strip()]
        for line in lines:
            if line.startswith('- ') or line.startswith('•'):
                doc.add_paragraph(line.strip('- •').strip(), style='List Bullet')
            elif line and not line.startswith('==='):
                doc.add_paragraph(line)
    
    def generate_complete_lesson(self, subject: str, grade: str, strand: str,
                                 substrand: str, topic: str, number: int = 1,
                                 phenomenon: str = None, total_lessons: int = 22,
                                 include_overview: bool = True) -> Dict:
        """Generate complete lesson"""
        
        print(f"\n{'='*70}")
        print(f"LESSON {number}: {topic}")
        print(f"{'='*70}\n")
        
        overview_content = ""
        if include_overview and number == 1:
            print("  → Generating sub-strand overview...")
            overview_content = self.generate_substrand_overview(
                subject, grade, strand, substrand
            )
            print("  ✓ Overview generated")
        
        print("  → Generating lesson content...")
        content = self.generate_lesson_content(
            subject, grade, strand, substrand, topic, number, phenomenon, total_lessons
        )
        print(f"  ✓ Content generated ({len(content)} chars)")
        
        if '===SECTION C:' in content:
            print("  ✓ Section C present")
        else:
            print("  ✗ WARNING: Section C missing!")
        
        print("  → Creating Word document...")
        filepath = self.create_word_document(
            content, overview_content, subject, grade, substrand,
            topic, number, total_lessons
        )
        print(f"  ✓ Document created: {Path(filepath).name}")
        
        return {
            'lesson_number': number,
            'topic': topic,
            'filepath': filepath,
            'status': 'success'
        }


def main():
    """Test"""
    generator = ProductionLessonGenerator()
    
    result = generator.generate_complete_lesson(
        subject="Chemistry",
        grade="Grade 10",
        strand="Matter",
        substrand="Atomic Structure",
        topic="Test Lesson",
        number=1,
        total_lessons=18
    )
    
    print(f"\n✓ Generated: {result['filepath']}")


if __name__ == "__main__":
    main()
