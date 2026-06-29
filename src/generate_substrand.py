#!/usr/bin/env python3
"""
generate_substrand.py — CBE Lesson Plan Content Generator
==========================================================
Generates a complete sub-strand data module for the universal generator
by calling the Claude API with KICD curriculum content and teacher templates.

Usage:
    python3 src/generate_substrand.py \\
        --subject biology \\
        --substrand 1.4 \\
        --output bio_1_4 \\
        --lessons 6 \\
        [--template "data/raw/CBE LESSON TEMPLATES/Biology 10.1.4 ...docx"] \\
        [--run]

Outputs: generators/data/<output>_data.js
Then optionally runs: node generators/generate.js <output>
"""

import argparse
import json
import os
import re
import sys
import time
from pathlib import Path

# ── Bootstrap ────────────────────────────────────────────────────────────────

PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT / 'src'))

try:
    from dotenv import load_dotenv
    load_dotenv(PROJECT_ROOT / '.env')
except ImportError:
    pass

try:
    import anthropic
except ImportError:
    print("ERROR: anthropic package not found. Run: pip install anthropic")
    sys.exit(1)

MODEL = os.environ.get('CLAUDE_MODEL', 'claude-sonnet-4-6')
CLIENT = anthropic.Anthropic(api_key=os.environ.get('ANTHROPIC_API_KEY'))

# ── Curriculum extraction ─────────────────────────────────────────────────────

def extract_curriculum_pdf(pdf_path: str, substrand_id: str) -> str:
    """Extract sub-strand section from KICD curriculum PDF."""
    try:
        import pdfminer.high_level as pdfminer
        text = pdfminer.extract_text(pdf_path)
    except Exception as e:
        print(f"  WARNING: Could not extract PDF: {e}")
        return ""

    # Search for sub-strand section using various patterns
    subject_num = substrand_id.split('.')[0]
    sub_num     = substrand_id.split('.')[1] if '.' in substrand_id else ''

    # Try to find the section
    patterns = [
        f"{subject_num}.{sub_num} ",
        f"Sub Strand {subject_num}.{sub_num}",
        f"SUB STRAND {subject_num}.{sub_num}",
    ]

    start_idx = -1
    for pattern in patterns:
        idx = text.find(pattern, 20000)   # skip past table of contents
        if idx > 0:
            # Verify it's the content section (not TOC) by checking for SLO language
            snippet = text[idx:idx+500]
            if any(kw in snippet for kw in ['learner should', 'Learner should', 'Learning Outcomes', 'By the end']):
                start_idx = idx
                break
            elif start_idx == -1:
                start_idx = idx   # use first hit as fallback

    if start_idx == -1:
        print(f"  WARNING: Sub-strand {substrand_id} not found in curriculum PDF")
        return ""

    # Extract ~4000 chars which should cover the full sub-strand entry
    raw = text[start_idx:start_idx + 4500]
    # Clean up the fragmented PDF text
    raw = re.sub(r'  +', ' ', raw)        # collapse multiple spaces
    raw = re.sub(r'\n +\n', '\n\n', raw)  # clean blank lines
    return raw.strip()


def extract_template_docx(docx_path: str) -> dict:
    """Extract content from teacher-authored SoW template docx."""
    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        print(f"  WARNING: Could not read template: {e}")
        return {}

    result = {
        'paragraphs': [],
        'table_content': [],
        'phenomenon': '',
        'lesson_sequence': '',
        'learning_outcomes': '',
        'core_competencies': [],
        'core_values': [],
        'evidence_activities': '',
        'final_explanation_notes': '',
    }

    # Extract paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            result['paragraphs'].append(text)
            # Capture specific fields
            if 'Learning Outcomes' in text and len(text) > 50:
                result['learning_outcomes'] = text
            if 'Student-facing objective' in text:
                result['student_objective'] = text

    # Detect competencies and values from paragraphs
    capture_comp = False
    capture_val  = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if 'CORE COMPETENCIES' in text:
            capture_comp = True
            capture_val  = False
            continue
        if 'CORE VALUES' in text:
            capture_val  = True
            capture_comp = False
            continue
        if capture_comp and text and ':' in text:
            result['core_competencies'].append(text.split(':')[0].strip())
        if capture_val and text and len(text) > 5:
            result['core_values'].append(text)

    # Extract table content
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                ct = cell.text.strip()
                if ct:
                    row_data.append(ct)
            if row_data:
                result['table_content'].append(row_data)

    # Find phenomenon (usually in row with "PHENOMENON")
    for row in result['table_content']:
        for i, cell in enumerate(row):
            if 'PHENOMENON' in cell.upper() and i + 1 < len(row):
                result['phenomenon'] = row[i + 1][:500]
                break

    # Find lesson sequence
    for row in result['table_content']:
        for i, cell in enumerate(row):
            if 'LESSON SEQUENCE' in cell.upper() and i + 1 < len(row):
                result['lesson_sequence'] = row[i + 1][:800]
                break

    # Find evidence activities
    for row in result['table_content']:
        for i, cell in enumerate(row):
            if 'EVIDENCE GATHERED' in cell.upper() and i + 1 < len(row):
                result['evidence_activities'] = row[i + 1][:800]
                break

    # Find final explanation notes
    for row in result['table_content']:
        for i, cell in enumerate(row):
            if 'FINAL EXPLANATION' in cell.upper() and i + 1 < len(row):
                result['final_explanation_notes'] = row[i + 1][:400]
                break

    return result


# ── System prompt ─────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are an expert curriculum designer for Kenya's Grade 10 CBE (Competency-Based Curriculum).
You generate structured lesson plan content for the ARES Education offline learning platform.

YOUR OUTPUTS:
- Always respond with valid JSON only — no markdown fences, no explanations, no preamble
- Follow the exact JSON schema provided in each request
- Use Kenya-relevant contexts, examples, scientists, food items, and place names throughout
- Embed NGSS Science and Engineering Practices where indicated
- All KICD-mandated text (learning outcomes, competencies, values, PCIs) must be used verbatim
- Write teacher actions as specific, actionable instructions (include WAIT TIME, cold-call counts, specific questions)
- Connect every lesson phase back to the anchoring phenomenon

LESSON STRUCTURE (5 phases):
1. Predict Phase — students predict before instruction
2. Observe Phase — evidence gathering (experiment, video, reading, simulation)
3. Explain Phase — sensemaking, applying knowledge
4. Driving Question Board (DQB) Creation — tracking growing understanding
5. Model Building Phase — cumulative model revision

KENYA CONTEXT RULES:
- Use Kenyan food examples (ugali, sukuma wiki, nyama choma, githeri, mandazi, maize, beans)
- Reference Kenyan scientists, athletes, farmers where relevant
- Use Kenyan place names (Nairobi, Mombasa, Kisumu, Kericho, Rift Valley, Lake Victoria)
- Prefer local Kenyan contexts over Western examples
"""

# ── Claude API call ───────────────────────────────────────────────────────────

# ── Tool schemas for structured generation ────────────────────────────────────
# Generating via tool use (rather than free-text JSON) guarantees well-formed
# output and biases the model to the ares-contract shape. additionalProperties
# is False so stray keys (e.g. legacy 'storyline', top-level 'safetyNotes')
# cannot appear; required sets mirror docs/SCHEMA.md / ares-contract.schema.json.

def _s(desc: str = "") -> dict:
    return {"type": "string", "description": desc} if desc else {"type": "string"}

UNIT_TOOL_SCHEMA = {
    "type": "object", "additionalProperties": False,
    "properties": {
        "gradeLevel": _s(), "subject": _s(), "strand": _s(), "substrand": _s(),
        "totalDuration": _s(),
        "content": _s("KICD sub-strand content; bullet list, one topic per line prefixed with •"),
        "learningOutcomes": _s(), "coreCompetencies": _s(), "values": _s(),
        "sep": _s(), "pcis": _s(), "careers": _s(), "focus": _s(),
        "phenomenon": _s(), "supportingPhenomena": _s(),
        "drivingQuestion": _s(), "storylineThread": _s(),
    },
    "required": ["gradeLevel", "subject", "strand", "substrand", "totalDuration",
                 "content", "learningOutcomes", "coreCompetencies",
                 "phenomenon", "drivingQuestion", "storylineThread"],
}

LESSON_TOOL_SCHEMA = {
    "type": "object", "additionalProperties": False,
    "properties": {
        "number": {"type": "integer"},
        "title": _s(), "duration": _s(), "substrand": _s(), "aresKeywords": _s(),
        "slo": {
            "type": "object", "additionalProperties": False,
            "properties": {"purpose": _s(), "knowledge": _s(), "skills": _s(),
                           "attitudes": _s(), "keyInquiry": _s(),
                           "purposeInStoryline": _s(), "safetyNotes": _s()},
            "required": ["purpose", "knowledge", "skills", "attitudes",
                         "keyInquiry", "purposeInStoryline", "safetyNotes"],
        },
        "overview": _s(),
        "framework": {
            "type": "array",
            "items": {
                "type": "object", "additionalProperties": False,
                "properties": {"phase": _s(), "learnerExperience": _s(),
                               "teacherMoves": _s(), "sensemakingStrategy": _s(),
                               "formativeAssessment": _s()},
                "required": ["phase", "learnerExperience", "teacherMoves",
                             "sensemakingStrategy", "formativeAssessment"],
            },
        },
        "teacherReflection": _s(),
        "summaryTablePrompt": {
            "type": "object", "additionalProperties": False,
            "properties": {"observed": _s(), "learned": _s(), "explained": _s()},
            "required": ["observed", "learned", "explained"],
        },
    },
    "required": ["number", "title", "duration", "substrand", "slo", "overview",
                 "framework", "teacherReflection", "summaryTablePrompt"],
}

FE_TOOL_SCHEMA = {
    "type": "object", "additionalProperties": False,
    "properties": {
        "subjectLabel": _s(), "instructions": _s(),
        "sections": {"type": "array", "items": {
            "type": "object", "additionalProperties": False,
            "properties": {"title": _s(), "prompt": _s(), "exemplar": _s()},
            "required": ["title", "prompt", "exemplar"]}},
        "rubric": {"type": "array", "items": {
            "type": "object", "additionalProperties": False,
            "properties": {"criterion": _s(), "excellent": _s(),
                           "proficient": _s(), "developing": _s()},
            "required": ["criterion", "excellent", "proficient", "developing"]}},
    },
    "required": ["subjectLabel", "instructions", "sections"],
}

ST_TOOL_SCHEMA = {
    "type": "object", "additionalProperties": False,
    "properties": {
        "subStrand": _s(), "drivingQuestion": _s(),
        "lessons": {"type": "array", "items": {
            "type": "object", "additionalProperties": False,
            "properties": {"number": {"type": "integer"}, "title": _s(),
                           "observed": _s(), "learned": _s(), "explained": _s()},
            "required": ["number", "title", "observed", "learned", "explained"]}},
    },
    "required": ["subStrand", "drivingQuestion", "lessons"],
}


def _schema_violations(obj, schema, path: str = "") -> list:
    """Dependency-free structural check against the (simple) tool schemas above.
    Catches what crashes the docx layer or breaks the contract: a field declared
    array/object arriving as a string, or a missing/empty required field."""
    t = schema.get("type")
    types = t if isinstance(t, list) else ([t] if t else [])

    def kind(v):
        if isinstance(v, bool):  return "boolean"
        if isinstance(v, str):   return "string"
        if isinstance(v, int):   return "integer"
        if isinstance(v, float): return "number"
        if isinstance(v, list):  return "array"
        if isinstance(v, dict):  return "object"
        if v is None:            return "null"
        return "unknown"

    k = kind(obj)
    if types and k not in types and not (k == "integer" and "number" in types):
        return [f"{path or 'root'}: expected {types}, got {k}"]

    errs = []
    if "object" in types and isinstance(obj, dict):
        for r in schema.get("required", []):
            if r not in obj or obj[r] in (None, ""):
                errs.append(f"{path}{r}: missing/empty required field")
        for key, sub in schema.get("properties", {}).items():
            if key in obj and obj[key] is not None:
                errs += _schema_violations(obj[key], sub, f"{path}{key}.")
    if "array" in types and isinstance(obj, list):
        item = schema.get("items")
        if item:
            for i, el in enumerate(obj):
                errs += _schema_violations(el, item, f"{path}[{i}].")
    return errs


def call_claude(user_prompt: str, max_tokens: int = 8000, retries: int = 3,
                schema: dict | None = None, tool_name: str = "emit_data") -> dict | None:
    """Call Claude and return a dict.

    If `schema` (a JSON Schema) is given, the model returns the data through a
    forced tool call and the SDK hands back an already-parsed dict — so
    malformed-JSON failures (unescaped quotes, missing delimiters, fence noise)
    cannot occur. Without a schema, falls back to free-text JSON parsing.
    """
    for attempt in range(retries):
        try:
            kwargs = dict(
                model=MODEL,
                max_tokens=max_tokens,
                system=SYSTEM_PROMPT,
                messages=[{"role": "user", "content": user_prompt}],
            )
            if schema is not None:
                kwargs["tools"] = [{
                    "name": tool_name,
                    "description": "Return the requested lesson-plan data as structured fields.",
                    "input_schema": schema,
                }]
                kwargs["tool_choice"] = {"type": "tool", "name": tool_name}

            response = CLIENT.messages.create(**kwargs)

            # Explicit truncation detection (previously surfaced only as a parse error)
            if getattr(response, "stop_reason", None) == "max_tokens":
                print(f"    Truncated (stop_reason=max_tokens, max_tokens={max_tokens}, attempt {attempt+1})")
                if attempt < retries - 1:
                    time.sleep(2)
                    continue
                return None

            if schema is not None:
                data = None
                for block in response.content:
                    if getattr(block, "type", None) == "tool_use" and block.name == tool_name:
                        data = block.input
                        break
                if data is None:
                    print(f"    No tool_use block returned (attempt {attempt+1})")
                    if attempt < retries - 1:
                        time.sleep(2)
                        continue
                    return None
                violations = _schema_violations(data, schema)
                if violations:
                    # e.g. model stuffed an array field into a string — re-roll
                    print(f"    Tool output off-schema (attempt {attempt+1}): {violations[:3]}")
                    if attempt < retries - 1:
                        time.sleep(2)
                        continue
                    return None
                return data

            raw = response.content[0].text.strip()
            raw = re.sub(r'^```(?:json)?\s*', '', raw)
            raw = re.sub(r'\s*```$', '', raw)
            raw = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', raw)
            return json.loads(raw)

        except json.JSONDecodeError as e:
            print(f"    JSON parse error (attempt {attempt+1}): {e}")
            if attempt < retries - 1:
                time.sleep(2)
        except anthropic.APIError as e:
            print(f"    API error (attempt {attempt+1}): {e}")
            if attempt < retries - 1:
                time.sleep(5)
        except Exception as e:
            print(f"    Unexpected error (attempt {attempt+1}): {e}")
            if attempt < retries - 1:
                time.sleep(2)

    return None


# ── Unit generation ───────────────────────────────────────────────────────────

def generate_unit(curriculum_text: str, template: dict, args) -> dict | None:
    """Generate the UNIT sub-strand overview data."""
    print("  Generating UNIT data...")

    prompt = f"""Generate the sub-strand overview (UNIT) data for:
Subject: {args.subject.capitalize()}
Grade: 10
Sub-strand: {args.substrand} — {args.substrand_name}
Number of lessons: {args.lessons}

KICD CURRICULUM CONTENT (use verbatim for SLOs, competencies, values, PCIs):
{curriculum_text}

TEACHER TEMPLATE — PHENOMENON:
{template.get('phenomenon', '')}

TEACHER TEMPLATE — LESSON SEQUENCE OUTLINE:
{template.get('lesson_sequence', '')}

TEACHER TEMPLATE — LEARNING OUTCOMES:
{template.get('learning_outcomes', '')}

Return ONLY this JSON structure (no other text):
{{
  "gradeLevel": "10",
  "subject": "{args.subject.capitalize()}",
  "strand": "Strand {args.substrand.split('.')[0]}.0: [official KICD strand name]",
  "substrand": "Sub-Strand {args.substrand}: {args.substrand_name}",
  "totalDuration": "{args.lessons} lessons × 40 minutes = {args.lessons * 40} minutes total",
  "content": "bullet list of sub-strand content topics, one per line prefixed with •",
  "learningOutcomes": "By the end of the sub-strand, the learner should be able to:\\na) ...\\nb) ...\\nc) ...\\nd) ...\\ne) ...",
  "coreCompetencies": "• Competency name: description of how developed in this sub-strand\\n• ...",
  "values": "• Value name: description\\n• ...",
  "sep": "• SEP name: how embedded\\n• ...",
  "pcis": "• PCI: description\\n• ...",
  "careers": "• Career: how this sub-strand connects\\n• ...",
  "focus": "2-3 sentence overview of the unit focus and approach",
  "totalDuration": "{args.lessons} lessons (approximately {args.lessons * 2} periods × 40 minutes = {args.lessons * 80} minutes)",
  "phenomenon": "Full description of the anchoring phenomenon including what students observe and why it is puzzling",
  "supportingPhenomena": "• 2-4 additional supporting phenomena\\n• ...",
  "drivingQuestion": "DRIVING QUESTION: [the main driving question]\\n\\nKICD KEY INQUIRY QUESTIONS:\\n1. [verbatim from KICD]\\n2. [verbatim from KICD]",
  "storylineThread": "Lesson 1: [brief description]\\nLesson 2: ...\\n[continue for all {args.lessons} lessons]"
}}"""

    return call_claude(prompt, max_tokens=6000, schema=UNIT_TOOL_SCHEMA)


# ── Lesson generation ─────────────────────────────────────────────────────────

LESSON_SCHEMA = {
    "number": 0,
    "title": "Lesson title — subtitle",
    "duration": "40 minutes",
    "substrand": "Sub-Strand X.Y: Name",
    "aresKeywords": "4-5 specific biology content keywords for ARES search",
    "slo": {
        "purpose": "One sentence describing lesson purpose in storyline",
        "knowledge": "• bullet 1\\n• bullet 2\\n• bullet 3",
        "skills": "• bullet 1\\n• bullet 2",
        "attitudes": "• bullet 1\\n• bullet 2",
        "keyInquiry": "The key question this lesson addresses",
        "purposeInStoryline": "How this lesson advances the storyline",
        "safetyNotes": "Safety considerations or 'No specific hazards.'"
    },
    "overview": "2-3 paragraph prose overview of the lesson...",
    "framework": [
        {
            "phase": "Predict Phase",
            "learnerExperience": "What students do (2-4 sentences)",
            "teacherMoves": "Specific teacher actions with exact quotes, WAIT TIME notes, cold-call counts",
            "sensemakingStrategy": "Name and description of strategy used",
            "formativeAssessment": "What teacher looks for and how to assess"
        },
        {"phase": "Observe Phase", "learnerExperience": "...", "teacherMoves": "...", "sensemakingStrategy": "...", "formativeAssessment": "..."},
        {"phase": "Explain Phase", "learnerExperience": "...", "teacherMoves": "...", "sensemakingStrategy": "...", "formativeAssessment": "..."},
        {"phase": "Driving Question Board (DQB) Creation", "learnerExperience": "...", "teacherMoves": "...", "sensemakingStrategy": "...", "formativeAssessment": "..."},
        {"phase": "Model Building Phase", "learnerExperience": "...", "teacherMoves": "...", "sensemakingStrategy": "...", "formativeAssessment": "..."}
    ],
    "teacherReflection": "3-5 numbered reflection questions for teacher after teaching",
    "summaryTablePrompt": {
        "observed": "What students observed in this lesson",
        "learned": "Key learning from this lesson",
        "explained": "How this explains the phenomenon"
    }
}


def generate_lesson(num: int, curriculum_text: str, template: dict,
                    unit: dict, prev_summaries: list, args) -> dict | None:
    """Generate a single lesson."""
    print(f"  Generating Lesson {num}/{args.lessons}...")

    # Build context from previous lessons
    prev_context = ""
    if prev_summaries:
        prev_context = "PREVIOUS LESSONS (for storyline continuity):\n"
        for s in prev_summaries:
            prev_context += f"  Lesson {s['number']}: {s['title']} — {s.get('summary', '')}\n"

    lesson_pos = "opening lesson that launches the phenomenon" if num == 1 else \
                 "final lesson with Final Explanation" if num == args.lessons else \
                 f"lesson {num} of {args.lessons} in the evidence-gathering sequence"

    prompt = f"""Generate Lesson {num} ({lesson_pos}) for:
Subject: {args.subject.capitalize()} Grade 10
Sub-strand: {args.substrand} {args.substrand_name}
Driving question: {unit.get('drivingQuestion', '')}
Phenomenon: {unit.get('phenomenon', '')}
Storyline thread for this lesson: {_get_lesson_storyline(unit, num)}

KICD CURRICULUM CONTENT:
{curriculum_text}

TEACHER TEMPLATE — EVIDENCE ACTIVITIES:
{template.get('evidence_activities', '')}

{prev_context}

RULES:
- All learner experiences must connect back to the phenomenon
- Teacher moves must include specific quoted phrases, WAIT TIME (10-15 seconds minimum), cold-call counts
- Sensemaking strategies: name the strategy and explain how it builds understanding
- Formative assessment: specific, observable indicators
- Kenya-relevant contexts throughout (local food examples, Kenyan scientists, local contexts)
- Lesson {num} should {"introduce the phenomenon and open the DQB" if num == 1 else "build on previous evidence and advance the driving question"}
{"- Final lesson: focus on Final Explanation, model comparison L1 vs final, DQB completion" if num == args.lessons else ""}

Use the emit_data tool to return the lesson with rich, real content in every field.
Field structure for reference (do NOT return any field as a stringified blob):
{json.dumps(LESSON_SCHEMA, indent=2)}

"framework" MUST be an array of 5 separate phase objects (Predict, Observe, Explain,
DQB/Driving Question Board, Model Building) — each an object with phase, learnerExperience,
teacherMoves, sensemakingStrategy, formativeAssessment. Never return framework as a string.
Set "number" to {num}.
Set "substrand" to "Sub-Strand {args.substrand}: {args.substrand_name}".
"""

    result = call_claude(prompt, max_tokens=8192, schema=LESSON_TOOL_SCHEMA)
    if result and 'lesson' in result:
        return result['lesson']
    return result   # in case Claude returns the lesson directly


def _get_lesson_storyline(unit: dict, num: int) -> str:
    """Extract this lesson's storyline description from unit data."""
    storyline = unit.get('storylineThread') or unit.get('storyline', '')
    if not storyline:
        return ''
    lines = [l for l in storyline.split('\n') if l.strip()]
    for line in lines:
        if line.strip().startswith(f'Lesson {num}:'):
            return line.strip()
    return lines[num - 1] if num <= len(lines) else ''


# ── Final Explanation & Summary Table generation ──────────────────────────────

def generate_final_explanation(curriculum_text: str, unit: dict,
                                lessons: list, args,
                                fe_template: dict | None = None) -> dict | None:
    """Generate Final Explanation document data."""
    print("  Generating Final Explanation...")

    lesson_titles = "\n".join(f"  Lesson {l['number']}: {l['title']}" for l in lessons)

    fe_template_section = ""
    if fe_template and fe_template.get('paragraphs'):
        fe_template_section = (
            "\nTEACHER TEMPLATE — FINAL EXPLANATION REFERENCE "
            "(use as structural and content guidance; improve quality where possible):\n"
            + "\n".join(fe_template['paragraphs'][:40])
        )

    prompt = f"""AUDIENCE: This Final Explanation document is for STUDENTS to use as a model for answering the driving question. Use clear, student-accessible language. Write in second person when appropriate.

Generate a Final Explanation assessment document for:
Subject: {args.subject.capitalize()} Grade 10
Sub-strand: {args.substrand} {args.substrand_name}
Driving question: {unit.get('drivingQuestion', '')}
Phenomenon: {unit.get('phenomenon', '')}

Lesson sequence completed:
{lesson_titles}
{fe_template_section}

Return ONLY this JSON:
{{
  "subjectLabel": "{args.substrand_name.upper()}",
  "instructions": "Multi-sentence instructions for students including what to use (Summary Table, experiments conducted, models), what to include (phenomenon connection, Kenya examples, evidence), and word count requirement (minimum 300 words)",
  "sections": [
    {{
      "title": "SECTION 1: [TITLE IN CAPS]",
      "prompt": "Specific prompt telling student what to explain in this section",
      "exemplar": "2-3 paragraph model answer demonstrating expected quality, using Kenyan contexts"
    }}
  ],
  "rubric": [
    {{
      "criterion": "Criterion name",
      "excellent": "Excellent (4) descriptor",
      "proficient": "Proficient (3) descriptor",
      "developing": "Developing (1-2) descriptor"
    }}
  ]
}}

Include 4-5 sections covering the main content areas.
Include 4-5 rubric criteria covering key scientific concepts.
"""

    return call_claude(prompt, max_tokens=8192, schema=FE_TOOL_SCHEMA)


def generate_summary_table(unit: dict, lessons: list, args,
                            st_template: dict | None = None) -> dict | None:
    """Generate teacher reference Summary Table data."""
    print("  Generating Summary Table...")

    # Pre-compute outside f-string to avoid dict/brace escaping issues
    lesson_refs = json.dumps([
        {
            "number": l["number"],
            "title": l["title"],
            "summary": l.get("summaryTablePrompt", {}).get("learned", "")[:150]
        }
        for l in lessons
    ], indent=2, ensure_ascii=False)

    dq = unit.get("drivingQuestion", "").split("\n")[0]
    ss = f"Sub-Strand {args.substrand}: {args.substrand_name}"

    st_template_section = ""
    if st_template and st_template.get('paragraphs'):
        st_template_section = (
            "\nTEACHER TEMPLATE — SUMMARY TABLE REFERENCE "
            "(use as structural and content guidance; improve quality where possible):\n"
            + "\n".join(st_template['paragraphs'][:30])
            + "\n"
        )

    prompt = (
        f"AUDIENCE: This Summary Table is a TEACHER REFERENCE — a pre-filled answer key "
        f"for use while teaching. Write in third-person teacher voice; include pedagogical "
        f"notes a teacher would find helpful.\n\n"
        f"Generate a teacher reference Summary Table for:\n"
        f"Subject: {args.subject.capitalize()} Grade 10\n"
        f"Sub-strand: {ss}\n"
        f"Driving question: {dq}\n"
        f"{st_template_section}\n"
        f"Return ONLY this JSON (no other text):\n"
        f'{{ "subStrand": "{ss}", "drivingQuestion": "{dq}", "lessons": ['
        f' {{ "number": 1, "title": "title", "observed": "...", "learned": "...", "explained": "..." }}'
        f' ] }}\n\n'
        f"Generate one entry per lesson for all {args.lessons} lessons.\n"
        f"Lesson titles and summaries for reference:\n"
        f"{lesson_refs}\n"
    )

    return call_claude(prompt, max_tokens=8192, schema=ST_TOOL_SCHEMA)


# ── Data file writer ──────────────────────────────────────────────────────────

# Version of ares-contract.schema.json that generated files declare conformance to.
SCHEMA_VERSION = '1.0.0'


def write_data_file(output_name: str, meta: dict, unit: dict,
                    lessons: list, fe: dict | None, st: dict | None):
    """Write the assembled data file to generators/data/."""
    output_path = PROJECT_ROOT / 'generators' / 'data' / f'{output_name}_data.js'

    # ── Normalise to ares-contract canonical shape (defensive: covers model
    #    wobble and applies to both sync and batch-collect callers) ──────────
    if 'storyline' in unit and not unit.get('storylineThread'):
        unit['storylineThread'] = unit.pop('storyline')
    else:
        unit.pop('storyline', None)
    if not meta.get('substrand_id') or not meta.get('substrand_name'):
        _sm = re.match(r'\s*Sub-Strand\s+([\d.]+):\s*(.+)', unit.get('substrand', ''))
        if _sm:
            meta.setdefault('substrand_id', _sm.group(1))
            meta.setdefault('substrand_name', _sm.group(2).strip())
    def _coerce_int(v):
        """Contract types lesson 'number' as integer; coerce stringified ints."""
        if isinstance(v, bool):
            return v
        if isinstance(v, str) and v.strip().lstrip('-').isdigit():
            return int(v.strip())
        return v
    for _lesson in lessons:
        if isinstance(_lesson, dict) and 'safetyNotes' in _lesson:
            _sn = _lesson.pop('safetyNotes')
            _slo = _lesson.get('slo')
            if isinstance(_slo, dict) and not _slo.get('safetyNotes'):
                _slo['safetyNotes'] = _sn
        if isinstance(_lesson, dict) and 'number' in _lesson:
            _lesson['number'] = _coerce_int(_lesson['number'])
    if isinstance(st, dict) and isinstance(st.get('lessons'), list):
        for _stl in st['lessons']:
            if isinstance(_stl, dict) and 'number' in _stl:
                _stl['number'] = _coerce_int(_stl['number'])

    def js_val(obj):
        """Convert Python object to JS-compatible string."""
        return json.dumps(obj, indent=2, ensure_ascii=False)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(f"'use strict';\n")
        f.write(f"/**\n")
        f.write(f" * {output_name}_data.js\n")
        f.write(f" * Generated by generate_substrand.py\n")
        f.write(f" * Run: node generators/generate.js {output_name}\n")
        f.write(f" */\n\n")

        # META
        f.write(f"const META = {js_val(meta)};\n\n")

        # UNIT
        f.write(f"const UNIT = {js_val(unit)};\n\n")

        # LESSONS
        f.write(f"const LESSONS = {js_val(lessons)};\n\n")

        # FINAL_EXPLANATION
        if fe:
            f.write(f"const FINAL_EXPLANATION = {js_val(fe)};\n\n")
        else:
            f.write(f"const FINAL_EXPLANATION = null;\n\n")

        # SUMMARY_TABLE
        if st:
            f.write(f"const SUMMARY_TABLE = {js_val(st)};\n\n")
        else:
            f.write(f"const SUMMARY_TABLE = null;\n\n")

        f.write(f"const schemaVersion = '{SCHEMA_VERSION}';\n\n")
        f.write(f"module.exports = {{ schemaVersion, META, UNIT, LESSONS, FINAL_EXPLANATION, SUMMARY_TABLE }};\n")

    print(f"  Written: {output_path}")
    return output_path


# ── Subject/substrand lookup ──────────────────────────────────────────────────

SUBSTRAND_NAMES = {
    'biology': {
        '1.1': 'Cell Structure',
        '1.2': 'Chemicals of Life',
        '1.3': 'Cell Biology',
        '2.1': 'Plant Nutrition',
        '2.2': 'Plant Transport',
        '2.3': 'Plant Gaseous Exchange and Respiration',
        '3.1': 'Animal Nutrition',
        '3.2': 'Animal Transport',
        '3.3': 'Animal Gaseous Exchange and Respiration',
    },
    'chemistry': {
        '1.1': 'Introduction to Chemistry',
        '1.2': 'The Atom',
        '1.3': 'The Periodic Table',
        '1.4': 'Chemical Bonding',
        '1.5': 'Periodicity',
        '2.1': 'Introduction to Salts',
        '3.1': 'Acids and Bases',
    },
    'physics': {
        '1.1': 'Pressure',
        '1.2': 'Mechanical Properties of Materials',
        '1.3': 'Temperature and Thermal Expansion',
        '1.4': 'Energy, Work, Power and Machines',
        '1.5': 'Moments of Equilibrium',
        '2.1': 'Properties of Waves',
        '3.1': 'Radioactivity and Stability of Isotopes',
        '3.2': 'Current Electricity',
        '3.3': 'Introduction to Electronics',
        '3.4': 'Electrostatics',
        '4.1': 'Greenhouse Effect and Climate Change',
        '4.2': 'Introduction to Space Physics',
    },
    'mathematics': {
        '1.1': 'Real Numbers',
        '1.2': 'Indices',
        '1.3': 'Quadratic Equations',
        '1.4': 'Congruence',
        '2.1': 'Similarity and Enlargement',
        '2.2': 'Area of Polygons',
        '2.3': 'Area of Part of a Circle',
        '2.4': 'Surface Area and Volume of Solids',
        '3.1': 'Trigonometry I',
        '3.2': 'Rotation',
        '3.3': 'Vectors I',
        '3.4': 'Linear Motion',
        '4.1': 'Statistics I',
        '4.2': 'Probability I',
    },
}

CURRICULUM_PDF_MAP = {
    'biology':     'data/raw/curriculum_pdfs/Grade10_Biology_CBE_Curriculum.pdf',
    'chemistry':   'data/raw/curriculum_pdfs/Grade10_Chemistry_CBE_Curriculum.pdf',
    'physics':     'data/raw/curriculum_pdfs/Grade10_Physics_CBE_Curriculum.pdf',
    'mathematics': 'data/raw/curriculum_pdfs/Grade10_Mathematics_Curriculum.pdf',
}

LESSON_COUNTS = {
    'biology': {'1.1': 6, '1.2': 14, '1.3': 20, '1.4': 24,
                '2.1': 12, '2.2': 22, '2.3': 22,
                '3.1': 12, '3.2': 24, '3.3': 24},
}


# ── v2 template lookup ────────────────────────────────────────────────────────

V2_TEMPLATE_DIR = PROJECT_ROOT / 'data' / 'raw' / 'CBE LESSON TEMPLATES' / 'v2_owner_inventory'

_SUBJECT_FOLDER = {
    'biology': 'Biology', 'chemistry': 'Chemistry',
    'physics': 'Physics', 'mathematics': 'Maths',
}


def _v2_output_dir(subject: str, substrand_id: str, substrand_name: str) -> str:
    """Return the outputDir path segment for v2 output, e.g. 'v2/Biology/SS1.1_Cell_Structure'."""
    subj_folder = _SUBJECT_FOLDER.get(subject, subject.capitalize())
    ss_folder = f"SS{substrand_id}_{substrand_name.replace(' ', '_')}"
    return f"v2/{subj_folder}/{ss_folder}"


def find_v2_templates(subject: str, substrand_id: str) -> dict:
    """Find up to three template docx files for this sub-strand from v2_owner_inventory.
    Returns {'lesson': Path|None, 'fe': Path|None, 'st': Path|None}.
    """
    subject_dir = V2_TEMPLATE_DIR / _SUBJECT_FOLDER.get(subject, subject.capitalize())
    matches = list(subject_dir.glob(f'SS{substrand_id}_*'))
    if not matches:
        return {'lesson': None, 'fe': None, 'st': None}

    found = {'lesson': None, 'fe': None, 'st': None}
    for f in matches[0].glob('*.docx'):
        fname = f.name.lower().replace('_', '').replace('-', '').replace(' ', '')
        if 'resourceguide' in fname:
            continue
        if 'finalexplanation' in fname:
            found['fe'] = f
        elif 'summarytable' in fname or 'answerkey' in fname:
            found['st'] = f
        else:
            found['lesson'] = f
    return found


def determine_lesson_count(curriculum_text: str, lesson_template: dict,
                            substrand_id: str, subject: str) -> tuple:
    """Return (count, source) where source is 'template_regex'|'pre_pass'|'default'.
    Clamps result to [6, 14].
    """
    # 1. Regex on template paragraph text
    if lesson_template:
        all_text = ' '.join(lesson_template.get('paragraphs', []))
        m = (re.search(r'\b(1[0-4]|[6-9])\s+lessons?\b', all_text, re.IGNORECASE) or
             re.search(r'\b(1[0-4]|[6-9])\s+periods?\b', all_text, re.IGNORECASE))
        if m:
            return (int(m.group(1)), 'template_regex')

    # 2. Pre-pass API call — small, cached in batch checkpoint
    COUNT_SCHEMA = {
        'type': 'object', 'additionalProperties': False,
        'properties': {
            'lesson_count': {'type': 'integer', 'minimum': 6, 'maximum': 14},
        },
        'required': ['lesson_count'],
    }
    excerpt = curriculum_text[:2000] if curriculum_text else ''
    template_excerpt = ' '.join(lesson_template.get('paragraphs', [])[:10]) if lesson_template else ''
    prompt = (
        f"How many lessons (6–14) should Sub-Strand {substrand_id} of Grade 10 "
        f"{subject.capitalize()} have?\n\n"
        f"CURRICULUM EXCERPT:\n{excerpt}\n\n"
        f"TEMPLATE EXCERPT:\n{template_excerpt}\n\n"
        f"Consider topic complexity, number of learning outcomes, and any KICD guidance. "
        f"Target range 6–14, average 8–10."
    )
    result = call_claude(prompt, max_tokens=512, retries=2,
                         schema=COUNT_SCHEMA, tool_name='propose_lesson_count')
    if result and isinstance(result.get('lesson_count'), int):
        return (max(6, min(14, result['lesson_count'])), 'pre_pass')

    return (8, 'default')


# ── Main ──────────────────────────────────────────────────────────────────────

def _checkpoint_path(output_name: str) -> Path:
    return PROJECT_ROOT / 'generators' / 'data' / f'.{output_name}_checkpoint.json'

def _save_checkpoint(output_name: str, lessons: list):
    path = _checkpoint_path(output_name)
    with open(path, 'w') as f:
        json.dump(lessons, f, ensure_ascii=False)

def _load_checkpoint(output_name: str) -> list | None:
    path = _checkpoint_path(output_name)
    if path.exists():
        try:
            with open(path) as f:
                data = json.load(f)
            print(f"  Found checkpoint with {len(data)} lesson(s)")
            return data
        except Exception:
            return None
    return None



# ── Message Batches API ───────────────────────────────────────────────────────

def _batch_id_path(output_name: str) -> Path:
    return PROJECT_ROOT / 'generators' / 'data' / f'.{output_name}_batch_id.txt'


def submit_batch(requests: list) -> str:
    """Submit all requests as a single Message Batches API job.
    Returns the batch_id for later collection.
    """
    print(f"  Submitting batch of {len(requests)} requests...")
    batch = CLIENT.beta.messages.batches.create(
        requests=requests,
        betas=["output-300k-2026-03-24"],   # 300K output tokens per request
    )
    print(f"  Batch ID: {batch.id}")
    print(f"  Status:   {batch.processing_status}")
    return batch.id


def poll_batch(batch_id: str, poll_interval: int = 60) -> object:
    """Poll until the batch is complete. Returns the finished batch object."""
    import sys
    print(f"  Polling batch {batch_id} (checking every {poll_interval}s)...")
    while True:
        batch = CLIENT.beta.messages.batches.retrieve(batch_id)
        counts = batch.request_counts
        done   = counts.succeeded + counts.errored + counts.expired
        total  = counts.processing + done
        print(f"    {batch.processing_status}: {done}/{total} done "
              f"({counts.succeeded} succeeded, {counts.errored} errored)",
              end='\r', flush=True)
        if batch.processing_status == 'ended':
            print()  # newline after \r
            return batch
        time.sleep(poll_interval)


def collect_batch_results(batch_id: str) -> dict:
    """Retrieve all results from a completed batch.
    Returns dict mapping custom_id → parsed dict (or None on error).
    Handles both tool_use (schema-enforced) and free-text JSON responses.
    """
    results = {}
    for result in CLIENT.beta.messages.batches.results(batch_id):
        cid = result.custom_id
        if result.result.type != 'succeeded':
            print(f"  WARNING: {cid} failed: {result.result.type}")
            results[cid] = None
            continue

        content = result.result.message.content

        # Prefer tool_use block (schema-enforced path)
        tool_data = None
        for block in content:
            if getattr(block, 'type', None) == 'tool_use':
                tool_data = block.input
                break

        if tool_data is not None:
            schema = LESSON_TOOL_SCHEMA if cid.startswith('lesson_') else FE_TOOL_SCHEMA
            violations = _schema_violations(tool_data, schema)
            if violations:
                print(f"  WARNING: Schema violations for {cid}: {violations[:3]}")
                results[cid] = None
            else:
                results[cid] = tool_data
        else:
            # Fallback: free-text JSON (legacy batches submitted without tool_use)
            raw = content[0].text.strip() if content else ''
            raw = re.sub(r'^```(?:json)?\s*', '', raw)
            raw = re.sub(r'\s*```$', '', raw)
            raw = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', raw)
            try:
                results[cid] = json.loads(raw)
            except json.JSONDecodeError as e:
                print(f"  WARNING: JSON parse failed for {cid}: {e}")
                results[cid] = None
    return results


def build_batch_requests(curriculum_text: str, lesson_template: dict,
                          unit: dict, args,
                          fe_template: dict | None = None) -> list:
    """Build all lesson + FE requests for batch submission.
    Note: prev_summaries context is omitted in batch mode since all
    requests are submitted simultaneously. The unit storyline thread
    provides sufficient continuity.
    """
    requests = []

    for lesson_num in range(1, args.lessons + 1):
        lesson_pos = "opening lesson that launches the phenomenon" if lesson_num == 1 else \
                     "final lesson with Final Explanation" if lesson_num == args.lessons else \
                     f"lesson {lesson_num} of {args.lessons} in the evidence-gathering sequence"

        prompt = (
            f"Generate Lesson {lesson_num} ({lesson_pos}) for:\n"
            f"Subject: {args.subject.capitalize()} Grade 10\n"
            f"Sub-strand: {args.substrand} {args.substrand_name}\n"
            f"Driving question: {unit.get('drivingQuestion', '')}\n"
            f"Phenomenon: {unit.get('phenomenon', '')}\n"
            f"Storyline thread for this lesson: {_get_lesson_storyline(unit, lesson_num)}\n\n"
            f"KICD CURRICULUM CONTENT:\n{curriculum_text}\n\n"
            f"TEACHER TEMPLATE — EVIDENCE ACTIVITIES:\n"
            f"{lesson_template.get('evidence_activities', '')}\n\n"
            f"RULES:\n"
            f"- All learner experiences must connect back to the phenomenon\n"
            f"- Teacher moves must include specific quoted phrases, WAIT TIME (10-15 seconds), cold-call counts\n"
            f"- Sensemaking strategies: name the strategy and explain how it builds understanding\n"
            f"- Formative assessment: specific, observable indicators\n"
            f"- Kenya-relevant contexts throughout\n"
            f"- Lesson {lesson_num} should "
            f"{'introduce the phenomenon and open the DQB' if lesson_num == 1 else 'build on previous evidence and advance the driving question'}\n"
            f"{'- Final lesson: focus on Final Explanation, model comparison L1 vs final, DQB completion' if lesson_num == args.lessons else ''}\n\n"
            f"Use the emit_data tool to return the lesson. Set number to {lesson_num}. "
            f"Set substrand to 'Sub-Strand {args.substrand}: {args.substrand_name}'."
        )

        requests.append({
            "custom_id": f"lesson_{lesson_num}",
            "params": {
                "model": MODEL,
                "max_tokens": 8192,
                "system": SYSTEM_PROMPT,
                "messages": [{"role": "user", "content": prompt}],
                "tools": [{
                    "name": "emit_data",
                    "description": "Return the requested lesson-plan data as structured fields.",
                    "input_schema": LESSON_TOOL_SCHEMA,
                }],
                "tool_choice": {"type": "tool", "name": "emit_data"},
            },
        })

    # Final Explanation request
    fe_template_section = ""
    if fe_template and fe_template.get('paragraphs'):
        fe_template_section = (
            "\nTEACHER TEMPLATE — FINAL EXPLANATION REFERENCE "
            "(use as structural and content guidance; improve quality where possible):\n"
            + "\n".join(fe_template['paragraphs'][:40])
            + "\n"
        )

    fe_prompt = (
        f"AUDIENCE: This Final Explanation document is for STUDENTS to use as a model "
        f"for answering the driving question. Use clear, student-accessible language. "
        f"Write in second person when appropriate.\n\n"
        f"Generate a Final Explanation assessment document for:\n"
        f"Subject: {args.subject.capitalize()} Grade 10\n"
        f"Sub-strand: {args.substrand} {args.substrand_name}\n"
        f"Driving question: {unit.get('drivingQuestion', '')}\n"
        f"Phenomenon: {unit.get('phenomenon', '')}\n"
        f"Number of lessons in sequence: {args.lessons}\n"
        f"{fe_template_section}\n"
        f"Use the emit_data tool to return the Final Explanation with 4-5 sections and 4-5 rubric criteria."
    )

    requests.append({
        "custom_id": "final_explanation",
        "params": {
            "model": MODEL,
            "max_tokens": 8192,
            "system": SYSTEM_PROMPT,
            "messages": [{"role": "user", "content": fe_prompt}],
            "tools": [{
                "name": "emit_data",
                "description": "Return the requested lesson-plan data as structured fields.",
                "input_schema": FE_TOOL_SCHEMA,
            }],
            "tool_choice": {"type": "tool", "name": "emit_data"},
        },
    })

    return requests


def run_collect(output_name: str, args):
    """Collect results from a previously submitted batch and assemble data file."""
    id_path = _batch_id_path(output_name)
    if not id_path.exists():
        print(f"ERROR: No batch ID found for '{output_name}'")
        print(f"  Expected: {id_path}")
        print(f"  Submit a batch first with: --batch")
        sys.exit(1)

    batch_id = id_path.read_text().strip()
    batch_meta_path = id_path.with_suffix('.json')
    batch_meta = json.loads(batch_meta_path.read_text()) if batch_meta_path.exists() else {}

    print(f"Collecting batch: {batch_id}")

    # Check status
    batch = CLIENT.beta.messages.batches.retrieve(batch_id)
    print(f"  Status: {batch.processing_status}")

    if batch.processing_status != 'ended':
        if args.wait:
            batch = poll_batch(batch_id)
        else:
            counts = batch.request_counts
            print(f"  Not ready yet ({counts.processing} still processing)")
            print(f"  Run again with --wait to poll, or check back later")
            return

    # Collect results
    print("  Collecting results...")
    results = collect_batch_results(batch_id)
    print(f"  Got {len(results)} results ({sum(1 for v in results.values() if v)} succeeded)")

    # Reconstruct unit from saved metadata
    unit = batch_meta.get('unit', {})
    meta = batch_meta.get('meta', {})
    n_lessons = batch_meta.get('lessons', args.lessons if hasattr(args, 'lessons') else 6)

    # Extract lessons
    lessons = []
    for i in range(1, n_lessons + 1):
        cid  = f"lesson_{i}"
        data = results.get(cid)
        if data:
            lesson = data.get('lesson', data)   # handle both {lesson: {...}} and {...}
            lessons.append(lesson)
        else:
            print(f"  WARNING: Lesson {i} missing or failed — using stub")
            lessons.append({
                "number": i, "title": f"Lesson {i}",
                "duration": "40 minutes",
                "substrand": f"Sub-Strand {meta.get('substrand_id', '')}: {meta.get('substrand_name', '')}",
                "aresKeywords": "",
                "slo": {"purpose":"","knowledge":"","skills":"","attitudes":"",
                        "keyInquiry":"","purposeInStoryline":"","safetyNotes":""},
                "overview": "",
                "framework": [
                    {"phase": ph, "learnerExperience":"","teacherMoves":"",
                     "sensemakingStrategy":"","formativeAssessment":""}
                    for ph in ["Predict Phase","Observe Phase","Explain Phase",
                               "Driving Question Board (DQB) Creation","Model Building Phase"]
                ],
                "teacherReflection": "",
                "summaryTablePrompt": {"observed":"","learned":"","explained":""},
            })

    # Extract FE
    fe = results.get('final_explanation')
    if not fe:
        print("  WARNING: Final Explanation missing or failed")

    # Generate Summary Table synchronously (needs real lesson data)
    print("  Generating Summary Table (synchronous — needs lesson data)...")

    # Build args-like namespace for generate_summary_table
    class _Args:
        pass
    st_args = _Args()
    st_args.subject     = meta.get('subject', 'Biology').lower()
    st_args.substrand   = meta.get('substrand_id', '')
    st_args.substrand_name = meta.get('substrand_name', '')
    st_args.lessons     = n_lessons

    # Re-load ST template (needed for audience-aware generation)
    _st_path = find_v2_templates(st_args.subject, st_args.substrand).get('st')
    st_template = extract_template_docx(str(_st_path)) if _st_path else None

    st = generate_summary_table(unit, lessons, st_args, st_template=st_template)
    if st:
        print("  Summary Table generated ✓")

    # Write data file
    print("\n  Writing data file...")
    output_path = write_data_file(output_name, meta, unit, lessons, fe, st)

    # Optionally run generator
    if hasattr(args, 'run') and args.run:
        print("\n  Running generator...")
        import subprocess
        result = subprocess.run(
            ['node', 'generators/generate.js', output_name],
            cwd=PROJECT_ROOT, capture_output=True, text=True,
        )
        print(result.stdout)
        if result.returncode != 0:
            print("Generator errors:", result.stderr)

    # Clean up batch ID file
    id_path.unlink(missing_ok=True)
    batch_meta_path.unlink(missing_ok=True)

    print(f"\n✓ Done! Data file: {output_path}")

def main():
    parser = argparse.ArgumentParser(description='Generate CBE sub-strand content via Claude API')
    parser.add_argument('--subject',   required=False, default=None,
                        choices=['biology', 'chemistry', 'physics', 'mathematics'])
    parser.add_argument('--substrand', required=False, default=None,
                        help='Sub-strand number e.g. 1.4')
    parser.add_argument('--output',    required=False, default=None,
                        help='Output name e.g. bio_1_4')
    parser.add_argument('--lessons',   type=int, default=None,
                        help='Number of lessons (default: from KICD curriculum)')
    parser.add_argument('--template',  default=None,
                        help='Path to teacher template docx (auto-detected if omitted)')
    parser.add_argument('--batch',     action='store_true',
                        help='Submit all content as a Message Batches API job (50%% cheaper)')
    parser.add_argument('--collect',   default=None, metavar='NAME',
                        help='Collect results from a previously submitted batch')
    parser.add_argument('--wait',      action='store_true',
                        help='With --collect: poll until batch is complete')
    parser.add_argument('--resume',    action='store_true',
                        help='Resume from checkpoint if interrupted')
    parser.add_argument('--run',       action='store_true',
                        help='Run node generators/generate.js after creating data file')
    parser.add_argument('--unit-only', action='store_true',
                        help='Generate UNIT only (for testing)')
    args = parser.parse_args()

    # Validate required args when not in collect mode
    if not args.collect:
        missing = [f for f, v in [('--subject', args.subject),
                                   ('--substrand', args.substrand),
                                   ('--output', args.output)] if v is None]
        if missing:
            parser.error(f"the following arguments are required: {', '.join(missing)}")

    # ── Batch collect mode ────────────────────────────────────────────────────
    if args.collect:
        run_collect(args.collect, args)
        return

    # Resolve substrand name
    args.substrand_name = SUBSTRAND_NAMES.get(args.subject, {}).get(
        args.substrand, f'Sub-Strand {args.substrand}'
    )

    # ── Extract source content + templates ───────────────────────────────────

    print("\n1. Extracting source content...")
    curriculum_pdf  = str(PROJECT_ROOT / CURRICULUM_PDF_MAP.get(args.subject, ''))
    curriculum_text = extract_curriculum_pdf(curriculum_pdf, args.substrand)
    print(f"  Curriculum text: {len(curriculum_text)} chars")

    # Find v2 templates (lesson + FE + ST); --template CLI arg overrides lesson only
    _v2 = find_v2_templates(args.subject, args.substrand)
    lesson_template_path = Path(args.template) if args.template else _v2['lesson']
    fe_template_path     = _v2['fe']
    st_template_path     = _v2['st']

    lesson_template = extract_template_docx(str(lesson_template_path)) if lesson_template_path else {}
    fe_template     = extract_template_docx(str(fe_template_path)) if fe_template_path else None
    st_template     = extract_template_docx(str(st_template_path)) if st_template_path else None

    if lesson_template_path:
        print(f"  Lesson template: {lesson_template_path.name} "
              f"({len(lesson_template.get('paragraphs', []))} paragraphs)")
    if fe_template_path:
        print(f"  FE template:     {fe_template_path.name}")
    if st_template_path:
        print(f"  ST template:     {st_template_path.name}")

    # Resolve lesson count (CLI --lessons wins; otherwise probe template/curriculum)
    lesson_count_source = 'cli'
    if args.lessons is None:
        args.lessons, lesson_count_source = determine_lesson_count(
            curriculum_text, lesson_template, args.substrand, args.subject)
        print(f"  Lesson count: {args.lessons} (source: {lesson_count_source})")

    print(f"\nGenerating: {args.subject.capitalize()} Grade 10 Sub-Strand {args.substrand}: {args.substrand_name}")
    print(f"  Lessons: {args.lessons}")
    print(f"  Output:  generators/data/{args.output}_data.js")
    print(f"  Model:   {MODEL}")

    # ── Generate content via Claude API ──────────────────────────────────────

    print("\n2. Generating content via Claude API...")

    # UNIT
    unit = generate_unit(curriculum_text, lesson_template, args)
    if not unit:
        print("ERROR: Failed to generate UNIT data")
        sys.exit(1)
    print(f"  UNIT generated ✓")

    if args.unit_only:
        print("\n[unit-only mode — stopping here]")
        print(json.dumps(unit, indent=2))
        return

    # ── Batch submit mode ─────────────────────────────────────────────────────
    if args.batch:
        print("\n2b. Building batch requests...")
        requests = build_batch_requests(curriculum_text, lesson_template, unit, args,
                                         fe_template=fe_template)
        print(f"  Built {len(requests)} requests "
              f"({args.lessons} lessons + 1 Final Explanation)")

        batch_id = submit_batch(requests)

        # Save batch ID and metadata for later collection
        id_path = _batch_id_path(args.output)
        id_path.write_text(batch_id)
        meta_path = id_path.with_suffix('.json')
        _subj_cap = args.subject.capitalize()
        meta_path.write_text(json.dumps({
            "unit": unit,
            "meta": {
                "subject":        _subj_cap,
                "grade":          10,
                "substrand_id":   args.substrand,
                "substrand_name": args.substrand_name,
                "outputDir":      _v2_output_dir(args.subject, args.substrand, args.substrand_name),
                "filePrefix":     f"{_subj_cap}_{args.substrand_name.replace(' ', '_')}",
                "titleDoc":       f"{_subj_cap.upper()} GRADE 10: {args.substrand_name.upper()}",
                "subtitleDoc":    f"CBE Phenomenon-Driven Lesson Sequence — Sub-Strand {args.substrand} ({args.lessons} Lessons)",
                "col3Label":      "Teacher Moves",
                "col5Label":      "Formative Assessment Strategy",
            },
            "lessons": args.lessons,
            "lesson_count_source": lesson_count_source,
        }, ensure_ascii=False, indent=2))

        print(f"\n✓ Batch submitted!")
        print(f"  Batch ID saved to: {id_path}")
        print(f"  Collect results when ready:")
        print(f"    python3 src/generate_substrand.py --collect {args.output} --run")
        print(f"  Or poll automatically:")
        print(f"    python3 src/generate_substrand.py --collect {args.output} --wait --run")
        return

    # LESSONS — with checkpoint resume support
    lessons        = []
    prev_summaries = []
    start_lesson   = 1

    if args.resume:
        saved = _load_checkpoint(args.output)
        if saved:
            lessons        = saved
            prev_summaries = [{"number": l["number"], "title": l["title"],
                                "summary": l.get("summaryTablePrompt", {}).get("learned", "")[:150]}
                               for l in lessons]
            start_lesson   = len(lessons) + 1
            print(f"  Resuming from lesson {start_lesson} ({len(lessons)} already done)")

    for lesson_num in range(start_lesson, args.lessons + 1):
        lesson = generate_lesson(lesson_num, curriculum_text, lesson_template,
                                  unit, prev_summaries, args)
        if not lesson:
            print(f"  WARNING: Failed to generate Lesson {lesson_num} — using stub")
            lesson = {
                "number": lesson_num,
                "title": f"Lesson {lesson_num}",
                "duration": "40 minutes",
                "substrand": f"Sub-Strand {args.substrand}: {args.substrand_name}",
                "aresKeywords": args.substrand_name.lower(),
                "slo": {"purpose": "", "knowledge": "", "skills": "", "attitudes": "",
                        "keyInquiry": "", "purposeInStoryline": "", "safetyNotes": ""},
                "overview": "",
                "framework": [
                    {"phase": ph, "learnerExperience": "", "teacherMoves": "",
                     "sensemakingStrategy": "", "formativeAssessment": ""}
                    for ph in ["Predict Phase", "Observe Phase", "Explain Phase",
                                "Driving Question Board (DQB) Creation", "Model Building Phase"]
                ],
                "teacherReflection": "",
                "summaryTablePrompt": {"observed": "", "learned": "", "explained": ""},
            }
        lessons.append(lesson)

        # Save checkpoint after each lesson so we can resume if interrupted
        _save_checkpoint(args.output, lessons)

        # Track lesson summary for next lesson's context
        prev_summaries.append({
            "number": lesson_num,
            "title": lesson.get("title", f"Lesson {lesson_num}"),
            "summary": lesson.get("summaryTablePrompt", {}).get("learned", "")[:150],
        })

        # Brief pause to avoid rate limits
        time.sleep(1)

    # FINAL EXPLANATION
    fe = generate_final_explanation(curriculum_text, unit, lessons, args,
                                     fe_template=fe_template)
    if fe:
        print("  Final Explanation generated ✓")
    else:
        print("  WARNING: Final Explanation generation failed — using null")

    # SUMMARY TABLE
    st = generate_summary_table(unit, lessons, args, st_template=st_template)
    if st:
        print("  Summary Table generated ✓")
    else:
        print("  WARNING: Summary Table generation failed — using null")

    # ── Build META ─────────────────────────────────────────────────────────────

    subject_cap = args.subject.capitalize()

    meta = {
        "subject":     subject_cap,
        "grade":       10,
        "substrand_id":   args.substrand,
        "substrand_name": args.substrand_name,
        "outputDir":   _v2_output_dir(args.subject, args.substrand, args.substrand_name),
        "filePrefix":  f"{subject_cap}_{args.substrand_name.replace(' ', '_')}",
        "titleDoc":    f"{subject_cap.upper()} GRADE 10: {args.substrand_name.upper()}",
        "subtitleDoc": f"CBE Phenomenon-Driven Lesson Sequence — Sub-Strand {args.substrand} ({args.lessons} Lessons)",
        "col3Label":   "Teacher Moves",
        "col5Label":   "Formative Assessment Strategy",
    }

    # ── Write data file ─────────────────────────────────────────────────────

    print("\n3. Writing data file...")
    output_path = write_data_file(args.output, meta, unit, lessons, fe, st)

    # ── Optionally run generator ─────────────────────────────────────────────

    if args.run:
        print(f"\n4. Running generator...")
        import subprocess
        result = subprocess.run(
            ['node', 'generators/generate.js', args.output],
            cwd=PROJECT_ROOT,
            capture_output=True, text=True,
        )
        print(result.stdout)
        if result.returncode != 0:
            print("Generator errors:", result.stderr)

    print(f"\n✓ Done!")
    print(f"  Data file: {output_path}")
    print(f"  To generate docx: node generators/generate.js {args.output}")


if __name__ == '__main__':
    main()
