"""
Simple DOCX Lesson Generator - No fancy formatting, just clean tables
"""
import os
import sys
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches

sys.path.insert(0, str(Path(__file__).parent))

load_dotenv()


class SimpleLessonGenerator:
    """Generate lessons in simple Word format"""
    
    def __init__(self):
        print("Initializing Simple DOCX Generator...")
        import anthropic
        self.claude = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        print("✓ Ready")
    
    def generate_content(self, subject, grade, topic, number, phenomenon=None):
        """Generate lesson content"""
        print(f"\nGenerating Lesson {number}: {topic}")
        
        prompt = f"""Generate a CBE lesson plan for Kenya:
- Subject: {subject}, Grade: {grade}
- Topic: {topic}
- Phenomenon: {phenomenon or 'N/A'}

Format as plain text with clear sections:

SECTION A: LEARNING OUTCOMES
Knowledge: [3-4 points]
Skills: [3-4 points]
Attitudes: [2-3 points]
Practices: [2-3 points]

SECTION B: OVERVIEW
Key Question: [question]
Purpose: [purpose]
Safety: [notes or n/a]

SECTION C: LESSON ACTIVITIES (3-5 activities)
Activity 1:
- What students do: [description]
- Resources: [list, include ARES placeholders]
- Teacher facilitation: [questions, wait time]
- Student thinking: [Think-Pair-Share, etc.]
- Assessment: [checkboxes, quick checks]

Activity 2:
[Same format]

Activity 3:
[Same format]

SECTION D: REFLECTION
1. [Question 1]
2. [Question 2]
3. [Question 3]

SECTION E: SUMMARY PROMPT
[Prompt for students' summary table]

Use Kenyan context (KSh, KPLC). Include ARES resources. Use Think-Pair-Share."""
        
        response = self.claude.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=6000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        print("✓ Content generated")
        return response.content[0].text
    
    def create_docx(self, content, number, topic, output_dir="data/outputs/docx"):
        """Create simple Word document"""
        print("→ Creating Word document...")
        
        doc = Document()
        
        # Title
        doc.add_heading(f'Lesson {number}: {topic}', 0)
        doc.add_paragraph('Kenya Competency-Based Curriculum')
        doc.add_paragraph()
        
        # Add content sections
        lines = content.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('SECTION'):
                doc.add_heading(line.replace('SECTION ', ''), 2)
                current_section = line
            elif line.startswith('Activity '):
                doc.add_heading(line, 3)
            elif line.startswith('- ') or line.startswith('• '):
                doc.add_paragraph(line, style='List Bullet')
            elif line[0].isdigit() and '. ' in line:
                doc.add_paragraph(line, style='List Number')
            elif ':' in line and len(line.split(':')[0].split()) <= 4:
                # Likely a label
                p = doc.add_paragraph()
                parts = line.split(':', 1)
                p.add_run(parts[0] + ':').bold = True
                if len(parts) > 1:
                    p.add_run(' ' + parts[1])
            else:
                doc.add_paragraph(line)
        
        # Save
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        filename = f"Lesson_{number:03d}_{topic.replace(' ', '_')}.docx"
        filepath = Path(output_dir) / filename
        
        doc.save(str(filepath))
        print(f"✓ Saved: {filepath}")
        return str(filepath)
    
    def generate_lesson(self, subject, grade, topic, number=1, phenomenon=None):
        """Generate complete lesson"""
        content = self.generate_content(subject, grade, topic, number, phenomenon)
        filepath = self.create_docx(content, number, topic)
        return {'filepath': filepath, 'status': 'success'}


def main():
    print("\n" + "="*70)
    print("SIMPLE CBE LESSON GENERATOR - WORD FORMAT")
    print("="*70)
    
    gen = SimpleLessonGenerator()
    
    result = gen.generate_lesson(
        subject="Physics",
        grade="Grade 10",
        topic="Introduction to Current Electricity",
        number=1,
        phenomenon="Why do electrical devices get hot?"
    )
    
    print(f"\n✅ SUCCESS! Saved to: {result['filepath']}")
    print("\nOpen in Microsoft Word to view and format as needed.")


if __name__ == "__main__":
    main()
